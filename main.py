import telegram
from telegram.ext import Updater, CommandHandler, MessageHandler, Filters, CallbackContext, CallbackQueryHandler
from telegram import InlineKeyboardButton, InlineKeyboardMarkup, Update, ReplyKeyboardRemove
from telegram import ReplyKeyboardMarkup, KeyboardButton
from telegram import ChatAction
from flask import Flask, request
from collections import defaultdict
# from docx import Document  # moved to lazy import
# from docx.shared import Pt  # moved to lazy import
# from docx.oxml.ns import qn  # moved to lazy import
# from docx.enum.text import WD_PARAGRAPH_ALIGNMENT  # moved to lazy import
import os
import re
import uuid
from groq import Groq
import re
from collections import defaultdict
from telegram.ext import Dispatcher
from telegram import ReplyKeyboardRemove
from urllib.parse import quote
import time
import re
from typing import Dict, Any, List

TOKEN = os.getenv("BOT_TOKEN")
if not TOKEN:
    raise RuntimeError("BOT_TOKEN environment variable is not set")
bot = telegram.Bot(token=TOKEN)

app = Flask(__name__)

@app.route("/_health", methods=["GET"])
def _health():
    return "ok", 200


user_data = {}

# متن دکمه  ها

BACK_BTN = "⬅️ بازگشت"
AI_RESUME   = "AI_RESUME"   # کال‌بک دکمه‌ی بازگشت از AI
AI_ASK_TEXT = "❓ سؤال دارم"
AI_Q_LIMIT = 5



# --- Contact Config (ویرایش کن) ---
CONTACT_MOBILE_IR = "09128687292"     # شماره موبایل برای تماس (فرمت داخلی ایران)
CONTACT_MOBILE_INTL = "989128687292"  # همان شماره ولی بدون صفر و با 98 برای واتساپ
DEFAULT_WHATSAPP_TEXT = "سلام، برای راهنمایی و ثبت صورتجلسه راهنمایی می‌خواستم."
THANKYOU_BRAND = "ثبت کوشا"           # نام برند شما

TTL_SECONDS = 7 * 24 * 60 * 60
PHONE_TTL_SECONDS = 7 * 24 * 3600
phones = {}        # chat_id -> {"phone": "+98912...", "ts": epoch}
phone_index = {}   # "+98912..." -> set(chat_id,...)

USER_PHONE: Dict[int, Dict[str, Any]] = {}      # chat_id -> {"phone": str, "saved_at": ts, "meta": {...}}
ACTIVITY_LOG: Dict[int, List[Dict[str, Any]]] = {}  # chat_id -> [{"ts": ts, "event": str, "meta": dict}, ...]

FA_TO_EN_DIGITS = str.maketrans("۰۱۲۳۴۵۶۷۸۹", "0123456789")

GROQ_MODEL_QUALITY = "llama-3.3-70b-versatile" # کیفیت بالاتر
GROQ_MODEL = GROQ_MODEL_QUALITY
GROQ_MODEL = os.getenv("GROQ_MODEL", "llama-3.3-70b-versatile")

TOPIC_EXTEND_ROLES = "تمدید سمت اعضا"


groq_client = Groq(api_key=os.environ.get("GROQ_API_KEY"))

def ask_groq(user_text: str, system_prompt: str = None, max_tokens: int = 1024) -> str:
    if system_prompt is None:
        system_prompt = (
            "You are an assistant answering in Persian (Farsi). "
            "متخصص قانون تجارت ایران و ثبت شرکت‌ها هستی. جواب‌ها کوتاه و کاربردی باشند."
        )

    resp = groq_client.chat.completions.create(
        model=GROQ_MODEL,  # همیشه 70B
        messages=[
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_text},
        ],
        temperature=0.2,
        max_tokens=max_tokens,
    )
    return resp.choices[0].message.content.strip()



# --- AI Landing Options (labels must match exactly) ---
AI_OPT_MINUTES   = "⚡️ ساخت انواع صورتجلسات در چند دقیقه"
AI_OPT_QA        = "💬 مشاوره مجازی قانون تجارت و ثبت شرکت"
AI_OPT_COMP_TYPE = "🏢 راهنمای انتخاب نوع شرکت"
AI_OPT_NAME      = "🧠 پیشنهاد هوشمند نام شرکت"
AI_OPT_CONTRACT  = "📝 تولید قرارداد آماده"
AI_OPT_FORMAL    = "✍️ تبدیل متن ساده به متن رسمی/حقوقی"


def back_keyboard():
    # import محلی تا وابستگی بالا ایجاد نشه
    from telegram import ReplyKeyboardMarkup, KeyboardButton
    rows = [[KeyboardButton(BACK_BTN)]]
    return ReplyKeyboardMarkup(rows, resize_keyboard=True, one_time_keyboard=False)

def ai_consult_keyboard():
    from telegram import ReplyKeyboardMarkup, KeyboardButton
    rows = [[KeyboardButton(AI_BACK_TO_MENU)]]
    return ReplyKeyboardMarkup(rows, resize_keyboard=True)



def ai_services_keyboard():
    from telegram import ReplyKeyboardMarkup, KeyboardButton
    rows = [
        [KeyboardButton(AI_OPT_MINUTES)],
        [KeyboardButton(AI_OPT_QA), KeyboardButton(AI_OPT_COMP_TYPE)],
        [KeyboardButton(AI_OPT_NAME), KeyboardButton(AI_OPT_CONTRACT)],
        [KeyboardButton(AI_OPT_FORMAL)],
    ]
    # اگر بک دکمهٔ سراسری داری، می‌تونی اینجا هم اضافه‌اش کنی
    return ReplyKeyboardMarkup(rows, resize_keyboard=True, one_time_keyboard=False)

def send_ai_services_menu(chat_id, context):
    context.bot.send_message(
        chat_id=chat_id,
        text="🤖 لطفاً یکی از خدمات زیر را انتخاب کنید:",
        reply_markup=ai_services_keyboard()
    )



# تابع ساخت کیبورد اصلی که فقط دکمه بازگشت داره
def main_keyboard():
    return ReplyKeyboardMarkup(
        [[KeyboardButton(AI_ASK_TEXT), KeyboardButton(BACK_BTN)]],
        resize_keyboard=True,
        one_time_keyboard=False
    )

def base_reply_keyboard():
    return ReplyKeyboardMarkup(
        [["🔙 بازگشت به ادامه مراحل"]],
        resize_keyboard=True
    )
    
fields = [
    "نوع شرکت", "نام شرکت", "شماره ثبت", "شناسه ملی", "سرمایه", "تاریخ", "ساعت",
    "مدیر عامل", "نایب رییس", "رییس", "منشی", "آدرس جدید", "کد پستی", "وکیل"
]


persian_number_fields = ["شماره ثبت", "شناسه ملی", "سرمایه", "کد پستی"]



NEWSPAPERS = [
    "اطلاعات","ایران","شرق","جمهوری اسلامی","همشهری",
    "آفتاب یزد","کیهان","اعتماد","دنیای اقتصاد","فرهیختگان",
    "جهان صنعت","خراسان","گل","هفت صبح","جوان",
    "جهان اقتصاد","قدس","فرصت","آرمان امروز"
]




def small_keyboard():
    # انتخاب سریع تعداد شرکا
    return ReplyKeyboardMarkup([["1 نفر","2 نفر","3 نفر","4 نفر یا بیشتر"], [BACK_BTN]], resize_keyboard=True)

def yes_no_keyboard():
    return ReplyKeyboardMarkup([["بله ✅","خیر ❌"], [BACK_BTN]], resize_keyboard=True)

def board_need_keyboard():
    return ReplyKeyboardMarkup([["هیئت‌مدیره 3 نفره می‌خواهم","مدیریت ساده/یک‌نفره کافیست"], [BACK_BTN]], resize_keyboard=True)

def transfer_need_keyboard():
    return ReplyKeyboardMarkup([["خیلی مهم است","اهمیت متوسط","اهمیت ندارد"], [BACK_BTN]], resize_keyboard=True)
def is_persian_number(text):
    return all('۰' <= ch <= '۹' or ch.isspace() for ch in text)





# تبدیل اعداد فارسی به انگلیسی
def fa_to_en_number(text):
    table = str.maketrans('۰۱۲۳۴۵۶۷۸۹', '0123456789')
    return text.translate(table)


DOCX_IMPORTED = False
Document = Pt = qn = None

def _lazy_import_docx():
    global DOCX_IMPORTED, Document, Pt, qn, WD_PARAGRAPH_ALIGNMENT
    if DOCX_IMPORTED:
        return
    from docx import Document as _Document
    from docx.shared import Pt as _Pt
    from docx.oxml.ns import qn as _qn
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT as _WD
    Document, Pt, qn, WD_PARAGRAPH_ALIGNMENT = _Document, _Pt, _qn, _WD
    DOCX_IMPORTED = True

        # -------------------------------
        # توابع گرفتن شماره موبایل
        # -------------------------------

def fa_to_en(s: str) -> str:
    return (s or "").translate(FA_TO_EN_DIGITS)

def normalize_phone(s: str) -> str:
    s = fa_to_en_number(s or "")
    s = re.sub(r"\D+", "", s)           # فقط رقم
    if s.startswith("0"):               # 09... => 989...
        s = "98" + s[1:]
    if len(s) == 10 and s.startswith("9"):
        s = "98" + s
    if not s.startswith("98") and not s.startswith("+98"):
        # اگر کاربر فرمت دیگری داد، همان را نگه می‌داریم
        pass
    if not s.startswith("+"):
        s = "+" + s
    return s

def request_phone_keyboard():
    return ReplyKeyboardMarkup(
        [[KeyboardButton("📱 ارسال شماره موبایل", request_contact=True)],
         [KeyboardButton(BACK_BTN)]],
        resize_keyboard=True, one_time_keyboard=True
    )

def cleanup_phones():
    now = time.time()
    for cid, info in list(phones.items()):
        if now - info["ts"] > PHONE_TTL_SECONDS:
            phone_index.get(info["phone"], set()).discard(cid)
            phones.pop(cid, None)

def save_phone(chat_id: int, phone: str, context: CallbackContext):
    p = normalize_phone(phone)
    if not p:
        return

    cleanup_phones()
    phones[chat_id] = {"phone": p, "ts": int(time.time())}
    phone_index.setdefault(p, set()).add(chat_id)

    # در هر دو ساختار ذخیره شود
    USER_PHONE[chat_id] = {"phone": p, "saved_at": time.time(), "meta": {}}

    context.user_data["phone"] = p
    context.user_data.pop("awaiting", None)

    print("✅ phone saved for", chat_id, ":", p)

    context.bot.send_message(
        chat_id=chat_id,
        text=f"✅ شماره شما ثبت شد: {p}",
        reply_markup=main_keyboard()
    )


    print("DBG: save_phone called with", phone)
    print("DBG: USER_PHONE now:", USER_PHONE)



def normalize_phone(s: str) -> str:
    s = fa_to_en(s)
    s = re.sub(r"\D+", "", s)  # فقط رقم‌ها
    # پترن‌های قابل قبول: 09xxxxxxxxx یا 9xxxxxxxxx یا 989xxxxxxxxx یا +989xxxxxxxxx
    if s.startswith("0098"):
        s = s[4:]
    if s.startswith("98"):
        s = s[2:]
    if s.startswith("0"):
        s = s[1:]
    # حالا باید 10 رقمی و با 9 شروع شود
    if len(s) == 10 and s.startswith("9"):
        return "+989" + s[1:]
    return ""  # نامعتبر

def handle_contact(update: Update, context: CallbackContext):
    chat_id = update.effective_chat.id
    contact = update.message.contact
    if not contact or not contact.phone_number:
        context.bot.send_message(chat_id, "❗️شماره نامعتبر بود. دوباره دکمه «📱 ارسال شماره موبایل» را بزنید.")
        return

    save_phone(chat_id, contact.phone_number, context)

    # 👇 پاک کردن حالت انتظار شماره
    context.user_data["awaiting_phone"] = False
    context.user_data.pop("awaiting", None)

    # ✅ حالا منوی خدمات هوش مصنوعی را نشان بده
    user_data.setdefault(chat_id, {}).update({"step": 0, "onboarding_ai_shown": True})
    send_ai_services_menu(chat_id, context)



def is_valid_phone_text(s: str) -> bool:
    return bool(normalize_phone(s))

def set_user_phone(chat_id: int, phone_raw: str, meta: Dict[str, Any] | None = None) -> str:
    phone = normalize_phone(phone_raw)
    if not phone:
        return ""
    USER_PHONE[chat_id] = {
        "phone": phone,
        "saved_at": time.time(),
        "meta": meta or {}
    }
    prune_expired(chat_id)  # پاکسازی لاگ قدیمی همین کاربر
    return phone

def get_user_phone(chat_id: int) -> str:
    rec = USER_PHONE.get(chat_id)
    if rec and (time.time() - rec["saved_at"] <= TTL_SECONDS):
        return rec["phone"]
    # منقضی شده
    USER_PHONE.pop(chat_id, None)
    return ""

def log_activity(chat_id: int, event: str, meta: Dict[str, Any] | None = None) -> None:
    ACTIVITY_LOG.setdefault(chat_id, [])
    ACTIVITY_LOG[chat_id].append({
        "ts": time.time(),
        "event": event,
        "meta": meta or {}
    })
    prune_expired(chat_id)

def get_activity_last_week(chat_id: int) -> List[Dict[str, Any]]:
    now = time.time()
    return [e for e in ACTIVITY_LOG.get(chat_id, []) if now - e["ts"] <= TTL_SECONDS]

def prune_expired(chat_id: int | None = None) -> None:
    now = time.time()
    targets = [chat_id] if chat_id is not None else list(set(USER_PHONE.keys()) | set(ACTIVITY_LOG.keys()))
    for cid in targets:
        # phone
        if cid in USER_PHONE and now - USER_PHONE[cid]["saved_at"] > TTL_SECONDS:
            USER_PHONE.pop(cid, None)
        # activities
        if cid in ACTIVITY_LOG:
            ACTIVITY_LOG[cid] = [e for e in ACTIVITY_LOG[cid] if now - e["ts"] <= TTL_SECONDS]
            if not ACTIVITY_LOG[cid]:
                ACTIVITY_LOG.pop(cid, None)

REQUEST_PHONE_TEXT = "📱 لطفاً شماره موبایل خود را ارسال کنید (یا دکمه ارسال شماره را بزنید):"

def phone_request_keyboard():
    kb = [[KeyboardButton("ارسال شماره من", request_contact=True)]]
    return ReplyKeyboardMarkup(kb, resize_keyboard=True, one_time_keyboard=True)

def ask_for_phone(chat_id, context):
    context.user_data["awaiting_phone"] = True
    context.bot.send_message(
        chat_id=chat_id,
        text=REQUEST_PHONE_TEXT,
        reply_markup=phone_request_keyboard()
    )

def confirm_phone_and_continue(chat_id, context, phone: str):
    context.user_data["awaiting_phone"] = False
    context.user_data.pop("awaiting", None)

    context.bot.send_message(
        chat_id=chat_id,
        text=f"✅ شماره شما ثبت شد: {phone}\n👇 لطفاً یکی از خدمات هوش مصنوعی را انتخاب کنید:",
        reply_markup=ai_services_keyboard()
    )



    
def is_valid_persian_national_id(s: str) -> bool:
    """بررسی کند که ورودی دقیقاً ۱۰ رقم فارسی باشد"""
    if not s or len(s) != 10:
        return False
    return all('۰' <= ch <= '۹' for ch in s)

def is_valid_persian_date(s: str) -> bool:
    # الگوی YYYY/MM/DD با اعداد فارسی
    return bool(re.fullmatch(r"[۰-۹]{4}/[۰-۹]{2}/[۰-۹]{2}", s or ""))

def has_min_digits_fa(s: str, n: int = 10) -> bool:
    # تبدیل به انگلیسی و شمارش رقم‌ها
    en = fa_to_en_number(s or "")
    digits = "".join(ch for ch in en if ch.isdigit())
    return len(digits) >= n

def _meeting_title_by_jalali_date(date_str: str) -> str:
    """
    اگر ماه جلالی بین ۱ تا ۴ باشد → «مجمع عمومی عادی بطور سالیانه»
    در غیر این صورت → «مجمع عمومی عادی بطور فوق العاده»
    انتظار فرمت: YYYY/MM/DD با اعداد فارسی (مثل ۱۴۰۴/۰۵/۱۵)
    """
    if not date_str or date_str.count("/") != 2:
        return "مجمع عمومی عادی بطور فوق العاده"
    en = fa_to_en_number(date_str)
    try:
        _y, m, _d = [int(x) for x in en.split("/")]
        return "مجمع عمومی عادی بطور سالیانه" if 1 <= m <= 4 else "مجمع عمومی عادی بطور فوق العاده"
    except Exception:
        return "مجمع عمومی عادی بطور فوق العاده"

def newspapers_keyboard():
    from telegram import InlineKeyboardButton, InlineKeyboardMarkup
    btns = []
    row = []
    for idx, name in enumerate(NEWSPAPERS, start=1):
        row.append(InlineKeyboardButton(name, callback_data=f"newspaper:{idx}"))
        if len(row) == 3:
            btns.append(row); row = []
    if row:
        btns.append(row)
    # دکمه کنسل در صورت نیاز
    btns.append([InlineKeyboardButton("❌ انصراف", callback_data="newspaper:cancel")])
    return InlineKeyboardMarkup(btns)

def send_newspaper_menu(chat_id, context, prompt_text="روزنامهٔ کثیرالانتشار را انتخاب کنید:"):
    # علامت می‌زنیم که الان منتظر انتخاب روزنامه‌ایم (برای هندلر برگشت/دیباگ مفید است)
    ctx = context.user_data.setdefault(chat_id, {}) if isinstance(context.user_data, dict) else context.user_data
    ctx["awaiting"] = "newspaper"
    context.bot.send_message(chat_id=chat_id, text=prompt_text, reply_markup=newspapers_keyboard())


# ——— [B] هندلر انتخاب روزنامه (پچ‌شده) ———
def handle_newspaper_choice(update: Update, context: CallbackContext):
    query = update.callback_query
    chat_id = query.message.chat_id if hasattr(query.message, "chat_id") else query.message.chat.id
    payload = query.data  # مثل "newspaper:5"
    if not payload.startswith("newspaper:"):
        return

    try: query.answer()
    except: pass

    _, choice = payload.split(":", 1)

    # ← دیکشنری وضعیتِ اصلی پروژه
    d = user_data.setdefault(chat_id, {})

    if choice == "cancel":
        # هم در user_data و هم (در صورت استفاده) در context.user_data پاک کن
        d.pop("awaiting", None)
        try: context.bot.edit_message_reply_markup(chat_id=chat_id, message_id=query.message.message_id, reply_markup=None)
        except: pass
        context.bot.send_message(chat_id=chat_id, text="انتخاب روزنامه لغو شد.", reply_markup=main_keyboard())
        return

    # ایندکس معتبر؟
    try:
        idx = int(choice)
        name = NEWSPAPERS[idx - 1]
    except Exception:
        context.bot.send_message(chat_id=chat_id, text="انتخاب نامعتبر روزنامه.", reply_markup=main_keyboard())
        return

    # ذخیره در user_data (همان جایی که بقیهٔ سناریو می‌خوانند)
    d["روزنامه کثیرالانتشار"] = name
    d.pop("awaiting", None)

    # حذف کیبورد اینلاین پیام قبلی
    try:
        context.bot.edit_message_reply_markup(chat_id=chat_id, message_id=query.message.message_id, reply_markup=None)
    except:
        pass

    موضوع = d.get("موضوع صورتجلسه") or d.get("موضوع") or context.user_data.get("topic")
    step = d.get("step", 0)

    # از 17 → 18
    d["step"] = step + 1

    try:
        # ✅ بلافاصله بعد از انتخاب روزنامه، سؤال «وکیل» (step=18) را بپرس
        if موضوع == "تمدید سمت اعضا" and d["step"] == 18:
            label = "نام وکیل (سهامدار یا وکیل رسمی شرکت) را وارد کنید (مثال: آقای ... / خانم ...):"
            if 'remember_last_question' in globals():
                remember_last_question(context, label)
            context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard())
            return

        # فالبک امن
        context.bot.send_message(chat_id=chat_id, text=f"روزنامه انتخاب شد: {name}", reply_markup=main_keyboard())

    except Exception as e:
        context.bot.send_message(chat_id=chat_id, text=f"ثبت روزنامه انجام شد ولی در ادامه فرم مشکلی بود: {e}", reply_markup=main_keyboard())


def build_contact_html(phone_ir: str, phone_intl: str, wa_text: str = "") -> str:
    """
    خروجی: متن HTML شامل لینک تماس مستقیم (tel:) و واتساپ (wa.me)
    """
    tel_link = f"<a href='tel:{phone_ir}'>تماس تلفنی</a>"
    wa_base = f"https://wa.me/{phone_intl}"
    if wa_text:
        wa_link = f"<a href='{wa_base}?text={quote(wa_text)}'>چت در واتساپ</a>"
    else:
        wa_link = f"<a href='{wa_base}'>چت در واتساپ</a>"
    return f"📞 {tel_link}\n💬 {wa_link}"


def send_thank_you_message_chatid(chat_id, context,
                                  phone_ir=None, phone_intl=None,
                                  wa_text=None, brand=None):
    phone_ir = phone_ir or CONTACT_MOBILE_IR
    phone_intl = phone_intl or CONTACT_MOBILE_INTL  # بدون صفر
    wa_text = wa_text if wa_text is not None else DEFAULT_WHATSAPP_TEXT
    brand = brand or THANKYOU_BRAND

    # متن پیام پایانی (شماره بین‌المللی با + برای لمس مستقیم روی موبایل)
    msg = (
        "🎉 صورتجلسه شما آماده و ارسال شد!\n"
        f"از اینکه {brand} رو انتخاب کردید سپاسگزاریم 🙏\n\n"
        "☎️ برای مشاوره بیشتر یا ثبت صورتجلسه:\n"
        f"• شماره تماس: +{phone_intl}\n"
    )

    # دکمه واتساپ (http/https تنها اسکیماهای مجاز)
    wa_url = f"https://wa.me/{phone_intl}"
    if wa_text:
        wa_url += f"?text={quote(wa_text)}"
    keyboard = InlineKeyboardMarkup([
        [InlineKeyboardButton("💬 چت در واتساپ", url=wa_url)]
    ])

    # 1) ارسال پیام تشکر + دکمه
    context.bot.send_message(
        chat_id=chat_id,
        text=msg,
        disable_web_page_preview=True
    )

    # 2) ارسال Contact واقعی (قابل لمس و ذخیره در مخاطبین)
    try:
        context.bot.send_contact(
            chat_id=chat_id,
            phone_number=f"+{phone_intl}",  # حتماً با + شروع شود
            first_name=brand,
            last_name="پشتیبانی"
            # می‌توانی vCard هم اضافه کنی اگر خواستی
        )
    except Exception:
        # اگر کاربر اجازه دریافت مخاطب نداده بود، مشکلی نیست
        pass

    # 3) ارسال دکمه واتساپ به‌صورت جدا (اختیاری؛ اگر می‌خواهی کنار Contact هم باشد)
    try:
        context.bot.send_message(
            chat_id=chat_id,
            text="برای شروع چت در واتساپ روی دکمه زیر بزنید:",
            reply_markup=keyboard,
            disable_web_page_preview=True
        )
    except Exception:
        pass
        


def enter_ai_mode_reply(update: Update, context: CallbackContext, sys_prompt: str = None):
    chat_id = update.effective_chat.id
    context.user_data["ai_mode"] = True
    context.user_data["ai_sys_prompt"] = sys_prompt or (
        "شما کارشناس قانون تجارت ایران و امور ثبت شرکت‌ها هستید. پاسخ دقیق، مرحله‌به‌مرحله و با ذکر نکات اجرایی بده."
    )
    # اگر محدودیت هم داری:
    context.user_data["ai_q_count"] = 0
    context.user_data["ai_q_limit"] = globals().get("AI_Q_LIMIT", 5)

    msg = update.message.reply_text(
        "🧠 حالت هوشمند ما فعال شد.\nسؤالت رو بپرس",
        reply_markup=ReplyKeyboardRemove()
    )

    # ⛔️ فقط اگر از مسیر «مشاوره…» نیامده باشد، دکمه اینلاین را اضافه کن
    if not context.user_data.get("ai_skip_inline_back"):
        try:
            msg.edit_reply_markup(
                reply_markup=InlineKeyboardMarkup(
                    [[InlineKeyboardButton("↩️ بازگشت به ادامه تنظیم صورتجلسه", callback_data=AI_RESUME)]]
                )
            )
        except Exception as e:
            context.bot.send_message(
                chat_id=chat_id,
                text="برای بازگشت از دکمهٔ زیر استفاده کن:",
                reply_markup=InlineKeyboardMarkup(
                    [[InlineKeyboardButton("↩️ بازگشت به ادامه تنظیم صورتجلسه", callback_data=AI_RESUME)]]
                )
            )
            print("edit_reply_markup failed:", e)

    # 🧹 فلگ فقط برای همین بار بود؛ پاکش کن که دفعه بعد تاثیر نگذارد
    context.user_data.pop("ai_skip_inline_back", None)



# ==== Recommendation builder for comp_type ====
def build_comp_type_recommendation(data: dict) -> str:
    score = 0
    bullets = []

    # 1) تعداد شرکا
    t = data.get("CT_تعداد_شرکا","")
    if t in ("1 نفر","2 نفر"):
        score -= 1
        bullets.append("تعداد شرکا کم است → «مسئولیت محدود» معمولاً ساده‌تر و سریع‌تر ثبت می‌شود.")
    elif t == "3 نفر":
        score += 1
        bullets.append("حداقل 3 شریک دارید → «سهامی خاص» با پیش‌نیاز 3 سهامدار هم‌خوان است.")
    elif t == "4 نفر یا بیشتر":
        score += 1
        bullets.append("تعداد شرکا بالا → سهامی خاص برای تقسیم سهام و شفافیت مناسب‌تر است.")

    # 2) سرمایه
    cap = int(data.get("CT_سرمایه", 0) or 0)
    if cap >= 500_000_000:
        score += 1
        bullets.append("سرمایه نسبتاً بالا → سهامی خاص برای اعتبار/قراردادهای بزرگ مناسب‌تر است.")
    else:
        score -= 0.5
        bullets.append("سرمایه کم/متوسط → «مسئولیت محدود» با تشریفات ساده‌تر به‌صرفه‌تر است.")

    # 3) مناقصات/اعتبار بانکی
    if data.get("CT_مناقصات", False):
        score += 1.5
        bullets.append("برنامه جدی برای مناقصات/اعتبار بانکی → سهامی خاص اعتبار بیشتری نزد کارفرما/بانک دارد.")
    else:
        score -= 0.5
        bullets.append("تمرکز بر فعالیت‌های ساده/داخلی → «مسئولیت محدود» کفایت می‌کند.")

    # 4) هیئت‌مدیره
    if data.get("CT_هیئت_مدیره", False):
        score += 1
        bullets.append("تمایل به هیئت‌مدیره 3 نفره → «سهامی خاص» انتخاب طبیعی‌تری است.")
    else:
        score -= 0.5
        bullets.append("مدیریت ساده/یک‌نفره → «مسئولیت محدود» عملی‌تر و سریع‌تر است.")

    # 5) سهولت نقل‌وانتقال
    tr = data.get("CT_انتقال","اهمیت متوسط")
    if tr == "خیلی مهم است":
        score += 1
        bullets.append("سهولت نقل‌وانتقال مهم → «سهامی خاص» (سهام) انعطاف بهتری نسبت به سهم‌الشرکه دارد.")
    elif tr == "اهمیت ندارد":
        score -= 0.5
        bullets.append("نقل‌وانتقال مهم نیست → «مسئولیت محدود» نیز گزینه خوبی است.")

    if score >= 1:
        title = "پیشنهاد: «سهامی خاص» ✅"
        why = "با توجه به پاسخ‌ها، سهامی خاص از نظر اعتبار، مناقصات، تقسیم سهام و آینده‌نگری حرفه‌ای مناسب‌تر است."
        next_hint = "اکنون می‌توانید از منوی اصلی مسیر صورتجلسه‌های متناسب با سهامی خاص را شروع کنید."
    else:
        title = "پیشنهاد: «مسئولیت محدود» ✅"
        why = "با توجه به پاسخ‌ها، مسئولیت محدود به‌خاطر سرعت، هزینه کمتر و سادگی اداره مناسب‌تر است."
        next_hint = "اکنون می‌توانید از منوی اصلی مسیر صورتجلسه‌های متناسب با مسئولیت محدود را شروع کنید."

    bullet_text = "\n".join([f"• {b}" for b in bullets])
    cap_fmt = f"{int(cap):,}".replace(",", "،")
    recap = (
        f"🔎 خلاصه پاسخ‌های شما:\n"
        f"— تعداد شرکا: {t}\n"
        f"— سرمایه: {cap_fmt} ریال\n"
        f"— مناقصه/اعتبار بانکی: {'بله' if data.get('CT_مناقصات') else 'خیر'}\n"
        f"— هیئت‌مدیره 3 نفره: {'بله' if data.get('CT_هیئت_مدیره') else 'خیر'}\n"
        f"— اهمیت نقل‌وانتقال: {tr}\n"
    )

    return (
        f"{title}\n\n{why}\n\n"
        f"دلایل ارزیابی:\n{bullet_text}\n\n"
        f"{recap}\n"
        f"ℹ️ در صورت نیاز، تفاوت‌های حقوقی/عملیاتی را هم می‌توانم دقیق‌تر توضیح بدهم.\n"
        f"{next_hint}"
    )


def make_formal_text_with_ai(raw_text: str, style_hint: str = "⚖️ رسمی و روان") -> str:
    """
    raw_text: متن ساده کاربر
    style_hint: شدت/سبک رسمیت انتخاب‌شده

    خروجی: متن رسمی استاندارد فارسی با لحن حقوقی/اداری
    """

    # نگاشت سبک به دستور
    style_map = {
        "🔒 خیلی رسمی و حقوقی": "خیلی رسمی، حقوقی و کاملاً اداری؛ جمله‌بندی‌های محکم و ارجاعی.",
        "⚖️ رسمی و روان": "رسمی و روان؛ دقیق، فاقد اغراق، خوانا و استاندارد اداری.",
        "🤝 رسمی دوستانه": "رسمی اما صمیمی و محترمانه؛ کوتاه‌تر و قابل‌فهم برای مخاطب غیرحقوقی."
    }
    style_directive = style_map.get(style_hint, style_map["⚖️ رسمی و روان"])

    system_prompt = (
        "شما ویرایشگر حقوقی/اداری فارسی هستید. متن ورودی را بدون تغییر در حقایق، "
        "به یک متن رسمی و استاندارد تبدیل کنید. از تعارفات دوری کنید، اعداد را به "
        "صورت استاندارد و قابل ثبت بنویسید، شفاف و قابل استناد. اگر متن مبهم است، "
        "ابهام را با ساختار حقوقی حفظ کنید و چیزی اختراع نکنید."
    )

    user_prompt = f"""
# دستور سبک:
{style_directive}

# الزامات:
- به هیچ وجه اطلاعات جدید اختراع نکن.
- اگر اطلاعات ناقص است، آن را با عبارات حقوقیِ «حسب اطلاعات موجود/نظر به...» مدیریت کن.
- لحن: رسمی فارسی؛ علائم سجاوندی درست؛ پاراگراف‌بندی تمیز.
- اگر موضوع «اعلام/درخواست/ابلاغ» است، تیتر مناسبی در ابتدای متن پیشنهاد بده (در یک خط).
- اگر مناسب است، تاریخ/شماره/مرجع را به شکل فیلدِ قابل‌جایگذاری بیاور (مثلاً: «تاریخ: .... / شماره: ....»).
- خروجی فقط متن نهایی باشد.

# متن ورودی:
{raw_text}
    """.strip()

    # اینجا از همان تابع AI پروژه استفاده کن (ask_groq / ask_openai / هرچه داری)
    # فرض: ask_groq(prompt: str) -> str
    result = ask_groq(system_prompt, user_prompt)  # ← امضای تابع خودت را رعایت کن
    return result.strip()





def handle_ai_text(update, context):
    if not context.user_data.get("ai_mode"):
        return

    text = (update.message.text or "").strip()
    if text == AI_ASK_TEXT:
        return

    # فقط دکمه‌های «بازگشت» واقعاً از AI خارج کنند
    if text in (BACK_BTN, "🔙 بازگشت به ادامه مراحل"):
        # توصیه: اطمینان از پاک‌شدن حالت AI در resume_from_ai
        # داخل resume_from_ai حتماً این‌ها را داشته باش:
        # context.user_data.pop("ai_mode", None)
        # data.pop("FORMAL_RAW", None); data.pop("FORMAL_STYLE", None); data["step"] = 0
        resume_from_ai(update, context)
        return

    chat_id = update.effective_chat.id
    context.bot.send_chat_action(chat_id=chat_id, action=ChatAction.TYPING)

    ai_mode = context.user_data.get("ai_mode")

    # لیست گزینه‌های اصلی AI برای تشخیص "متن دکمه‌ی منو"
    AI_TOP_OPTIONS = (
        AI_OPT_MINUTES, AI_OPT_QA, AI_OPT_COMP_TYPE, AI_OPT_NAME, AI_OPT_CONTRACT, AI_OPT_FORMAL
    )

    # ---------------------------
    # شاخهٔ مخصوص formalizer (رسمی‌سازی متن)
    # ---------------------------
    if ai_mode == "formalizer":
        user_data.setdefault(chat_id, {})
        data = user_data[chat_id]
        step = data.get("step", 1)
        if step not in (1, 2):
            step = 1
            data["step"] = 1
    
        try:
            # --- گام 1: دریافت متن خام ---
            if step == 1:
                # اگر کاربر یکی از گزینه‌های منوی AI را زد، جلو نرو
                if text in AI_TOP_OPTIONS:
                    update.message.reply_text(
                        "برای تغییر سرویس، ابتدا «⬅️ بازگشت» را بزنید.",
                        reply_markup=back_keyboard()
                    )
                    return
    
                if not text:
                    # ❌ قبلاً: reply_text(" ", reply_markup=ReplyKeyboardRemove()) → BadRequest
                    # ✅ مستقیم همان پیام درخواست را با کیبورد بازگشت بفرست:
                    update.message.reply_text("📝 لطفاً متن ساده‌تان را ارسال کنید.", reply_markup=back_keyboard())
                    return
    
                data["FORMAL_RAW"] = text
                data["step"] = 2
    
                keyboard = [[
                    "🔒 خیلی رسمی و حقوقی",
                    "⚖️ رسمی و روان",
                    "🤝 رسمی دوستانه"
                ], [BACK_BTN]]
                update.message.reply_text(
                    "سبک/شدت رسمیت را انتخاب کنید:",
                    reply_markup=ReplyKeyboardMarkup(keyboard, resize_keyboard=True)
                )
                return
    
            # --- گام 2: دریافت سبک ---
            if step == 2:
                style = text
                valid_styles = ("🔒 خیلی رسمی و حقوقی", "⚖️ رسمی و روان", "🤝 رسمی دوستانه")
                if style not in valid_styles:
                    update.message.reply_text(
                        "لطفاً یکی از گزینه‌های سبک را انتخاب کنید.",
                        reply_markup=back_keyboard()
                    )
                    return

                raw = (data.get("FORMAL_RAW", "") or "").strip()
                if not raw:
                    data["step"] = 1
                    update.message.reply_text("❗️ابتدا متن ساده را ارسال کنید.",
                                              reply_markup=ReplyKeyboardMarkup([[BACK_BTN]], resize_keyboard=True))
                    return

                style_map = {
                    "🔒 خیلی رسمی و حقوقی": "خیلی رسمی، حقوقی و کاملاً اداری؛ جمله‌بندی‌های محکم و ارجاعی.",
                    "⚖️ رسمی و روان": "رسمی و روان؛ دقیق، فاقد اغراق، خوانا و استاندارد اداری.",
                    "🤝 رسمی دوستانه": "رسمی اما صمیمی و محترمانه؛ کوتاه‌تر و قابل‌فهم برای مخاطب غیرحقوقی."
                }
                style_directive = style_map[style]

                combined_prompt = f"""
[System]
شما ویرایشگر حقوقی/اداری فارسی هستید. متن ورودی را بدون تغییر در حقایق، به یک متن رسمی و استاندارد تبدیل کنید.
از تعارفات دوری کنید، اعداد را استاندارد بنویسید، متن شفاف و قابل استناد باشد. چیزی اختراع نکنید.

[Style]
{style_directive}

[Requirements]
- اطلاعات جدید اضافه نشود.
- اگر اطلاعات ناقص است، با عبارات حقوقی مانند «حسب اطلاعات موجود/نظر به...» مدیریت شود.
- لحن: رسمی فارسی؛ علائم سجاوندی درست؛ پاراگراف‌بندی تمیز.
- اگر موضوع «اعلام/درخواست/ابلاغ» است، تیتر مناسب در ابتدای متن پیشنهاد شود (در یک خط).
- در صورت لزوم، تاریخ/شماره/مرجع به‌صورت فیلد قابل‌جایگذاری بیاید (مثلاً: «تاریخ: .... / شماره: ....»).
- خروجی فقط متن نهایی باشد.

[Input]
{raw}
                """.strip()

                # تولید متن رسمی
                answer = ask_groq(combined_prompt, max_tokens=900)

                # پاسخ را تکه‌تکه بفرست (بدون دکمهٔ اینلاین)
                chunks = [answer[i:i+3500] for i in range(0, len(answer), 3500)]
                for ch in chunks:
                    update.message.reply_text(ch)

                # ارسال فایل Word
                try:
                    file_path = generate_word_file(answer)
                    with open(file_path, 'rb') as f:
                        context.bot.send_document(chat_id=chat_id, document=f, filename="متن_رسمی.docx")
                except Exception as fe:
                    print("WORD FILE ERROR:", fe)

                # ماندن در formalizer اما بازگشت به گام 1، با کیبورد مینیمال
                data["step"] = 1
                update.message.reply_text(
                    "✅ تمام شد. اگر متن دیگری دارید، همین‌جا بفرستید یا با «بازگشت» خارج شوید.",
                    reply_markup=ReplyKeyboardMarkup([[BACK_BTN]], resize_keyboard=True)
                )
                return

        except Exception as e:
            update.message.reply_text("❌ خطا در تولید متن رسمی. کمی بعد دوباره تلاش کنید.")
            print("FORMALIZER ERROR:", e)
            user_data[chat_id]["step"] = 1
            update.message.reply_text("📝 لطفاً دوباره متن ساده‌تان را بفرستید.",
                                      reply_markup=ReplyKeyboardMarkup([[BACK_BTN]], resize_keyboard=True))
            return

    # ---------------------------
    # رفتار پیش‌فرض برای سایر ai_mode ها (منطق قبلی‌ات)
    # ---------------------------
    try:
        answer = ask_groq(text, max_tokens=900)
        chunks = [answer[i:i+3500] for i in range(0, len(answer), 3500)]
        for ch in chunks:
            update.message.reply_text(ch)
    except Exception as e:
        update.message.reply_text("❌ خطا در دریافت پاسخ هوشمند. کمی بعد دوباره تلاش کنید.")
        print("GROQ ERROR:", e)






def resume_from_ai(update, context):
    # 1) اگر از اینلاین‌باتن بود، فقط answer کن
    q = getattr(update, "callback_query", None)
    if q:
        try: q.answer()
        except Exception: pass

    # 2) خاموش کردن حالت AI
    context.user_data["ai_mode"] = False

    # 3) chat_id امن
    chat_id = None
    if getattr(update, "effective_chat", None):
        chat_id = update.effective_chat.id
    elif q and getattr(q, "message", None):
        chat_id = q.message.chat_id
    if not chat_id:
        return

    # 4) ادامه‌ی مرحله قبلی: «خودِ آخرین سؤال» را بفرست
    last_q = context.user_data.get("last_question_text") or context.user_data.get("last_question")
    if last_q:
        # ⬅️ دقیقاً همان سؤال قبلی را دوباره به کاربر نشان بده
        context.bot.send_message(chat_id=chat_id, text=last_q, reply_markup=base_reply_keyboard())
        return

    # اگر چیزی ذخیره نشده بود، برگرد به انتخاب موضوع
    send_topic_menu(chat_id, context)



def generate_word_file(text: str, filepath: str = None):
    _lazy_import_docx()
    doc = Document()

    # تنظیم فونت B Nazanin اگر نصب باشد
    style = doc.styles['Normal']
    font = style.font
    font.name = 'B Nazanin'
    font.size = Pt(14)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'B Nazanin')

    # راست‌چین کردن و بولد کردن فقط خط اول
    lines = text.strip().split('\n')
    for i, line in enumerate(lines):
        if not line.strip():
            continue
        p = doc.add_paragraph()
        run = p.add_run(line.strip())
        if i == 0:
            run.bold = True
        p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    # مسیر ذخیره‌سازی فایل
    if not filepath:
        filename = f"soratjalase_{uuid.uuid4().hex}.docx"
        filepath = os.path.join("/tmp", filename)

    doc.save(filepath)
    return filepath
def send_topic_menu(chat_id, context):
    """منوی انتخاب «موضوع صورتجلسه» را نشان می‌دهد."""
    keyboard = [
        [InlineKeyboardButton("🏢 تغییر آدرس", callback_data='تغییر آدرس')],
        [InlineKeyboardButton("🔄 نقل و انتقال سهام", callback_data='نقل و انتقال سهام')],
        [InlineKeyboardButton("🧾 تغییر موضوع فعالیت", callback_data='تغییر موضوع فعالیت')],
        [InlineKeyboardButton("👔 انتخاب مدیران", callback_data='topic:board_election')],
        [InlineKeyboardButton("⏳ تمدید سمت اعضا", callback_data="topic:extend_roles")],
        [InlineKeyboardButton("📈 افزایش سرمایه", callback_data='افزایش سرمایه')],
        [InlineKeyboardButton("📉 کاهش سرمایه", callback_data='کاهش سرمایه')],
        [InlineKeyboardButton("🏷️ تغییر نام شرکت", callback_data='تغییر نام شرکت')],
        [InlineKeyboardButton("❌ انحلال شرکت", callback_data='انحلال شرکت')],
        [InlineKeyboardButton("💰 پرداخت سرمایه تعهدی شرکت", callback_data='پرداخت سرمایه تعهدی شرکت')],
    ]
    reply_markup = InlineKeyboardMarkup(keyboard)
    context.bot.send_message(
        chat_id=chat_id,
        text="💬 برای چه موضوعی صورتجلسه نیاز دارید؟\nلطفاً یکی از گزینه‌های زیر را انتخاب کنید:",
        reply_markup=reply_markup
    )

def send_company_type_menu(chat_id, context):
    """پس از انتخاب موضوع، منوی «نوع شرکت» را نشان می‌دهد."""
    keyboard = [
        [InlineKeyboardButton("سهامی خاص", callback_data='سهامی خاص')],
        [InlineKeyboardButton("مسئولیت محدود", callback_data='مسئولیت محدود')]
    ]
    reply_markup = InlineKeyboardMarkup(keyboard)
    context.bot.send_message(
        chat_id=chat_id,
        text="نوع شرکت را انتخاب کنید:",
        reply_markup=reply_markup
    )


def start(update: Update, context: CallbackContext):
    chat_id = update.message.chat_id
    user_data[chat_id] = {"step": 0}

    update.message.reply_text(
        "به ثبت کوشا خوش آمدید 🙌\n"
        "اینجا می‌توانید:\n"
        "• صورتجلسه‌های رسمی شرکت را هوشمند ، دقیق و با کمترین خطا بسازید (متن + فایل Word)\n"
        "• با «دستیار هوشمند» راهنمای نوع شرکت بگیرید، برای شرکت نام پیشنهاد شود، متن‌ها را رسمی کنید و قراردادهای پایه بسازید\n"
        "• پرسش‌های فوری درباره رویه‌های ثبت و مقررات را بپرسید"
    )

    # اگر قبلاً در ۷ روز اخیر شماره دارد، مستقیم منو را بده
    saved = get_user_phone(chat_id)
    if saved:
        context.user_data["awaiting_phone"] = False
        context.bot.send_message(chat_id=chat_id, text=f"📌 شماره تأییدشده شما: {saved}")
        send_ai_services_menu(chat_id, context)
        
    else:
        # در غیر این صورت، شماره را بگیریم
        ask_for_phone(chat_id, context)




def start_extend_roles_flow(update, context):
    chat_id = update.effective_chat.id
    user_data.setdefault(chat_id, {})
    d = user_data[chat_id]

    # مقداردهی اولیه سناریو
    d["موضوع صورتجلسه"] = TOPIC_EXTEND_ROLES
    d["نوع شرکت"] = "سهامی خاص"   # مهم: گارد «نوع شرکت را انتخاب کنید» را دور می‌زنیم
    d["step"] = 1

    # پاک‌سازی باقیمانده‌های احتمالی از اجرای قبلی همین سناریو
    for k in ["عضو_index", "سهامدار_index", "تعداد اعضای هیئت مدیره", "تعداد سهامداران"]:
        d.pop(k, None)
    for k in list(d.keys()):
        if k.startswith("عضو ") or k.startswith("سهامدار "):
            d.pop(k, None)

    # سؤال اول (هماهنگ با روال پروژه)
    label = get_label("نام شرکت") if 'get_label' in globals() else "نام شرکت را وارد کنید:"
    if 'remember_last_question' in globals():
        remember_last_question(context, label)

    kb = main_keyboard() if 'main_keyboard' in globals() else None
    if kb:
        context.bot.send_message(chat_id=chat_id, text=label, reply_markup=kb)
    else:
        context.bot.send_message(chat_id=chat_id, text=label)


def get_label(field, **kwargs):
    labels = {
        "نوع شرکت": "نوع شرکت را انتخاب کنید:",
        "نام شرکت": "نام شرکت را وارد کنید:",
        "شماره ثبت": "شماره ثبت شرکت را وارد کنید (اعداد فارسی):",
        "شناسه ملی": "شناسه ملی شرکت را وارد کنید (اعداد فارسی):",
        "سرمایه": "سرمایه ثبت‌شده شرکت را به ریال وارد کنید (اعداد فارسی):",
        "تاریخ": "تاریخ صورتجلسه را وارد کنید (مثلاً: ۱۴۰۴/۰۵/۱۵):",
        "ساعت": "ساعت جلسه را وارد کنید (اعداد فارسی):",
        "مدیر عامل": "مدیر عامل (رئیس جلسه) را وارد کنید:",
        "نایب رییس": "ناظر 1 جلسه را وارد کنید (از بین هیئت مدیره):",
        "رییس": "ناظر 2 جلسه را وارد کنید (از بین هیئت مدیره):",
        "منشی": "منشی جلسه را وارد کنید:",
        "آدرس جدید": "آدرس جدید شرکت را وارد کنید:",
        "کد پستی": "کد پستی آدرس جدید را وارد کنید (اعداد فارسی):",
        "وکیل": "نام وکیل (ثبت‌کننده صورتجلسه) را وارد کنید:",
        "شماره دفترخانه": "شماره دفترخانه فروشنده {i} را وارد کنید (مثلاً: 22 تهران):",
        "نام جدید شرکت": "نام جدید شرکت را وارد کنید:",

        # برچسب‌های مخصوص انحلال
        "علت انحلال": "علت انحلال را وارد کنید (مثلاً: مشکلات اقتصادی):",
        "نام مدیر تصفیه": "نام مدیر تصفیه را وارد کنید:",
        "کد ملی مدیر تصفیه": "کد ملی مدیر تصفیه را وارد کنید (اعداد فارسی):",
        "مدت مدیر تصفیه": "مدت مدیر تصفیه (سال) را وارد کنید (اعداد فارسی):",
        "آدرس مدیر تصفیه": "آدرس مدیر تصفیه و محل تصفیه را وارد کنید:",
        "تعداد سهامداران حاضر": "تعداد سهامداران حاضر را وارد کنید (عدد):",

        # برای مسیرهای دیگر که استفاده داری
        "تعداد شرکا": "تعداد شرکا را وارد کنید (بین ۲ تا ۷):",

        # 🔔 اطلاعیه ماده ۱۰۳
        "اطلاعیه_ماده103": (
            "یادآوری مهم — ماده ۱۰۳ قانون تجارت ⚖️\n"
            "نقل‌وانتقال سهم‌الشرکه در شرکت با مسئولیت محدود، از عقود تشریفاتی است و باید به موجب «سند رسمی» در دفترخانه انجام شود. 🏛️📄\n\n"
            "برای تکمیل این صورتجلسه، لازم است ابتدا {سند} را در یکی از دفاتر اسناد رسمی تنظیم کرده باشید؛ "
            "زیرا درج مشخصات آن در متن صورتجلسه الزامی است. ✍️🧾"
        ),
    }

    msg = labels.get(field, f"{field} را وارد کنید:")
    try:
        return msg.format(**kwargs)  # برای کلیدهایی که جای‌نگهدار دارند مثل {سند}، {i}، {k}
    except Exception:
        return msg

def cmd_ai(update, context):
    chat_id = update.effective_chat.id
    args_text = (update.message.text or "").split(" ", 1)
    query = args_text[1].strip() if len(args_text) > 1 else ""

    if not query:
        update.message.reply_text("سؤال را بعد از /ai بنویسید.")
        return

    try:
        answer = ask_groq(query, max_tokens=900)  # بدون انتخاب مدل
        for i in range(0, len(answer), 3500):
            update.message.reply_text(answer[i:i+3500])
    except Exception as e:
        update.message.reply_text("❌ خطا در دریافت پاسخ از Groq.")
        print("GROQ ERROR:", e)



# --- [A] کیبورد انتخاب سمت عضو هیئت‌مدیره ---
def roles_keyboard(member_index: int):
    kb = [
        [InlineKeyboardButton("رئیس هیئت‌مدیره", callback_data=f"role:{member_index}:chair")],
        [InlineKeyboardButton("نایب رئیس هیئت‌مدیره", callback_data=f"role:{member_index}:vice")],
        [InlineKeyboardButton("مدیرعامل", callback_data=f"role:{member_index}:ceo")],
        [InlineKeyboardButton("عضو هیئت‌مدیره", callback_data=f"role:{member_index}:member")],
        [InlineKeyboardButton("مدیرعامل و رئیس هیئت‌مدیره",   callback_data=f"role:{member_index}:ceo_chair")],
        [InlineKeyboardButton("مدیرعامل و نایب رئیس هیئت‌مدیره", callback_data=f"role:{member_index}:ceo_vice")],
        [InlineKeyboardButton("مدیرعامل و عضو هیئت‌مدیره",    callback_data=f"role:{member_index}:ceo_member")],
    ]
    return InlineKeyboardMarkup(kb)

# --- [B] کیبورد انتخاب حق‌امضا برای هر عضو ---
def sign_authority_keyboard(member_index: int):
    kb = [
        [InlineKeyboardButton("اوراق و اسناد بهادار و تعهد‌آور", callback_data=f"sig:{member_index}:b")],
        [InlineKeyboardButton("اوراق عادی و اداری", callback_data=f"sig:{member_index}:n")],
        [InlineKeyboardButton("هر دو گزینه", callback_data=f"sig:{member_index}:bn")],
        [InlineKeyboardButton("❌ حق امضا ندارد", callback_data=f"sig:{member_index}:none")],
    ]
    return InlineKeyboardMarkup(kb)

# --- [C] سؤال اضافی برای مدیرعامل: خارج از سهامداران هست؟ ---
def ceo_outside_keyboard(member_index: int):
    kb = [
        [InlineKeyboardButton("بله", callback_data=f"ceo_out:{member_index}:yes")],
        [InlineKeyboardButton("خیر", callback_data=f"ceo_out:{member_index}:no")],
    ]
    return InlineKeyboardMarkup(kb)


# --- [D] سازنده‌ی بند «حق‌امضا هوشمند» ---
def build_signature_clause_roles(d: dict) -> str:
    """
    خروجیِ یک‌جمله‌ای در صورت وجود هر دو دسته؛
    اگر فقط یکی موجود باشد همان یک جمله ساخته می‌شود.
    گزینه‌ی 'none' نادیده گرفته می‌شود.
    'متفق' قبل از لیست سمت‌ها می‌آید.
    """

    def fa_role_label(code: str) -> str:
        return {
            "chair":       "رئیس هیئت‌مدیره",
            "vice":        "نایب رئیس هیئت‌مدیره",
            "ceo":         "مدیرعامل",
            "member":      "عضو هیئت‌مدیره",
            "ceo_chair":   "مدیرعامل و رئیس هیئت‌مدیره",
            "ceo_vice":    "مدیرعامل و نایب رئیس هیئت‌مدیره",
            "ceo_member":  "مدیرعامل و عضو هیئت‌مدیره",
        }.get(code, code or "عضو هیئت‌مدیره")

    def uniq(seq):
        seen = set(); out = []
        for x in seq:
            if x not in seen:
                seen.add(x); out.append(x)
        return out

    def fmt(roles):
        roles = uniq(roles)
        if not roles:
            return ""
        if len(roles) == 1:
            return roles[0]
        # «متفق» قبل از لیست سمت‌ها
        return "متفق " + " و ".join(roles)

    total = int(fa_to_en_number(str(d.get("تعداد اعضای هیئت مدیره", 0)) or "0"))
    b_roles, n_roles = [], []

    for i in range(1, total + 1):
        r  = d.get(f"عضو {i} سمت کد")
        ch = d.get(f"عضو {i} حق‌امضا")  # b / n / bn / none
        if not r or not ch:
            continue
        label = fa_role_label(r)
        if ch in ("b", "bn"):
            b_roles.append(label)
        if ch in ("n", "bn"):
            n_roles.append(label)
        # اگر ch == "none" → در هیچ‌جا اضافه نکن

    b_txt = fmt(b_roles)
    n_txt = fmt(n_roles)

    # هیچ امضاکننده‌ای انتخاب نشده:
    if not b_txt and not n_txt:
        return ""

    # هر دو دسته وجود داشته باشند → یک جمله‌ی پیوسته (بدون خط جدید/فاصله اضافی)
    if b_txt and n_txt:
        return (
            "كليه اوراق و اسناد بهادار و تعهد‌آور شركت از قبيل چك، سفته، بروات، قراردادها و عقود اسلامي "
            f"با امضا {b_txt} همراه با مهر شرکت و مکاتبات عادی و اداری "
            f"با امضاء {n_txt} همراه با مهر شرکت معتبر می باشد"
        )

    # فقط بهادار/تعهدآور
    if b_txt:
        return (
            "كليه اوراق و اسناد بهادار و تعهد‌آور شركت از قبيل چك، سفته، بروات، قراردادها و عقود اسلامي "
            f"با امضاء {b_txt} همراه با مهر شرکت معتبر می باشد"
        )

    # فقط عادی/اداری
    return (
        f"مکاتبات عادی و اداری با امضاء {n_txt} همراه با مهر شرکت معتبر می باشد"
    )


def build_signatures_block(d: dict) -> str:
    """
    اسامی اعضای هیئت‌مدیره را دو‌تایی در هر خط چاپ می‌کند.
    اگر تعداد فرد باشد، نفر آخر در یک خط تنها می‌آید.
    فاصله‌ی بین دو اسم با NBSP پر می‌شود تا در کلاینت تلگرام جمع نشود.
    """
    try:
        total = int(fa_to_en_number(str(d.get("تعداد اعضای هیئت مدیره", 0)) or "0"))
    except Exception:
        total = 0

    names = []
    for i in range(1, total + 1):
        nm = (d.get(f"عضو {i} نام", "") or "").strip()
        if nm:
            names.append(nm)

    if not names:
        return ""

    NBSP = "\u00A0"  # non-breaking space
    GAP  = NBSP * 40  # مقدار فاصله بین دو اسم را می‌توانی کم/زیاد کنی (مثلاً 20 یا 40)

    lines = []
    for idx in range(0, len(names), 2):
        left  = names[idx]
        right = names[idx + 1] if idx + 1 < len(names) else ""
        if right:
            lines.append(f"     {left}{GAP}{right}")
        else:
            lines.append(f"     {left}")

    return "امضاء اعضای هیات مدیره\n\n" + "\n".join(lines)




def handle_inline_callbacks(update: Update, context: CallbackContext):
    q = update.callback_query
    if not q:
        return
    chat_id = q.message.chat_id if hasattr(q.message, "chat_id") else q.message.chat.id
    d = user_data.setdefault(chat_id, {})
    payload = q.data or ""
    try:
        q.answer()
    except Exception:
        pass

    # --- انتخاب موضوع: "👔 انتخاب مدیران" ---
    if payload == "topic:board_election":
        # پاکسازی کلیدهای مرتبط با سناریوهای قبلی
        for k in ["step", "board_index", "عضو_index", "سهامدار_index",
                  "تعداد اعضای هیئت مدیره", "تعداد سهامداران"]:
            d.pop(k, None)
        d["موضوع صورتجلسه"] = "انتخاب مدیران"

        # نمایش منوی نوع شرکت
        send_company_type_menu(chat_id, context)
        return

    # --- انتخاب نوع شرکت ---
    if payload in ("سهامی خاص", "مسئولیت محدود"):
        d["نوع شرکت"] = payload

        # شروع سناریوی انتخاب مدیران فقط برای سهامی خاص
        if d.get("موضوع صورتجلسه") == "انتخاب مدیران" and payload == "سهامی خاص":
            d["step"] = 1
            label = get_label("نام شرکت") if 'get_label' in globals() else "نام شرکت را وارد کنید:"
            if 'remember_last_question' in globals():
                remember_last_question(context, label)
            context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard())
            return

        # (در غیر اینصورت می‌تونی اینجا سناریوهای دیگر را آغاز کنی)

    # --- سایر payload ها ... ---

    # --- انتخاب سمت برای عضو i ---
    if payload.startswith("role:"):
        parts = payload.split(":", 2)  # "role:{i}:{code}"
        if len(parts) != 3:
            context.bot.send_message(chat_id=chat_id, text="انتخاب سمت نامعتبر بود.")
            return
        _, idx_str, code = parts
        try:
            i = int(idx_str)
        except ValueError:
            context.bot.send_message(chat_id=chat_id, text="شناسهٔ عضو نامعتبر بود.")
            return
    
        role_map = {
            "chair":       "رئیس هیئت‌مدیره",
            "vice":        "نایب رئیس هیئت‌مدیره",
            "ceo":         "مدیرعامل",
            "member":      "عضو هیئت‌مدیره",
            "ceo_chair":   "مدیرعامل و رئیس هیئت‌مدیره",
            "ceo_vice":    "مدیرعامل و نایب رئیس هیئت‌مدیره",
            "ceo_member":  "مدیرعامل و عضو هیئت‌مدیره",
        }
    
        d[f"عضو {i} سمت کد"] = code
        d[f"عضو {i} سمت"]    = role_map.get(code, "عضو هیئت‌مدیره")
    
        # ✅ پیام خلاصه همزمان با سؤال بعدی: «اسم شخص : سمت شخص»
        person_name  = d.get(f"عضو {i} نام", "")
        person_role  = d.get(f"عضو {i} سمت", "")
        info_line    = f"{person_name} : {person_role}"
    
        if code == "ceo":
            # فقط برای مدیرعامل سؤال اضافه می‌پرسیم
            context.bot.send_message(chat_id=chat_id, text=info_line)
            context.bot.send_message(
                chat_id=chat_id,
                text="آیا مدیرعامل خارج از سهامداران است؟",
                reply_markup=ceo_outside_keyboard(i)
            )
            return
    
        # سایر سمت‌ها (از جمله ترکیبی‌ها) → مستقیم برو سراغ حق‌امضا
        context.bot.send_message(chat_id=chat_id, text=info_line)
        context.bot.send_message(
            chat_id=chat_id,
            text=f"وضعیت حق‌امضا برای «{person_name}» را انتخاب کنید:",
            reply_markup=sign_authority_keyboard(i)
        )
        return


    # --- پاسخ به سؤال «مدیرعامل خارج از سهامداران؟» ---
    if payload.startswith("ceo_out:"):
        parts = payload.split(":", 2)   # "ceo_out:{i}:{yes|no}"
        if len(parts) != 3:
            context.bot.send_message(chat_id=chat_id, text="دادهٔ مدیرعامل نامعتبر بود.")
            return
        _, idx_str, yn = parts
        try:
            i = int(idx_str)
        except ValueError:
            context.bot.send_message(chat_id=chat_id, text="شناسهٔ عضو نامعتبر بود.")
            return
    
        is_out = (yn == "yes")
        d[f"عضو {i} مدیرعامل بیرون سهامداران؟"] = is_out
    
        # ✅ اگر مدیرعامل خارج از سهامداران است → حداقل ۴ عضو لازم است
        if is_out:
            cnt_str = str(d.get("تعداد اعضای هیئت مدیره", "") or "0")
            total = int(fa_to_en_number(cnt_str))
            if total < 4:
                # پاک‌سازی تمام داده‌های اعضا + خودِ تعداد
                for j in range(1, total + 1):
                    for key in (
                        f"عضو {j} نام",
                        f"عضو {j} کد ملی",
                        f"عضو {j} سمت",
                        f"عضو {j} سمت کد",
                        f"عضو {j} حق‌امضا",
                        f"عضو {j} مدیرعامل بیرون سهامداران؟",
                    ):
                        d.pop(key, None)
                d.pop("تعداد اعضای هیئت مدیره", None)
                d["board_index"] = 1
                d["step"] = 7  # ← برگشت به سؤال «تعداد اعضای هیئت‌مدیره»
    
                warn = (
                    "❗️از آنجا که «مدیرعامل خارج از سهامداران» انتخاب کردید، باید مشخصات حداقل 4 نفر را وارد کنید (یعنی تعداد اعضای هیئت‌مدیره حداقل 3 نفر بعلاوه 1 نفر مدیرعامل خارج از اعضای هیئت‌مدیره) .\n"
                    "تعداد فعلی کافی نیست. لطفاً تعداد اعضای هیئت‌مدیره را حداقل 4 نفر انتخاب کنید:"
                )
                context.bot.send_message(chat_id=chat_id, text=warn, reply_markup=main_keyboard())
                if 'remember_last_question' in globals():
                    remember_last_question(context, "تعداد اعضای هیئت‌مدیره را وارد کنید (اعداد فارسی):")
                context.bot.send_message(chat_id=chat_id, text="تعداد اعضای هیئت‌مدیره را وارد کنید (اعداد فارسی):", reply_markup=main_keyboard())
                return
    
        # در غیر این صورت یا اگر شرط برقرار بود → ادامهٔ فلو: پرسش حق‌امضا برای همین عضو
        person_name = d.get(f"عضو {i} نام", "")
        context.bot.send_message(
            chat_id=chat_id,
            text=f"وضعیت حق‌امضا برای «{person_name}» را انتخاب کنید:",
            reply_markup=sign_authority_keyboard(i)
        )
        return



    # --- حق‌امضا برای عضو i ---
    if payload.startswith("sig:"):
        parts = payload.split(":", 2)  # "sig:{i}:{b|n|bn|none}"
        if len(parts) != 3:
            context.bot.send_message(chat_id=chat_id, text="دادهٔ حق‌امضا نامعتبر بود."); return
        _, idx_str, choice = parts
        try:
            i = int(idx_str)
        except ValueError:
            context.bot.send_message(chat_id=chat_id, text="شناسهٔ عضو نامعتبر بود."); return
    
        if choice not in ("b", "n", "bn", "none"):   # ← گزینهٔ جدید
            context.bot.send_message(chat_id=chat_id, text="گزینهٔ حق‌امضا نامعتبر بود."); return
    
        d[f"عضو {i} حق‌امضا"] = choice
    
        total = int(fa_to_en_number(str(d.get("تعداد اعضای هیئت مدیره", 0)) or "0"))
        if i < total:
            d["board_index"] = i + 1
            fa_next = str(d["board_index"]).translate(str.maketrans("0123456789","۰۱۲۳۴۵۶۷۸۹"))
            label = f"نام عضو هیئت‌مدیره {fa_next} را وارد کنید (مثال: آقای ... / خانم ...):"
            if 'remember_last_question' in globals():
                remember_last_question(context, label)
            context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard())
        else:
            # --- پایان ورود حق‌امضا برای آخرین عضو ---
        
            # 1) حداقل یک امضاکننده برای «بهادار/تعهدآور» و حداقل یک امضاکننده برای «عادی/اداری»
            b_count = 0
            n_count = 0
            for j in range(1, total + 1):
                chj = d.get(f"عضو {j} حق‌امضا")
                if chj in ("b", "bn"):
                    b_count += 1
                if chj in ("n", "bn"):
                    n_count += 1
        
            if b_count < 1 or n_count < 1:
                # پاک‌سازی کامل اعضا + خودِ تعداد → بازگشت به سؤال «تعداد اعضای هیئت‌مدیره»
                for j in range(1, total + 1):
                    for key in (
                        f"عضو {j} نام",
                        f"عضو {j} کد ملی",
                        f"عضو {j} سمت",
                        f"عضو {j} سمت کد",
                        f"عضو {j} حق‌امضا",
                        f"عضو {j} مدیرعامل بیرون سهامداران؟",
                    ):
                        d.pop(key, None)
                d.pop("تعداد اعضای هیئت مدیره", None)
                d["board_index"] = 1
                d["step"] = 7
        
                warn = (
                    "❗️برای اعتبار صورتجلسه، باید حداقل یک امضاکننده برای «اوراق و اسناد بهادار و تعهد‌آور» "
                    "و حداقل یک امضاکننده برای «مکاتبات عادی و اداری» انتخاب شود.\n"
                    "اطلاعات اعضای هیئت‌مدیره پاک شد. لطفاً تعداد اعضای هیئت‌مدیره را دوباره وارد کنید (اعداد فارسی):"
                )
                context.bot.send_message(chat_id=chat_id, text=warn, reply_markup=main_keyboard())
                if 'remember_last_question' in globals():
                    remember_last_question(context, "تعداد اعضای هیئت‌مدیره را وارد کنید (اعداد فارسی):")
                context.bot.send_message(chat_id=chat_id, text="تعداد اعضای هیئت‌مدیره را وارد کنید (اعداد فارسی):", reply_markup=main_keyboard())
                return
        
            # 2) الزام وجودِ نقش‌ها: مدیرعامل + رئیس + نایب رئیس + عضو هیئت‌مدیره
            role_codes = []
            for j in range(1, total + 1):
                rc = d.get(f"عضو {j} سمت کد")
                if rc:
                    role_codes.append(rc)
        
            has_ceo    = any(rc in ("ceo", "ceo_chair", "ceo_vice", "ceo_member") for rc in role_codes)
            has_chair  = any(rc in ("chair", "ceo_chair") for rc in role_codes)
            has_vice   = any(rc in ("vice", "ceo_vice") for rc in role_codes)
            has_member = any(rc in ("member", "ceo_member") for rc in role_codes)
        
            if not (has_ceo and has_chair and has_vice and has_member):
                # پاک‌سازی کامل اعضا + خودِ تعداد → بازگشت به سؤال «تعداد اعضای هیئت‌مدیره»
                for j in range(1, total + 1):
                    for key in (
                        f"عضو {j} نام",
                        f"عضو {j} کد ملی",
                        f"عضو {j} سمت",
                        f"عضو {j} سمت کد",
                        f"عضو {j} حق‌امضا",
                        f"عضو {j} مدیرعامل بیرون سهامداران؟",
                    ):
                        d.pop(key, None)
                d.pop("تعداد اعضای هیئت مدیره", None)
                d["board_index"] = 1
                d["step"] = 7
        
                warn = (
                    "❗️ترکیب سمت‌ها ناقص است. باید حتماً «مدیرعامل»، «رئیس هیئت‌مدیره»، «نایب رئیس هیئت‌مدیره» و "
                    "«عضو هیئت‌مدیره» در میان اعضا انتخاب شوند.\n"
                    "نقش‌های ترکیبی که شامل هیئت‌مدیره هستند قابل قبول‌اند (مثلاً «مدیرعامل و رئیس هیئت‌مدیره»، "
                    "«مدیرعامل و نایب رئیس هیئت‌مدیره»، «مدیرعامل و عضو هیئت‌مدیره»).\n"
                    "اطلاعات اعضای هیئت‌مدیره پاک شد. :"
                )
                context.bot.send_message(chat_id=chat_id, text=warn, reply_markup=main_keyboard())
                if 'remember_last_question' in globals():
                    remember_last_question(context, "تعداد اعضای هیئت‌مدیره را وارد کنید (اعداد فارسی):")
                context.bot.send_message(chat_id=chat_id, text="تعداد اعضای هیئت‌مدیره را وارد کنید (اعداد فارسی):", reply_markup=main_keyboard())
                return
        
            # اگر هر دو شرط برقرار بود → ادامهٔ فلو (وکیل)
            d["step"] = 9
            label = get_label("وکیل") if 'get_label' in globals() else "نام وکیل را وارد کنید:"
            if 'remember_last_question' in globals():
                remember_last_question(context, label)
            context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard())
            return



    # فوروارد کردن بقیه payload ها به هندلرهای موجود (مثل روزنامه و ...)
    if payload.startswith("newspaper:"):
        handle_newspaper_choice(update, context)
        return

    if payload == AI_RESUME:
        resume_from_ai(update, context)
        return

    # اگر otp دکمه‌ای داری:
    try:
        otp_buttons_handler(update, context)
    except Exception:
        pass




def handle_message(update: Update, context: CallbackContext):
    try:
        # دیباگ:
        print("DBG: handle_message got message text:", getattr(update.message, "text", None))
        
        chat_id = update.message.chat_id
        text = (update.message.text or "").strip()
        data = user_data.setdefault(chat_id, {"step": 0})
    
        # --- گارد حالت AI: ابتدای تابع ---
        if context.user_data.get("ai_mode"):
            # اگر دکمه برگشت منویی داری، همین‌جا هندل کن (اختیاری)
            if text == BACK_BTN:
                context.user_data["ai_mode"] = False
                send_ai_services_menu(chat_id, context)
                return
    
            handle_ai_text(update, context)
            if not context.user_data.get("ai_mode"):
                return
            return

        # ========== گارد شماره موبایل (اولویت قبل از هر چیز) ==========
        # اگر در وضعیت انتظار شماره هستیم، فقط شماره را پردازش کن:
        if context.user_data.get("awaiting") == "phone":
            m = re.search(r"[۰-۹0-9]{10,}", (update.message.text or ""))
            if m:
                phone = set_user_phone(update.effective_chat.id, m.group(0), meta={
                    "first_name": getattr(update.message.from_user, "first_name", ""),
                    "last_name": getattr(update.message.from_user, "last_name", ""),
                    "username": getattr(update.message.from_user, "username", "")
                })
                confirm_phone_and_continue(update.effective_chat.id, context, phone)
                return

            context.bot.send_message(update.effective_chat.id,
                "شماره معتبر پیدا نشد. لطفاً با دکمه زیر شماره موبایل را بفرستید.",
                reply_markup=request_phone_keyboard())
            return

        if context.user_data.get("awaiting_phone"):
            # اگر کاربر Contact فرستاد
            if update.message.contact and update.message.contact.phone_number:
                phone_raw = update.message.contact.phone_number
                phone = set_user_phone(chat_id, phone_raw, meta={
                    "first_name": getattr(update.message.from_user, "first_name", ""),
                    "last_name": getattr(update.message.from_user, "last_name", ""),
                    "username": getattr(update.message.from_user, "username", "")
                })
                if phone:
                    confirm_phone_and_continue(chat_id, context, phone)
                    return
                else:
                    context.bot.send_message(
                        chat_id=chat_id,
                        text="❗️شماره معتبر نیست. لطفاً دوباره ارسال کنید.",
                        reply_markup=phone_request_keyboard()
                    )
                    return

            # اگر کاربر شماره را تایپ کرد
            if text and is_valid_phone_text(text):
                phone = set_user_phone(chat_id, text, meta={
                    "first_name": getattr(update.message.from_user, "first_name", ""),
                    "last_name": getattr(update.message.from_user, "last_name", ""),
                    "username": getattr(update.message.from_user, "username", "")
                })
                confirm_phone_and_continue(chat_id, context, phone)
                return

            # ورودی نامعتبر
            context.bot.send_message(
                chat_id=chat_id,
                text="❗️لطفاً شماره معتبر وارد کنید (مثال: 09xxxxxxxxx) یا دکمه «ارسال شماره من» را بزنید.",
                reply_markup=phone_request_keyboard()
            )
            return
      
        print("DBG: get_user_phone result:", get_user_phone(chat_id))
        print("DBG: context.user_data phone:", context.user_data.get("phone"))

        # اگر هنوز شماره ثبت نشده، درخواست شماره بده و جلوی ادامه‌ی فلو را بگیر:
        if not get_user_phone(chat_id):
            ask_for_phone(chat_id, context)
            return
        # ============================================================

        # اگر کاربر دکمه بازگشت زد
        if text == BACK_BTN:
            handle_back(update, context)
            return

        if text == AI_ASK_TEXT:
            enter_ai_mode_reply(update, context)  # همین تابع خودت که ai_mode را True می‌کند
            return


        # --- AI Landing Options ---
        if text in (AI_OPT_MINUTES, AI_OPT_QA, AI_OPT_COMP_TYPE, AI_OPT_NAME, AI_OPT_CONTRACT, AI_OPT_FORMAL):
            if text == AI_OPT_MINUTES:
                # تنظیم وضعیت برای ورود به فلو صورتجلسه
                
                data["step"] = 0
                data.pop("موضوع صورتجلسه", None)
                send_topic_menu(chat_id, context)
                return

            # اگر کاربر گزینه «پیشنهاد هوشمند نام شرکت» را انتخاب کرد
            if text == AI_OPT_NAME:
               
                data["ai_mode"] = "name_suggestion"
                data["step"] = 1
                context.bot.send_message(
                    chat_id=chat_id,
                    text="🧩 لطفاً بفرمایید چه کلمه‌ی اصلی برای نام شرکت در نظر دارید؟\n(مثلاً: آتی، پارس، نیک، آراد...)",
                    reply_markup=back_keyboard()
                )
                return

            # نگاشت مستقیم: «مشاوره …» ≡ «سؤال دارم» + حذف بک اینلاین
            if text == AI_OPT_QA:
                context.user_data["ai_skip_inline_back"] = True  # اگر نمی‌خواهی اینلاین‌بکِ صورتجلسه نمایش داده شود
                enter_ai_mode_reply(update, context)
                return

                
            # === شروع فلو «راهنمای انتخاب نوع شرکت» ===
            if text == AI_OPT_COMP_TYPE:
                data["ai_mode"] = "comp_type"
                data["step"] = 1
                # پاکسازی پاسخ‌های قبلی این راهنما
                for k in ["CT_تعداد_شرکا","CT_سرمایه","CT_مناقصات","CT_هیئت_مدیره","CT_انتقال"]:
                    data.pop(k, None)
        
                context.bot.send_message(
                    chat_id=chat_id,
                    text="برای شروع راهنما، تعداد مؤسسین/شرکا را انتخاب کنید:",
                    reply_markup=small_keyboard()
                )
                return


            
            # === شروع فلو «تبدیل متن ساده به متن رسمی/حقوقی» ===
            if text == AI_OPT_FORMAL:
                chat_id = update.effective_chat.id
                user_data.setdefault(chat_id, {})
                data = user_data[chat_id]
            
                context.user_data["ai_mode"] = "formalizer"
                data["step"] = 1
                for k in ["FORMAL_RAW", "FORMAL_STYLE"]:
                    data.pop(k, None)
            
                label = (
                    "📝 لطفاً متن ساده‌تان را ارسال کنید.\n"
                    "مثال: «یه متن می‌خوام برای اعلام تغییر ساعت کاری شرکت به اداره ثبت» یا متن کامل بند/نامه.\n\n"
                    "نکته: اطلاعات حقیقی/حقوقی موجود را کامل بنویسید تا متن رسمی دقیق تولید شود."
                )
                context.bot.send_message(chat_id=chat_id, text=label, reply_markup=back_keyboard())
                return



            
        
            # باقی گزینه‌ها که هنوز آماده نشده‌اند
            pending_map = {
                AI_OPT_CONTRACT:  "📝 «تولید قرارداد آماده» به‌زودی فعال می‌شود.",
                AI_OPT_FORMAL:    "✍️ «تبدیل متن ساده به متن رسمی/حقوقی» به‌زودی فعال می‌شود.",
            }
            context.bot.send_message(chat_id=chat_id, text=pending_map.get(text, "به‌زودی…"))
            send_ai_services_menu(chat_id, context)
            return

        

        # -------------------------------
        # فلو هوش مصنوعی: پیشنهاد نام شرکت (Groq API)
        # -------------------------------
        if data.get("ai_mode") == "name_suggestion":
            step = data.get("step", 0)
        
            # گام ۱: دریافت کلمه اصلی
            if step == 1:
                data["کلمه اصلی"] = text
                data["step"] = 2
                context.bot.send_message(
                    chat_id=chat_id,
                    text="💼 موضوع فعالیت شرکت در چه حوزه‌ای است؟\n(مثلاً: بازرگانی، فناوری، ساخت‌وساز، خدمات...)",
                    reply_markup=back_keyboard()
                )
                return
        
            # گام ۲: دریافت حوزه فعالیت و تولید نام‌های پیشنهادی
            if step == 2:
                data["حوزه فعالیت"] = text
        
                # آماده‌سازی پرامپت برای Groq
                user_prompt = (
                    f"پنج نام پیشنهادی سه‌ کلمه ای یا سه بخشی برای شرکت بساز که در آن از کلمه‌ی '{data['کلمه اصلی']}' استفاده شده باشد "
                    f"و حوزه‌ی فعالیت شرکت '{data['حوزه فعالیت']}' است. "
                    "نام‌ها باید فارسی، خوش‌آوا، معنادار و قابل ثبت باشند. "
                    "خروجی را فقط به صورت فهرست شماره‌دار بنویس."
                )
        
                system_prompt = (
                    "شما یک متخصص در نام‌گذاری شرکت‌های ایرانی هستید. "
                    "نام‌های پیشنهادی باید ریشه فارسی داشته باشد و کاملاً معنادار و حرفه‌ای باشند، دست‌کم دارای یک اسم خاص باشد، از ترکیب کلمه‌ی کاربر با مفاهیم مرتبط با حوزه‌ی فعالیت، از این کلمات استفاده نکن :(ملی، ایران، ملت، کشور، انتظام، نظام، نفت، پتروشیمی، دادگستر، بنیاد، سازمان، مرکز، بهزیست، بانک، بسیج، جهاد، آسانسور، آژانس، ایثار، ایثارگران، شاهد، شهید، آزاده، جانبازان، تکنو، فامیلی، فیلتر، نیک، باور، میکروبرد، تست، استار، تک، مدرن، پیک، امنیت، کارواش، فانتزی، شهروند، اقتصاد، میهن، جوانان، ایمان، دکوراسیون، هدف، سپاه، تکنیک، دیتا، تک نو، تکنولوژی، اورست، مونتاژ، ونوس، دفتر، ایتال، بیو، کنترل،‌ متریک، مترلژی، وب)،برخی کلمات در نام شرکت‌ها جزو سیلاب محسوب نمی‌شوند. هرچند استفاده از این واژه‌ها در اسم شرکت مانعی ندارد اما جزئی از نام شرکت شمرده نمی‌شوند و در نام‌گذاری باید به این نکته توجه کرد که اسم شخصیت حقوقی غیر از این کلمات شامل حداقل سه سیلاب (سه کلمه) باشد:
صفت‌هایی نظیرِ برتر، خوب، نیک و…
نام شهرها، رنگ‌ها و اعداد مانندِ تهران، فیروزه‌ای، چهل و…
 کلماتی که بیانگر موضوع فعالیت هستند، مثلِ ساختمانی، حقوقی، بازرگانی، صنعتی، خدماتی، توسعه، تجاری، مهندسی، مهندسی مشاور، فنی مهندسی، تولیدی و… "
                )
        
                # فراخوانی هوش مصنوعی متصل به Groq
                try:
                    # ❌ از این استفاده نکن: from ai_module import ask_groq
                    response = ask_groq(user_prompt, system_prompt, max_tokens=300)
                except Exception as e:
                    # لاگ دقیق تا بفهمی ایراد چیه
                    import traceback, sys
                    print("ERR: ask_groq failed:", e, file=sys.stderr)
                    traceback.print_exc()
                    response = f"❗️خطا در ارتباط با هوش مصنوعی:\n{e}"

        
                # ارسال نتیجه
                context.bot.send_message(
                    chat_id=chat_id,
                    text=f"🤖 بر اساس پاسخ شما، پنج نام پیشنهادی:\n\n{response}",
                    reply_markup=ai_services_keyboard()  # بازگشت به منوی اصلی AI
                )
        
                # پاکسازی حالت
                data.pop("ai_mode", None)
                data["step"] = 0
                return

        # -------------------------------
        # فلو هوش مصنوعی: راهنمای انتخاب نوع شرکت (Groq API)
        # -------------------------------

        if data.get("ai_mode") == "comp_type":
            step = int(data.get("step", 1))
        
            # دکمه بازگشت
            if text == BACK_BTN:
                prev = max(1, step - 1)
                data["step"] = prev
                if prev == 1:
                    context.bot.send_message(chat_id=chat_id, text="برای شروع راهنما، تعداد مؤسسین/شرکا را انتخاب کنید:", reply_markup=small_keyboard()); return
                if prev == 2:
                    context.bot.send_message(chat_id=chat_id, text="حدود سرمایه‌ی اولیه (به ریال) چقدر است؟ (فقط عدد، مثل 200000000)", reply_markup=back_keyboard()); return
                if prev == 3:
                    context.bot.send_message(chat_id=chat_id, text="برنامه شرکت برای «مناقصات/مزایدات بزرگ و اعتبار بانکی» پررنگ است؟", reply_markup=yes_no_keyboard()); return
                if prev == 4:
                    context.bot.send_message(chat_id=chat_id, text="هیئت‌مدیره 3 نفره می‌خواهید یا مدیریت ساده کافیست؟", reply_markup=board_need_keyboard()); return
                if prev == 5:
                    context.bot.send_message(chat_id=chat_id, text="سهولت نقل‌وانتقال مالکیت چقدر مهم است؟", reply_markup=transfer_need_keyboard()); return
        
            # STEP 1: تعداد شرکا
            if step == 1:
                valid = ["1 نفر","2 نفر","3 نفر","4 نفر یا بیشتر"]
                if text not in valid:
                    context.bot.send_message(chat_id=chat_id, text="لطفاً از گزینه‌ها انتخاب کنید.", reply_markup=small_keyboard()); return
                data["CT_تعداد_شرکا"] = text
                data["step"] = 2
                context.bot.send_message(chat_id=chat_id, text="حدود سرمایه‌ی اولیه (به ریال) چقدر است؟ (فقط عدد، مثل 200000000)", reply_markup=back_keyboard())
                return
        
            # STEP 2: سرمایه
            if step == 2:
                digits = "".join(ch for ch in text if ch.isdigit())
                if not digits:
                    context.bot.send_message(chat_id=chat_id, text="❗️یک عدد وارد کنید (مثل 200000000).", reply_markup=back_keyboard()); return
                data["CT_سرمایه"] = int(digits)
                data["step"] = 3
                context.bot.send_message(chat_id=chat_id, text="برنامه شرکت برای «مناقصات/مزایدات بزرگ و اعتبار بانکی» پررنگ است؟", reply_markup=yes_no_keyboard())
                return
        
            # STEP 3: مناقصات/اعتبار
            if step == 3:
                if text not in ["بله ✅","خیر ❌"]:
                    context.bot.send_message(chat_id=chat_id, text="یک گزینه را انتخاب کنید.", reply_markup=yes_no_keyboard()); return
                data["CT_مناقصات"] = (text == "بله ✅")
                data["step"] = 4
                context.bot.send_message(chat_id=chat_id, text="هیئت‌مدیره 3 نفره می‌خواهید یا مدیریت ساده کافیست؟", reply_markup=board_need_keyboard())
                return
        
            # STEP 4: هیئت‌مدیره
            if step == 4:
                if text not in ["هیئت‌مدیره 3 نفره می‌خواهم","مدیریت ساده/یک‌نفره کافیست"]:
                    context.bot.send_message(chat_id=chat_id, text="از بین دو گزینه انتخاب کنید.", reply_markup=board_need_keyboard()); return
                data["CT_هیئت_مدیره"] = (text == "هیئت‌مدیره 3 نفره می‌خواهم")
                data["step"] = 5
                context.bot.send_message(chat_id=chat_id, text="سهولت نقل‌وانتقال مالکیت چقدر مهم است؟", reply_markup=transfer_need_keyboard())
                return
        
            # STEP 5: سهولت نقل‌وانتقال
            if step == 5:
                if text not in ["خیلی مهم است","اهمیت متوسط","اهمیت ندارد"]:
                    context.bot.send_message(chat_id=chat_id, text="از بین گزینه‌ها انتخاب کنید.", reply_markup=transfer_need_keyboard()); return
                data["CT_انتقال"] = text
        
                # محاسبه نتیجه
                result_text = build_comp_type_recommendation(data)
        
                # خروج از AI این سناریو (فقط همین سناریو)
                data["ai_mode"] = None
                data["step"] = 0
        
                context.bot.send_message(chat_id=chat_id, text=result_text, reply_markup=main_keyboard())
                send_ai_services_menu(chat_id, context)  # اگر می‌خواهی بعدش دوباره منوی AI نشان داده شود
                return




        step = data.get("step", 0)
    
        موضوع = data.get("موضوع صورتجلسه")
        نوع_شرکت = data.get("نوع شرکت")
    
        if "موضوع صورتجلسه" not in data and not context.user_data.get("ai_mode"):
            context.bot.send_message(
                chat_id=chat_id,
                text="لطفاً ابتدا موضوع صورتجلسه را انتخاب کنید. برای شروع مجدد /start را ارسال کنید .",
                reply_markup=main_keyboard()
            )
            return

        
        # ===== ادامه‌ی منطق‌های قبلی شما از اینجا به بعد =====
        # ...

    
        # تعریف فیلدهای پایه برای تغییر آدرس مسئولیت محدود (در صورت نیاز)
        common_fields = ["نام شرکت", "شماره ثبت", "شناسه ملی", "سرمایه", "تاریخ", "ساعت", "آدرس جدید", "کد پستی", "وکیل"]
    
        # -------------------------------
        # تغییر نام شرکت - سهامی خاص
        # گام‌ها: 1 نام شرکت، 2 ثبت، 3 شناسه، 4 سرمایه، 5 تاریخ، 6 ساعت،
        # 7 مدیر عامل، 8 نایب رییس، 9 رییس، 10 منشی،
        # 11 نام جدید شرکت، 12 وکیل → خروجی
        # -------------------------------
        if موضوع == "تغییر نام شرکت" and نوع_شرکت == "سهامی خاص":
            if step == 1:
                data["نام شرکت"] = text
                data["step"] = 2
                label = get_label("شماره ثبت")
                remember_last_question(context, label)
                context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard())
                return
    
            if step == 2:
                if not is_persian_number(text):
                    context.bot.send_message(chat_id=chat_id, text="❗️شماره ثبت را فقط با اعداد فارسی وارد کنید.", reply_markup=main_keyboard())
                    return
                data["شماره ثبت"] = text
                data["step"] = 3
                label = get_label("شناسه ملی")
                remember_last_question(context, label)
                context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard())
                return
    
            if step == 3:
                if not is_persian_number(text):
                    context.bot.send_message(chat_id=chat_id, text="❗️شناسه ملی را فقط با اعداد فارسی وارد کنید.", reply_markup=main_keyboard())
                    return
                data["شناسه ملی"] = text
                data["step"] = 4
                label = get_label("سرمایه")
                remember_last_question(context, label)
                context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard())
                return
    
            if step == 4:
                if not is_persian_number(text):
                    context.bot.send_message(chat_id=chat_id, text="❗️سرمایه را فقط با اعداد فارسی وارد کنید.", reply_markup=main_keyboard())
                    return
                data["سرمایه"] = text
                data["step"] = 5
                label = get_label("تاریخ")
                remember_last_question(context, label)
                context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard())
                return
    
            if step == 5:
                if 'is_valid_persian_date' in globals():
                    if not is_valid_persian_date(text):
                        context.bot.send_message(chat_id=chat_id, text="❗️فرمت تاریخ صحیح نیست. نمونه: ۱۴۰۴/۰۵/۱۵", reply_markup=main_keyboard())
                        return
                else:
                    if text.count('/') != 2:
                        context.bot.send_message(chat_id=chat_id, text="❗️فرمت تاریخ صحیح نیست.", reply_markup=main_keyboard())
                        return
                data["تاریخ"] = text
                data["step"] = 6
                label = get_label("ساعت")
                remember_last_question(context, label)
                context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard())
                return
    
            if step == 6:
                if not is_persian_number(text):
                    context.bot.send_message(chat_id=chat_id, text="❗️ساعت را فقط با اعداد فارسی وارد کنید.", reply_markup=main_keyboard())
                    return
                data["ساعت"] = text
                data["step"] = 7
                label = get_label("مدیر عامل")
                remember_last_question(context, label)
                context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard())
                return
    
            if step == 7:
                data["مدیر عامل"] = text
                data["step"] = 8
                label = get_label("نایب رییس")
                remember_last_question(context, label)
                context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard())
                return
    
            if step == 8:
                data["نایب رییس"] = text
                data["step"] = 9
                label = get_label("رییس")
                remember_last_question(context, label)
                context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard())
                return
    
            if step == 9:
                data["رییس"] = text
                data["step"] = 10
                label = get_label("منشی")
                remember_last_question(context, label)
                context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard())
                return
    
            if step == 10:
                data["منشی"] = text
                data["step"] = 11
                label = get_label("نام جدید شرکت")
                remember_last_question(context, label)
                context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard())
                return
    
            if step == 11:
                data["نام جدید شرکت"] = text
                data["step"] = 12
                label = get_label("وکیل")
                remember_last_question(context, label)
                context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard())
                return
    
            if step == 12:
                data["وکیل"] = text
                send_summary(chat_id, context)
                data["step"] = 13
                return
    
            if step >= 13:
                context.bot.send_message(chat_id=chat_id, text="✅ اطلاعات ثبت شد. برای شروع مجدد /start را ارسال کنید.")
                return



        # -------------------------------
        # انتخاب مدیران - سهامی خاص
        # -------------------------------
        if data.get("موضوع صورتجلسه") == "انتخاب مدیران" and data.get("نوع شرکت") == "سهامی خاص":
            if step == 1:
                data["نام شرکت"] = text
                data["step"] = 2
                label = get_label("شماره ثبت")
                remember_last_question(context, label)
                context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard()); return
        
            if step == 2:
                if not is_persian_number(text):
                    context.bot.send_message(chat_id=chat_id, text="❗️شماره ثبت را فقط با اعداد فارسی وارد کنید.", reply_markup=main_keyboard()); return
                data["شماره ثبت"] = text
                data["step"] = 3
                label = get_label("شناسه ملی")
                remember_last_question(context, label)
                context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard()); return
        
            if step == 3:
                if not is_persian_number(text):
                    context.bot.send_message(chat_id=chat_id, text="❗️شناسه ملی را فقط با اعداد فارسی وارد کنید.", reply_markup=main_keyboard()); return
                data["شناسه ملی"] = text
                data["step"] = 4
                label = get_label("سرمایه")
                remember_last_question(context, label)
                context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard()); return
        
            if step == 4:
                if not is_persian_number(text):
                    context.bot.send_message(chat_id=chat_id, text="❗️سرمایه را فقط با اعداد فارسی وارد کنید.", reply_markup=main_keyboard()); return
                data["سرمایه"] = text
                data["step"] = 5
                label = get_label("تاریخ")
                remember_last_question(context, label)
                context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard()); return
        
            if step == 5:
                if not is_valid_persian_date(text):
                    context.bot.send_message(chat_id=chat_id, text="❗️فرمت تاریخ صحیح نیست. نمونه: ۱۴۰۴/۰۵/۱۵", reply_markup=main_keyboard()); return
                data["تاریخ"] = text
                data["step"] = 6
                label = get_label("ساعت")
                remember_last_question(context, label)
                context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard()); return
        
            if step == 6:
                if not is_persian_number(text):
                    context.bot.send_message(chat_id=chat_id, text="❗️ساعت را فقط با اعداد فارسی وارد کنید.", reply_markup=main_keyboard()); return
                data["ساعت"] = text
                data["step"] = 7
                label = "تعداد اعضای هیئت‌مدیره را وارد کنید (اعداد فارسی):"
                remember_last_question(context, label)
                context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard()); return
        
            if step == 7:
                if not is_persian_number(text):
                    context.bot.send_message(chat_id=chat_id, text="❗️عدد فارسی وارد کنید.", reply_markup=main_keyboard()); return
            
                count = int(fa_to_en_number(text))
                if count < 3:
                    context.bot.send_message(chat_id=chat_id, text="❗️حداقل سه عضو لازم است.", reply_markup=main_keyboard()); return
            
                data["تعداد اعضای هیئت مدیره"] = count
                data["board_index"] = 1
                data["step"] = 8
            
                fa1 = "1".translate(str.maketrans("0123456789", "۰۱۲۳۴۵۶۷۸۹"))
                label = f"نام عضو هیئت‌مدیره {fa1} را وارد کنید (مثال: آقای ... / خانم ...):"
                if 'remember_last_question' in globals():
                    remember_last_question(context, label)

                # 1) پیام اطلاع‌رسانی همزمان (در پیام جداگانه)
                context.bot.send_message(
                    chat_id=chat_id,
                    text=(
                        "اعضای هیئت‌مدیره الزاماً باید از میان سهامداران انتخاب شوند.\n"
                        "مدیرعامل لزوماً سهامدار نیست، اما اعضای هیئت‌مدیره باید سهامدار باشند."
                    )
                )
                
                # 2) سوال بعدی (نام عضو ۱)
                context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard())
                return

        
            # حلقه اعضای هیئت‌مدیره (نام → کدملی → انتخاب سمت (دکمه) → اگر ceo سوال اضافه → حق‌امضا (دکمه))
            if step == 8:
                i = data.get("board_index", 1)
                fa_i = str(i).translate(str.maketrans("0123456789","۰۱۲۳۴۵۶۷۸۹"))
                prefix = f"عضو {i}"
        
                if f"{prefix} نام" not in data:
                    data[f"{prefix} نام"] = text
                    label = f"کد ملی عضو هیئت‌مدیره {fa_i} را وارد کنید (اعداد فارسی):"
                    remember_last_question(context, label)
                    context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard()); return
        
                if f"{prefix} کد ملی" not in data:
                    if not is_persian_number(text):
                        context.bot.send_message(chat_id=chat_id, text="❗️کد ملی را فقط با اعداد فارسی وارد کنید.", reply_markup=main_keyboard()); return
                    data[f"{prefix} کد ملی"] = text
                    # حالا انتخاب سمت با دکمه‌ها
                    context.bot.send_message(chat_id=chat_id,
                                             text=f"سمت «{data.get(f'{prefix} نام','')}» را انتخاب کنید:",
                                             reply_markup=roles_keyboard(i))
                    return
        
                # بعد از انتخاب سمت، callback نقش → اگر ceo باشد سؤال اضافه → سپس callback حق‌امضا
                # بعد از ذخیره حق‌امضا در callback، یا به عضو بعدی می‌رویم یا به مرحله وکیل.
        
            if step == 9:
                data["وکیل"] = text
                # ساخت و ارسال خروجی
                text_out = render_board_election_text(data)
                try:
                    # پیام
                    for ofs in range(0, len(text_out), 3500):
                        context.bot.send_message(chat_id=chat_id, text=text_out[ofs:ofs+3500], reply_markup=main_keyboard())
                    # فایل Word
                    file_path = generate_word_file(text_out)
                    with open(file_path, 'rb') as f:
                        context.bot.send_document(chat_id=chat_id, document=f, filename="صورتجلسه انتخاب مدیران.docx")
                    os.remove(file_path)
                except Exception as e:
                    context.bot.send_message(chat_id=chat_id, text=f"❗️خطا در ساخت/ارسال فایل: {e}", reply_markup=main_keyboard())
                data["step"] = 10
                return
        
            if step >= 10:
                context.bot.send_message(chat_id=chat_id, text="✅ اطلاعات ثبت شد. برای شروع مجدد /start را ارسال کنید.", reply_markup=main_keyboard())
                return


        
        # تعریف فیلدهای پایه برای تغییر آدرس مسئولیت محدود
        common_fields = ["نام شرکت", "شماره ثبت", "شناسه ملی", "سرمایه", "تاریخ", "ساعت", "آدرس جدید", "کد پستی", "وکیل"]
    
        # -------------------------------
        # تغییر آدرس - مسئولیت محدود
        # -------------------------------
        if data.get("موضوع صورتجلسه") == "تغییر آدرس" and data.get("نوع شرکت") == "مسئولیت محدود":
            if step == 1:
                data["نام شرکت"] = text
                data["step"] = 2
                label = "شماره ثبت شرکت را وارد کنید:"
                remember_last_question(context, label)
                context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard())
                return
    
            if 2 <= step <= 9:
                field = common_fields[step - 1]
    
                if field == "تاریخ":
                    if text.count('/') != 2:
                        context.bot.send_message(chat_id=chat_id, text="❗️فرمت تاریخ صحیح نیست. لطفاً به صورت ۱۴۰۴/۰۴/۰۷ وارد کنید (با دو /).", reply_markup=main_keyboard())
                        return
    
                if field in persian_number_fields:
                    if not is_persian_number(text):
                        context.bot.send_message(chat_id=chat_id, text=f"لطفاً مقدار '{field}' را فقط با اعداد فارسی وارد کنید.", reply_markup=main_keyboard())
                        return
    
                data[field] = text
                data["step"] += 1
    
                if step == 9:
                    label = "تعداد شرکا را وارد کنید (بین ۲ تا ۷):"
                    remember_last_question(context, label)
                    context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard())
                    return
                else:
                    next_field = common_fields[step]
                    label = get_label(next_field)
                    remember_last_question(context, label)
                    context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard())
                    return
    
            if step == 10:
                if not text.isdigit():
                    context.bot.send_message(chat_id=chat_id, text="❗️لطفاً تعداد شرکا را فقط با عدد وارد کنید (بین ۲ تا ۷).", reply_markup=main_keyboard())
                    return
                count = int(text)
                if count < 2 or count > 7:
                    context.bot.send_message(chat_id=chat_id, text="❗️تعداد شرکا باید بین ۲ تا ۷ باشد. لطفاً مجدداً وارد کنید.", reply_markup=main_keyboard())
                    return
                data["تعداد شرکا"] = count
                data["step"] += 1
                data["current_partner"] = 1
                label = "نام شریک شماره ۱ را وارد کنید:"
                remember_last_question(context, label)
                context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard())
                return
    
            if step > 10:
                current_partner = data.get("current_partner", 1)
                count = data.get("تعداد شرکا", 0)
    
                if f"شریک {current_partner}" not in data:
                    data[f"شریک {current_partner}"] = text
                    label = f"میزان سهم الشرکه شریک شماره {current_partner} را به ریال وارد کنید (عدد فارسی):"
                    remember_last_question(context, label)
                    context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard())
                    return
                elif f"سهم الشرکه شریک {current_partner}" not in data:
                    if not is_persian_number(text):
                        context.bot.send_message(chat_id=chat_id, text="❗️لطفاً میزان سهم الشرکه را فقط با اعداد فارسی وارد کنید.", reply_markup=main_keyboard())
                        return
                    data[f"سهم الشرکه شریک {current_partner}"] = text
                    if current_partner < count:
                        data["current_partner"] = current_partner + 1
                        label = f"نام شریک شماره {current_partner + 1} را وارد کنید:"
                        remember_last_question(context, label)
                        context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard())
                        return
                    else:
                        send_summary(chat_id, context)
                        data["step"] = 11
                        return
    
            if step >= 11:
                context.bot.send_message(chat_id=chat_id, text="✅ اطلاعات قبلاً ثبت شده است. برای شروع مجدد /start را ارسال کنید.", reply_markup=main_keyboard())
                return

        # -------------------------------
        # تمدید سمت اعضا - سهامی خاص (داینامیک هیئت‌مدیره + سهامداران)
        # -------------------------------
        if موضوع == "تمدید سمت اعضا" and نوع_شرکت == "سهامی خاص":
            if step == 1:
                data["نام شرکت"] = text
                data["step"] = 2
                label = get_label("شماره ثبت")
                remember_last_question(context, label)
                context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard())
                return
        
            if step == 2:
                if not is_persian_number(text):
                    context.bot.send_message(chat_id=chat_id, text="❗️شماره ثبت را فقط با اعداد فارسی وارد کنید.", reply_markup=main_keyboard())
                    return
                data["شماره ثبت"] = text
                data["step"] = 3
                label = get_label("شناسه ملی")
                remember_last_question(context, label)
                context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard())
                return
        
            if step == 3:
                if not is_persian_number(text):
                    context.bot.send_message(chat_id=chat_id, text="❗️شناسه ملی را فقط با اعداد فارسی وارد کنید.", reply_markup=main_keyboard())
                    return
                data["شناسه ملی"] = text
                data["step"] = 4
                label = get_label("سرمایه")
                remember_last_question(context, label)
                context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard())
                return
        
            if step == 4:
                if not is_persian_number(text):
                    context.bot.send_message(chat_id=chat_id, text="❗️سرمایه را فقط با اعداد فارسی وارد کنید.", reply_markup=main_keyboard())
                    return
                data["سرمایه"] = text
                data["step"] = 5
                label = get_label("تاریخ")
                remember_last_question(context, label)
                context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard())
                return
        
            if step == 5:
                if 'is_valid_persian_date' in globals():
                    if not is_valid_persian_date(text):
                        context.bot.send_message(chat_id=chat_id, text="❗️فرمت تاریخ صحیح نیست. نمونه: ۱۴۰۴/۰۵/۱۵", reply_markup=main_keyboard())
                        return
                else:
                    if text.count('/') != 2:
                        context.bot.send_message(chat_id=chat_id, text="❗️فرمت تاریخ صحیح نیست.", reply_markup=main_keyboard())
                        return
                data["تاریخ"] = text
                data["step"] = 6
                label = get_label("ساعت")
                remember_last_question(context, label)
                context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard())
                return
        
            if step == 6:
                if not is_persian_number(text):
                    context.bot.send_message(chat_id=chat_id, text="❗️ساعت را فقط با اعداد فارسی وارد کنید.", reply_markup=main_keyboard())
                    return
                data["ساعت"] = text
                data["step"] = 7
                label = "نام مدیرعامل را وارد کنید (مثال: آقای ... / خانم ...):"
                remember_last_question(context, label)
                context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard())
                return
        
            if step == 7:
                data["مدیر عامل"] = text
                data["step"] = 8
                label = "نام نایب‌رییس (ناظر ۱) را وارد کنید (مثال: آقای ... / خانم ...):"
                remember_last_question(context, label)
                context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard())
                return
        
            if step == 8:
                data["نایب رییس"] = text
                data["step"] = 9
                label = "نام رییس (ناظر ۲) را وارد کنید (مثال: آقای ... / خانم ...):"
                remember_last_question(context, label)
                context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard())
                return
        
            if step == 9:
                data["رییس"] = text
                data["step"] = 10
                label = "نام منشی جلسه را وارد کنید (مثال: آقای ... / خانم ...):"
                remember_last_question(context, label)
                context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard())
                return
        
            if step == 10:
                data["منشی"] = text
                data["step"] = 11
                label = "تعداد اعضای هیئت‌مدیره را وارد کنید (اعداد فارسی):"
                remember_last_question(context, label)
                context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard())
                return
        
            # دریافت تعداد اعضای هیئت‌مدیره → حلقه نام/کدملی
            if step == 11:
                if not is_persian_number(text):
                    context.bot.send_message(chat_id=chat_id, text="❗️تعداد اعضای هیئت‌مدیره را فقط با اعداد فارسی وارد کنید.", reply_markup=main_keyboard())
                    return
                count = int(fa_to_en_number(text))
                if count < 1:
                    context.bot.send_message(chat_id=chat_id, text="❗️حداقل یک عضو لازم است.", reply_markup=main_keyboard())
                    return
                data["تعداد اعضای هیئت مدیره"] = count
                data["عضو_index"] = 1
                data["step"] = 12
                fa1 = "1".translate(str.maketrans("0123456789", "۰۱۲۳۴۵۶۷۸۹"))  # ۱
                label = f"نام عضو هیئت‌مدیره {fa1} را وارد کنید (مثال: آقای ... / خانم ...):"
                remember_last_question(context, label)
                context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard())
                return
        
            # حلقه اعضای هیئت‌مدیره: step == 12
            if step == 12:
                i = data.get("عضو_index", 1)
                fa_i = str(i).translate(str.maketrans("0123456789", "۰۱۲۳۴۵۶۷۸۹"))
                prefix = f"عضو {i}"
                if f"{prefix} نام" not in data:
                    data[f"{prefix} نام"] = text
                    label = f"کد ملی عضو هیئت‌مدیره {fa_i} را وارد کنید (اعداد فارسی):"
                    remember_last_question(context, label)
                    context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard())
                    return
                elif f"{prefix} کد ملی" not in data:
                    if not is_persian_number(text):
                        context.bot.send_message(chat_id=chat_id, text="❗️کد ملی را فقط با اعداد فارسی وارد کنید.", reply_markup=main_keyboard())
                        return
                    data[f"{prefix} کد ملی"] = text
                    total = data["تعداد اعضای هیئت مدیره"]
                    if i < total:
                        data["عضو_index"] = i + 1
                        fa_next = str(i+1).translate(str.maketrans("0123456789", "۰۱۲۳۴۵۶۷۸۹"))
                        label = f"نام عضو هیئت‌مدیره {fa_next} را وارد کنید (مثال: آقای ... / خانم ...):"
                        remember_last_question(context, label)
                        context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard())
                        return
                    else:
                        data["step"] = 13
                        label = "نام بازرس اصلی را وارد کنید (مثال: آقای ... / خانم ...):"
                        remember_last_question(context, label)
                        context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard())
                        return
        
            if step == 13:
                data["بازرس اصلی"] = text
                data["step"] = 14
                label = "کد ملی بازرس اصلی را وارد کنید (اعداد فارسی):"
                remember_last_question(context, label)
                context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard())
                return
        
            if step == 14:
                if not is_persian_number(text):
                    context.bot.send_message(chat_id=chat_id, text="❗️کد ملی را فقط با اعداد فارسی وارد کنید.", reply_markup=main_keyboard())
                    return
                data["کد ملی بازرس اصلی"] = text
                data["step"] = 15
                label = "نام بازرس علی‌البدل را وارد کنید (مثال: آقای ... / خانم ...):"
                remember_last_question(context, label)
                context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard())
                return
        
            if step == 15:
                data["بازرس علی البدل"] = text
                data["step"] = 16
                label = "کد ملی بازرس علی‌البدل را وارد کنید (اعداد فارسی):"
                remember_last_question(context, label)
                context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard())
                return
        
            if step == 16:
                if not is_persian_number(text):
                    context.bot.send_message(chat_id=chat_id, text="❗️کد ملی را فقط با اعداد فارسی وارد کنید.", reply_markup=main_keyboard())
                    return
                data["کد ملی بازرس علی البدل"] = text
                data["step"] = 17
                remember_last_question(context, "روزنامهٔ کثیرالانتشار را انتخاب کنید:")
                send_newspaper_menu(chat_id, context, "روزنامهٔ کثیرالانتشار را انتخاب کنید:")
                return

        
            if step == 17:
                data["روزنامه کثیرالانتشار"] = text
                data["step"] = 18
                label = "نام وکیل (سهامدار یا وکیل رسمی شرکت) را وارد کنید (مثال: آقای ... / خانم ...):"
                remember_last_question(context, label)
                context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard())
                return
        
            if step == 18:
                data["وکیل"] = text
                data["step"] = 19
                label = "تعداد سهامداران حاضر را وارد کنید (عدد فارسی):"
                remember_last_question(context, label)
                context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard())
                return
        
            # دریافت تعداد سهامداران → حلقه نام/تعداد
            if step == 19:
                if not text.isdigit() and not is_persian_number(text):
                    context.bot.send_message(chat_id=chat_id, text="❗️عدد وارد کنید.", reply_markup=main_keyboard())
                    return
                count = int(fa_to_en_number(text))
                if count < 1:
                    context.bot.send_message(chat_id=chat_id, text="❗️حداقل یک سهامدار لازم است.", reply_markup=main_keyboard())
                    return
                data["تعداد سهامداران"] = count
                data["سهامدار_index"] = 1
                data["step"] = 20
                fa1 = "1".translate(str.maketrans("0123456789", "۰۱۲۳۴۵۶۷۸۹"))
                label = f"نام سهامدار شماره {fa1} را وارد کنید (مثال: آقای ... / خانم ...):"
                remember_last_question(context, label)
                context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard())
                return
        
            if step == 20:
                i = data.get("سهامدار_index", 1)
                fa_i = str(i).translate(str.maketrans("0123456789", "۰۱۲۳۴۵۶۷۸۹"))
                prefix = f"سهامدار {i}"
            
                if f"{prefix} نام" not in data:
                    data[f"{prefix} نام"] = text
                    label = f"تعداد سهام سهامدار {fa_i} را وارد کنید (اعداد فارسی):"
                    remember_last_question(context, label)
                    context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard())
                    return
            
                elif f"{prefix} تعداد" not in data:
                    if not is_persian_number(text):
                        context.bot.send_message(chat_id=chat_id, text="❗️تعداد سهام را فقط با اعداد فارسی وارد کنید.", reply_markup=main_keyboard())
                        return
            
                    data[f"{prefix} تعداد"] = text
                    total_holders = data["تعداد سهامداران"]
            
                    if i < total_holders:
                        data["سهامدار_index"] = i + 1
                        fa_next = str(i+1).translate(str.maketrans("0123456789", "۰۱۲۳۴۵۶۷۸۹"))
                        label = f"نام سهامدار شماره {fa_next} را وارد کنید (مثال: آقای ... / خانم ...):"
                        remember_last_question(context, label)
                        context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard())
                        return
            
                    # ====== اینجا به آخرین سهامدار رسیدیم — ساخت خروجی نهایی ======
                    try:
                        total_board = int(fa_to_en_number(str(data.get("تعداد اعضای هیئت مدیره", 0))))  # ← مشکل اصلی اینجا حل شد
                        meeting_title = (_meeting_title_by_jalali_date(data.get("تاریخ", "")))
                        # اگر تابع بالا در کدت نیست، از عنوان پیش‌فرض استفاده کن:
                    except NameError:
                        meeting_title = "صورتجلسه مجمع عمومی فوق‌العاده"
                        total_board = int(fa_to_en_number(str(data.get("تعداد اعضای هیئت مدیره", 0))))
            
                    # بلوک اعضای هیئت‌مدیره
                    board_parts = []
                    for j in range(1, total_board + 1):
                        nm  = data.get(f"عضو {j} نام", "")
                        nid = data.get(f"عضو {j} کد ملی", "")
                        board_parts.append(nm if not nid else f"{nm} به شماره ملی {nid}")
                    board_block = " ".join(board_parts).strip()
            
                    # جدول سهامداران
                    holders_lines = []
                    for j in range(1, data["تعداد سهامداران"] + 1):
                        nm = data.get(f"سهامدار {j} نام", "")
                        sh = data.get(f"سهامدار {j} تعداد", "")
                        holders_lines.append(f"{j}\n\t{nm}\t{sh}\t")
                    holders_block = "\n".join(holders_lines)
            
                    # ساخت متن خروجی
                    try:
                        text_out = f"""
            {meeting_title} شرکت {data.get("نام شرکت","")} ){نوع_شرکت}(
            شماره ثبت شرکت :     {data.get("شماره ثبت","")}
            شناسه ملی :      {data.get("شناسه ملی","")}
            سرمایه ثبت شده : {data.get("سرمایه","")} ریال
            
            {meeting_title} شرکت {data.get("نام شرکت","")} ){نوع_شرکت}( ثبت شده به شماره {data.get("شماره ثبت","")} در تاریخ {data.get("تاریخ","")} ساعت {data.get("ساعت","")} با حضور کلیه سهامداران در محل قانونی شرکت تشکیل گردید.
            الف: در اجرای ماده 101 لایحه اصلاحی قانون تجارت
            ـ  {data.get("مدیر عامل","")}                                   به سمت رئیس جلسه 
            ـ  {data.get("نایب رییس","")}                                  به سمت ناظر 1 جلسه 
            ـ  {data.get("رییس","")}                                        به سمت ناظر 2 جلسه 
            ـ  {data.get("منشی","")}                                        به سمت منشی جلسه انتخاب شدند
            ب: در خصوص دستور جلسه، 1ـ انتخاب مدیران 2ـ انتخاب بازرسین 3ـ انتخاب روزنامه کثیرالانتشار
            ب ـ 1ـ اعضای هیات مدیره عبارتند از {board_block} برای مدت دو سال انتخاب و با امضاء ذیل صورتجلسه قبولی خود را اعلام می دارند. 
            ب ـ 2ـ با رعایت ماده 147 لایحه اصلاحی قانون تجارت {data.get("بازرس اصلی","")} به شماره ملی {data.get("کد ملی بازرس اصلی","")} به سمت بازرس اصلی و {data.get("بازرس علی البدل","")} به شماره ملی {data.get("کد ملی بازرس علی البدل","")} به سمت بازرس علی البدل برای مدت یک سال مالی انتخاب شدند.
            ب ـ 3ـ روزنامه کثیرالانتشار {data.get("روزنامه کثیرالانتشار","")} جهت نشر آگهی های شرکت انتخاب شد.
            ج: اینجانبان اعضاء هیات مدیره و بازرسین ضمن قبولی سمت خود اقرار می نمائیم که هیچگونه سوء پیشینه کیفری نداشته و ممنوعیت اصل 141 قانون اساسی و مواد 111 و 147 لایحه اصلاحی قانون تجارت را نداریم. 
            د: به {data.get("وکیل","")} احدی از سهامداران یا وکیل رسمی شرکت وکالت داده می شود که ضمن مراجعه به اداره ثبت شرکت ها نسبت به ثبت صورتجلسه و پرداخت حق الثبت و امضاء ذیل دفاتر ثبت اقدام نماید.
            امضاء اعضاء هیات رئیسه: 
            رئیس جلسه :  {data.get("مدیر عامل","")}                                   ناظر1 جلسه : {data.get("نایب رییس","")}                               
            
            ناظر2جلسه : {data.get("رییس","")}                                       منشی جلسه: {data.get("منشی","")}
            
            امضاء اعضای هیات مدیره:
            { "                           ".join([data.get(f"عضو {k} نام","") for k in range(1, total_board+1)]) }
            امضاء بازرسین:
            {data.get("بازرس اصلی","")}                                    {data.get("بازرس علی البدل","")}
            
            صورت سهامداران حاضر در {meeting_title} مورخه {data.get("تاریخ","")}
            {data.get("نام شرکت","")}
            ردیف\tنام و نام خانوادگی\tتعداد سهام\tامضا سهامداران
            {holders_block}
            """
                    except Exception as e:
                        context.bot.send_message(chat_id=chat_id, text=f"❗️خطا در ساخت متن: {e}", reply_markup=main_keyboard())
                        data["step"] = 20
                        return
            
                    # ارسال متن به صورت تکه‌تکه (حد ۴۰۹۶ کاراکتر تلگرام)
                    try:
                        for ofs in range(0, len(text_out), 3500):
                            context.bot.send_message(chat_id=chat_id, text=text_out[ofs:ofs+3500], reply_markup=main_keyboard())
                    except Exception as e:
                        context.bot.send_message(chat_id=chat_id, text=f"❗️خطا در ارسال متن: {e}", reply_markup=main_keyboard())
            
                    # فایل Word
                    try:
                        file_path = generate_word_file(text_out)  # فرض بر این است که قبلاً در پروژه‌ات داریش
                        with open(file_path, 'rb') as f:
                            context.bot.send_document(chat_id=chat_id, document=f, filename="صورتجلسه تمدید سمت اعضا.docx")
                        os.remove(file_path)
                    except Exception as e:
                        context.bot.send_message(chat_id=chat_id, text=f"❗️خطا در ساخت/ارسال فایل Word: {e}", reply_markup=main_keyboard())
            
                    # قفل کردن فرم
                    data["step"] = 21
                    return


    
        # -------------------------------
        # تغییر نام شرکت - مسئولیت محدود
        # -------------------------------
        if موضوع == "تغییر نام شرکت" and نوع_شرکت == "مسئولیت محدود":
            if step == 1:
                data["نام شرکت"] = text
                data["step"] = 2
                label = get_label("شماره ثبت")
                remember_last_question(context, label)
                context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard())
                return
    
            if step == 2:
                if not is_persian_number(text):
                    context.bot.send_message(chat_id=chat_id, text="❗️شماره ثبت را فقط با اعداد فارسی وارد کنید.", reply_markup=main_keyboard())
                    return
                data["شماره ثبت"] = text
                data["step"] = 3
                label = get_label("شناسه ملی")
                remember_last_question(context, label)
                context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard())
                return
    
            if step == 3:
                if not is_persian_number(text):
                    context.bot.send_message(chat_id=chat_id, text="❗️شناسه ملی را فقط با اعداد فارسی وارد کنید.", reply_markup=main_keyboard())
                    return
                data["شناسه ملی"] = text
                data["step"] = 4
                label = get_label("سرمایه")
                remember_last_question(context, label)
                context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard())
                return
    
            if step == 4:
                if not is_persian_number(text):
                    context.bot.send_message(chat_id=chat_id, text="❗️سرمایه را فقط با اعداد فارسی وارد کنید.", reply_markup=main_keyboard())
                    return
                data["سرمایه"] = text
                data["step"] = 5
                label = get_label("تاریخ")
                remember_last_question(context, label)
                context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard())
                return
    
            if step == 5:
                if 'is_valid_persian_date' in globals():
                    if not is_valid_persian_date(text):
                        context.bot.send_message(chat_id=chat_id, text="❗️فرمت تاریخ صحیح نیست. نمونه: ۱۴۰۴/۰۵/۱۵", reply_markup=main_keyboard())
                        return
                else:
                    if text.count('/') != 2:
                        context.bot.send_message(chat_id=chat_id, text="❗️فرمت تاریخ صحیح نیست.", reply_markup=main_keyboard())
                        return
                data["تاریخ"] = text
                data["step"] = 6
                label = get_label("ساعت")
                remember_last_question(context, label)
                context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard())
                return
    
            if step == 6:
                if not is_persian_number(text):
                    context.bot.send_message(chat_id=chat_id, text="❗️ساعت را فقط با اعداد فارسی وارد کنید.", reply_markup=main_keyboard())
                    return
                data["ساعت"] = text
                data["step"] = 7
                label = get_label("نام جدید شرکت")
                remember_last_question(context, label)
                context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard())
                return
    
            if step == 7:
                data["نام جدید شرکت"] = text
                data["step"] = 8
                label = get_label("تعداد شرکا")
                remember_last_question(context, label)
                context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard())
                return
    
            if step == 8:
                if not text.isdigit():
                    context.bot.send_message(chat_id=chat_id, text="❗️عدد وارد کنید.", reply_markup=main_keyboard())
                    return
                count = int(text)
                if count < 2:
                    context.bot.send_message(chat_id=chat_id, text="❗️حداقل دو شریک لازم است.", reply_markup=main_keyboard())
                    return
                data["تعداد شرکا"] = count
                data["current_partner"] = 1
                data["step"] = 9
                label = get_label("نام شریک", i=1)
                remember_last_question(context, label)
                context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard())
                return
    
            if step == 9:
                i = data["current_partner"]
                data[f"شریک {i}"] = text
                data["step"] = 10
                label = get_label("سهم الشرکه شریک", i=i)
                remember_last_question(context, label)
                context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard())
                return
    
            if step == 10:
                i = data["current_partner"]
                if not is_persian_number(text):
                    context.bot.send_message(chat_id=chat_id, text="❗️سهم‌الشرکه را فقط با اعداد فارسی وارد کنید.", reply_markup=main_keyboard())
                    return
                data[f"سهم الشرکه شریک {i}"] = text
                if i < data["تعداد شرکا"]:
                    data["current_partner"] = i + 1
                    data["step"] = 9
                    label = get_label("نام شریک", i=i+1)
                    remember_last_question(context, label)
                    context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard())
                else:
                    data["step"] = 11
                    label = get_label("وکیل")
                    remember_last_question(context, label)
                    context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard())
                return
    
            if step == 11:
                data["وکیل"] = text
                send_summary(chat_id, context)
                data["step"] = 12
                return
    
            if step >= 12:
                context.bot.send_message(chat_id=chat_id, text="✅ اطلاعات ثبت شد. برای شروع مجدد /start را ارسال کنید.", reply_markup=main_keyboard())
                return
    
        # ✅ تغییر موضوع فعالیت - مسئولیت محدود
        if موضوع == "تغییر موضوع فعالیت" and نوع_شرکت == "مسئولیت محدود":
            if step == 1:
                data["نام شرکت"] = text
                data["step"] = 2
                label = "شماره ثبت شرکت را وارد کنید:"
                remember_last_question(context, label)
                context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard())
                return
    
            if step == 2:
                if not is_persian_number(text):
                    context.bot.send_message(chat_id=chat_id, text="❗️شماره ثبت را فقط با اعداد فارسی وارد کنید.", reply_markup=main_keyboard())
                    return
                data["شماره ثبت"] = text
                data["step"] = 3
                label = "شناسه ملی شرکت را وارد کنید:"
                remember_last_question(context, label)
                context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard())
                return
    
            if step == 3:
                if not is_persian_number(text):
                    context.bot.send_message(chat_id=chat_id, text="❗️شناسه ملی را فقط با اعداد فارسی وارد کنید.", reply_markup=main_keyboard())
                    return
                data["شناسه ملی"] = text
                data["step"] = 4
                label = "سرمایه شرکت به ریال را وارد کنید (اعداد فارسی):"
                remember_last_question(context, label)
                context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard())
                return
    
            if step == 4:
                if not is_persian_number(text):
                    context.bot.send_message(chat_id=chat_id, text="❗️سرمایه را فقط با اعداد فارسی وارد کنید.", reply_markup=main_keyboard())
                    return
                data["سرمایه"] = text
                data["step"] = 5
                label = "تاریخ صورتجلسه را وارد کنید (مثلاً: ۱۴۰۴/۰۵/۱۵):"
                remember_last_question(context, label)
                context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard())
                return
    
            if step == 5:
                if text.count('/') != 2:
                    context.bot.send_message(chat_id=chat_id, text="❗️فرمت تاریخ صحیح نیست.", reply_markup=main_keyboard())
                    return
                data["تاریخ"] = text
                data["step"] = 6
                label = "ساعت جلسه را وارد کنید:"
                remember_last_question(context, label)
                context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard())
                return
    
            if step == 6:
                if not is_persian_number(text):
                    context.bot.send_message(chat_id=chat_id, text="❗️ساعت را فقط با اعداد فارسی وارد کنید.", reply_markup=main_keyboard())
                    return
                data["ساعت"] = text
                data["step"] = 7
                label = "تعداد شرکا را وارد کنید:"
                remember_last_question(context, label)
                context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard())
                return
    
            if step == 7:
                if not text.isdigit():
                    context.bot.send_message(chat_id=chat_id, text="❗️عدد وارد کنید.", reply_markup=main_keyboard())
                    return
                count = int(text)
                data["تعداد شرکا"] = count
                data["current_partner"] = 1
                data["step"] = 8
                label = "نام شریک شماره ۱ را وارد کنید:"
                remember_last_question(context, label)
                context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard())
                return
    
            if step == 8:
                i = data["current_partner"]
                data[f"شریک {i}"] = text
                data["step"] = 9
                label = f"سهم الشرکه شریک شماره {i} را وارد کنید (عدد فارسی):"
                remember_last_question(context, label)
                context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard())
                return
    
            if step == 9:
                i = data["current_partner"]
                if not is_persian_number(text):
                    context.bot.send_message(chat_id=chat_id, text="❗️سهم الشرکه را فقط با اعداد فارسی وارد کنید.", reply_markup=main_keyboard())
                    return
                data[f"سهم الشرکه شریک {i}"] = text
                if i < data["تعداد شرکا"]:
                    data["current_partner"] += 1
                    data["step"] = 8
                    label = f"نام شریک شماره {i+1} را وارد کنید:"
                    remember_last_question(context, label)
                    context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard())
                else:
                    data["step"] = 10
                    # مرحله بعدی با دکمه‌های اینلاین است؛ این را در last_question ذخیره نکن تا در بازگشت از AI مشکلی نباشد.
                    keyboard = [
                        [InlineKeyboardButton("➕ اضافه می‌گردد", callback_data='الحاق')],
                        [InlineKeyboardButton("🔄 جایگزین می‌گردد", callback_data='جایگزین')]
                    ]
                    context.bot.send_message(chat_id=chat_id, text="❓آیا موضوعات جدید به موضوع قبلی اضافه می‌شوند یا جایگزین آن؟", reply_markup=InlineKeyboardMarkup(keyboard))
                return
    
            # در CallbackHandler مربوط به این مرحله، نیازی به remember_last_question نیست (ورودی از طریق دکمه است)
            if data.get("step") == 10 and update.callback_query:
                answer = update.callback_query.data
                update.callback_query.answer()
                if answer in ["الحاق", "جایگزین"]:
                    data["نوع تغییر موضوع"] = answer
                    data["step"] = 11
                    label = "موضوع جدید فعالیت شرکت را وارد کنید:"
                    remember_last_question(context, label)
                    context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard())
                return
    
            if step == 11:
                data["موضوع جدید"] = text
                data["step"] = 12
                label = "نام وکیل (ثبت‌کننده صورتجلسه) را وارد کنید:"
                remember_last_question(context, label)
                context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard())
                return
    
            if step == 12:
                data["وکیل"] = text
                send_summary(chat_id, context)
                return
    
        # ✅ تغییر موضوع فعالیت – سهامی خاص
        if موضوع == "تغییر موضوع فعالیت" and نوع_شرکت == "سهامی خاص":
            if step == 1:
                data["نام شرکت"] = text
                data["step"] = 2
                label = "شماره ثبت شرکت را وارد کنید (اعداد فارسی):"
                remember_last_question(context, label)
                context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard())
                return
    
            if step == 2:
                if not is_persian_number(text):
                    context.bot.send_message(chat_id=chat_id, text="❗️شماره ثبت را فقط با اعداد فارسی وارد کنید.", reply_markup=main_keyboard())
                    return
                data["شماره ثبت"] = text
                data["step"] = 3
                label = "شناسه ملی شرکت را وارد کنید (اعداد فارسی):"
                remember_last_question(context, label)
                context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard())
                return
    
            if step == 3:
                if not is_persian_number(text):
                    context.bot.send_message(chat_id=chat_id, text="❗️شناسه ملی را فقط با اعداد فارسی وارد کنید.", reply_markup=main_keyboard())
                    return
                data["شناسه ملی"] = text
                data["step"] = 4
                label = "سرمایه ثبت‌شده شرکت (به ریال، اعداد فارسی):"
                remember_last_question(context, label)
                context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard())
                return
    
            if step == 4:
                if not is_persian_number(text):
                    context.bot.send_message(chat_id=chat_id, text="❗️سرمایه را فقط با اعداد فارسی وارد کنید.", reply_markup=main_keyboard())
                    return
                data["سرمایه"] = text
                data["step"] = 5
                label = "تاریخ صورتجلسه را وارد کنید (مثلاً: ۱۴۰۴/۰۵/۱۵):"
                remember_last_question(context, label)
                context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard())
                return
    
            if step == 5:
                if text.count('/') != 2:
                    context.bot.send_message(chat_id=chat_id, text="❗️فرمت تاریخ صحیح نیست.", reply_markup=main_keyboard())
                    return
                data["تاریخ"] = text
                data["step"] = 6
                label = "ساعت جلسه را وارد کنید (اعداد فارسی):"
                remember_last_question(context, label)
                context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard())
                return
    
            if step == 6:
                if not is_persian_number(text):
                    context.bot.send_message(chat_id=chat_id, text="❗️ساعت را فقط با اعداد فارسی وارد کنید.", reply_markup=main_keyboard())
                    return
                data["ساعت"] = text
                data["step"] = 7
                label = "مدیر عامل (رئیس جلسه) را وارد کنید:"
                remember_last_question(context, label)
                context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard())
                return
    
            if step == 7:
                data["مدیر عامل"] = text
                data["step"] = 8
                label = "ناظر 1 جلسه (نایب رئیس) را وارد کنید:"
                remember_last_question(context, label)
                context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard())
                return
    
            if step == 8:
                if text == data["مدیر عامل"]:
                    context.bot.send_message(chat_id=chat_id, text="❗️ناظر 1 نمی‌تواند با مدیر عامل یکی باشد. شخص دیگری را وارد کنید.", reply_markup=main_keyboard())
                    return
                data["نایب رییس"] = text
                data["step"] = 9
                label = "ناظر 2 جلسه (رییس) را وارد کنید:"
                remember_last_question(context, label)
                context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard())
                return
    
            if step == 9:
                if text == data["مدیر عامل"] or text == data["نایب رییس"]:
                    context.bot.send_message(chat_id=chat_id, text="❗️ناظر 2 نمی‌تواند با مدیر عامل یا ناظر 1 یکی باشد.", reply_markup=main_keyboard())
                    return
                data["رییس"] = text
                data["step"] = 10
                label = "منشی جلسه را وارد کنید:"
                remember_last_question(context, label)
                context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard())
                return
    
            if step == 10:
                data["منشی"] = text
                data["step"] = 11
                label = "تعداد سهامداران حاضر را وارد کنید:"
                remember_last_question(context, label)
                context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard())
                return
    
            if step == 11:
                if not text.isdigit():
                    context.bot.send_message(chat_id=chat_id, text="❗️عدد وارد کنید.", reply_markup=main_keyboard())
                    return
                count = int(text)
                if count < 1:
                    context.bot.send_message(chat_id=chat_id, text="❗️حداقل یک سهامدار باید وجود داشته باشد.", reply_markup=main_keyboard())
                    return
                data["تعداد سهامداران"] = count
                data["سهامدار_index"] = 1
                data["step"] = 12
                label = "نام سهامدار شماره ۱ را وارد کنید:"
                remember_last_question(context, label)
                context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard())
                return
    
            if step == 12:
                i = data.get("سهامدار_index", 1)
                prefix = f"سهامدار {i}"
                if f"{prefix} نام" not in data:
                    data[f"{prefix} نام"] = text
                    label = f"تعداد سهام {prefix} را وارد کنید (اعداد فارسی):"
                    remember_last_question(context, label)
                    context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard())
                    return
                elif f"{prefix} تعداد" not in data:
                    if not is_persian_number(text):
                        context.bot.send_message(chat_id=chat_id, text="❗️تعداد سهام را فقط با اعداد فارسی وارد کنید.", reply_markup=main_keyboard())
                        return
                    data[f"{prefix} تعداد"] = text
                    if i < data["تعداد سهامداران"]:
                        data["سهامدار_index"] = i + 1
                        label = f"نام سهامدار شماره {i+1} را وارد کنید:"
                        remember_last_question(context, label)
                        context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard())
                        return
                    else:
                        # پس از تکمیل سهامداران، انتخاب الحاق/جایگزین
                        keyboard = [
                            [InlineKeyboardButton("➕ اضافه می‌گردد", callback_data='الحاق')],
                            [InlineKeyboardButton("🔄 جایگزین می‌گردد", callback_data='جایگزین')]
                        ]
                        data["step"] = 13
                        context.bot.send_message(chat_id=chat_id, text="❓آیا موضوعات جدید به موضوع قبلی اضافه می‌شوند یا جایگزین آن؟",
                                                 reply_markup=InlineKeyboardMarkup(keyboard))
                        return
    
            if step == 14:
                data["موضوع جدید"] = text
                data["step"] = 15
                label = "نام وکیل (شخص ثبت‌کننده صورتجلسه) را وارد کنید:"
                remember_last_question(context, label)
                context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard())
                return
    
            if step == 15:
                data["وکیل"] = text
                send_summary(chat_id, context)
                return
    
        # -------------------------------
        # انحلال شرکت - مسئولیت محدود
        # -------------------------------
        if موضوع == "انحلال شرکت" and نوع_شرکت == "مسئولیت محدود":
            if step == 1:
                data["نام شرکت"] = text
                data["step"] = 2
                label = "شماره ثبت شرکت را وارد کنید (اعداد فارسی):"
                remember_last_question(context, label)
                context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard())
                return
    
            if step == 2:
                if not is_persian_number(text):
                    context.bot.send_message(chat_id=chat_id, text="❗️شماره ثبت را فقط با اعداد فارسی وارد کنید.", reply_markup=main_keyboard())
                    return
                data["شماره ثبت"] = text
                data["step"] = 3
                label = "شناسه ملی شرکت را وارد کنید (اعداد فارسی):"
                remember_last_question(context, label)
                context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard())
                return
    
            if step == 3:
                if not is_persian_number(text):
                    context.bot.send_message(chat_id=chat_id, text="❗️شناسه ملی را فقط با اعداد فارسی وارد کنید.", reply_markup=main_keyboard())
                    return
                data["شناسه ملی"] = text
                data["step"] = 4
                label = "سرمایه ثبت‌شده شرکت (ریال، اعداد فارسی):"
                remember_last_question(context, label)
                context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard())
                return
    
            if step == 4:
                if not is_persian_number(text):
                    context.bot.send_message(chat_id=chat_id, text="❗️سرمایه را فقط با اعداد فارسی وارد کنید.", reply_markup=main_keyboard())
                    return
                data["سرمایه"] = text
                data["step"] = 5
                label = "تاریخ صورتجلسه را وارد کنید (مثلاً: ۱۴۰۴/۰۵/۱۵):"
                remember_last_question(context, label)
                context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard())
                return
    
            if step == 5:
                if text.count('/') != 2:
                    context.bot.send_message(chat_id=chat_id, text="❗️فرمت تاریخ صحیح نیست.", reply_markup=main_keyboard())
                    return
                data["تاریخ"] = text
                data["step"] = 6
                label = "ساعت جلسه را وارد کنید (اعداد فارسی):"
                remember_last_question(context, label)
                context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard())
                return
    
            if step == 6:
                if not is_persian_number(text):
                    context.bot.send_message(chat_id=chat_id, text="❗️ساعت را فقط با اعداد فارسی وارد کنید.", reply_markup=main_keyboard())
                    return
                data["ساعت"] = text
                data["step"] = 7
                label = "تعداد شرکا را وارد کنید (عدد):"
                remember_last_question(context, label)
                context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard())
                return
    
            if step == 7:
                if not text.isdigit():
                    context.bot.send_message(chat_id=chat_id, text="❗️عدد وارد کنید.", reply_markup=main_keyboard())
                    return
                count = int(text)
                if count < 2:
                    context.bot.send_message(chat_id=chat_id, text="❗️حداقل دو شریک لازم است.", reply_markup=main_keyboard())
                    return
                data["تعداد شرکا"] = count
                data["current_partner"] = 1
                data["step"] = 8
                label = "نام شریک شماره ۱ را وارد کنید:"
                remember_last_question(context, label)
                context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard())
                return
    
            if step == 8:
                i = data["current_partner"]
                data[f"شریک {i}"] = text
                data["step"] = 9
                label = f"سهم‌الشرکه شریک شماره {i} را به ریال وارد کنید (اعداد فارسی):"
                remember_last_question(context, label)
                context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard())
                return
    
            if step == 9:
                i = data["current_partner"]
                if not is_persian_number(text):
                    context.bot.send_message(chat_id=chat_id, text="❗️سهم‌الشرکه را فقط با اعداد فارسی وارد کنید.", reply_markup=main_keyboard())
                    return
                data[f"سهم الشرکه شریک {i}"] = text
                if i < data["تعداد شرکا"]:
                    data["current_partner"] = i + 1
                    data["step"] = 8
                    label = f"نام شریک شماره {i+1} را وارد کنید:"
                    remember_last_question(context, label)
                    context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard())
                else:
                    data["step"] = 10
                    label = "علت انحلال را وارد کنید (مثلاً: مشکلات اقتصادی، توافق شرکا و ...):"
                    remember_last_question(context, label)
                    context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard())
                return
    
            if step == 10:
                data["علت انحلال"] = text
                data["step"] = 11
                label = "نام مدیر تصفیه را وارد کنید:"
                remember_last_question(context, label)
                context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard())
                return
    
            if step == 11:
                data["نام مدیر تصفیه"] = text
                data["step"] = 12
                label = "کد ملی مدیر تصفیه را وارد کنید (اعداد فارسی):"
                remember_last_question(context, label)
                context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard())
                return
    
            if step == 12:
                if not is_valid_persian_national_id(text):
                    context.bot.send_message(chat_id=chat_id, text="❗️کد ملی باید دقیقاً ۱۰ رقم فارسی باشد.", reply_markup=main_keyboard())
                    return
                data["کد ملی مدیر تصفیه"] = text
                data["step"] = 13
                label = "مدت مدیر تصفیه (سال) را وارد کنید (اعداد فارسی):"
                remember_last_question(context, label)
                context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard())
                return
    
            if step == 13:
                if not is_persian_number(text):
                    context.bot.send_message(chat_id=chat_id, text="❗️مدت را فقط با اعداد فارسی وارد کنید.", reply_markup=main_keyboard())
                    return
                data["مدت مدیر تصفیه"] = text
                data["step"] = 14
                label = "آدرس مدیر تصفیه و محل تصفیه را وارد کنید:"
                remember_last_question(context, label)
                context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard())
                return
    
            if step == 14:
                data["آدرس مدیر تصفیه"] = text
                data["step"] = 15
                label = "نام وکیل (ثبت‌کننده صورتجلسه) را وارد کنید:"
                remember_last_question(context, label)
                context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard())
                return
    
            if step == 15:
                data["وکیل"] = text
                send_summary(chat_id, context)
                data["step"] = 16
                return
    
            if step >= 16:
                context.bot.send_message(chat_id=chat_id, text="✅ اطلاعات قبلاً ثبت شده است. برای شروع مجدد /start را ارسال کنید.", reply_markup=main_keyboard())
                return
    
        # -------------------------------
        # انحلال شرکت - سهامی خاص
        # -------------------------------
        if موضوع == "انحلال شرکت" and نوع_شرکت == "سهامی خاص":
            if step == 1:
                data["نام شرکت"] = text
                data["step"] = 2
                label = "شماره ثبت شرکت را وارد کنید (اعداد فارسی):"
                remember_last_question(context, label)
                context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard())
                return
    
            if step == 2:
                if not is_persian_number(text):
                    context.bot.send_message(chat_id=chat_id, text="❗️شماره ثبت را فقط با اعداد فارسی وارد کنید.", reply_markup=main_keyboard())
                    return
                data["شماره ثبت"] = text
                data["step"] = 3
                label = "شناسه ملی شرکت را وارد کنید (اعداد فارسی):"
                remember_last_question(context, label)
                context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard())
                return
    
            if step == 3:
                if not is_persian_number(text):
                    context.bot.send_message(chat_id=chat_id, text="❗️شناسه ملی را فقط با اعداد فارسی وارد کنید.", reply_markup=main_keyboard())
                    return
                data["شناسه ملی"] = text
                data["step"] = 4
                label = "سرمایه ثبت‌شده (به ریال، اعداد فارسی):"
                remember_last_question(context, label)
                context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard())
                return
    
            if step == 4:
                if not is_persian_number(text):
                    context.bot.send_message(chat_id=chat_id, text="❗️سرمایه را فقط با اعداد فارسی وارد کنید.", reply_markup=main_keyboard())
                    return
                data["سرمایه"] = text
                data["step"] = 5
                label = "تاریخ صورتجلسه را وارد کنید (مثلاً ۱۴۰۴/۰۵/۱۵):"
                remember_last_question(context, label)
                context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard())
                return
    
            if step == 5:
                if text.count('/') != 2:
                    context.bot.send_message(chat_id=chat_id, text="❗️فرمت تاریخ صحیح نیست.", reply_markup=main_keyboard())
                    return
                data["تاریخ"] = text
                data["step"] = 6
                label = "ساعت جلسه را وارد کنید (اعداد فارسی):"
                remember_last_question(context, label)
                context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard())
                return
    
            if step == 6:
                if not is_persian_number(text):
                    context.bot.send_message(chat_id=chat_id, text="❗️ساعت را فقط با اعداد فارسی وارد کنید.", reply_markup=main_keyboard())
                    return
                data["ساعت"] = text
                data["step"] = 7
                label = "مدیر عامل (رئیس جلسه) را وارد کنید:"
                remember_last_question(context, label)
                context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard())
                return
    
            if step == 7:
                data["مدیر عامل"] = text
                data["step"] = 8
                label = "ناظر 1 جلسه (از بین هیئت مدیره) را وارد کنید:"
                remember_last_question(context, label)
                context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard())
                return
    
            if step == 8:
                if text == data["مدیر عامل"]:
                    context.bot.send_message(chat_id=chat_id, text="❗️ناظر 1 نمی‌تواند با مدیر عامل یکی باشد.", reply_markup=main_keyboard())
                    return
                data["نایب رییس"] = text
                data["step"] = 9
                label = "ناظر 2 جلسه (از بین هیئت مدیره) را وارد کنید:"
                remember_last_question(context, label)
                context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard())
                return
    
            if step == 9:
                if text == data["مدیر عامل"] or text == data["نایب رییس"]:
                    context.bot.send_message(chat_id=chat_id, text="❗️ناظر 2 نمی‌تواند با مدیر عامل یا ناظر 1 یکی باشد.", reply_markup=main_keyboard())
                    return
                data["رییس"] = text
                data["step"] = 10
                label = "منشی جلسه را وارد کنید:"
                remember_last_question(context, label)
                context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard())
                return
    
            if step == 10:
                data["منشی"] = text
                data["step"] = 11
                label = "علت انحلال را وارد کنید (مثلاً: مشکلات اقتصادی ، توافق شرکا و ...):"
                remember_last_question(context, label)
                context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard())
                return
    
            if step == 11:
                data["علت انحلال"] = text
                data["step"] = 12
                label = "نام مدیر تصفیه را وارد کنید:"
                remember_last_question(context, label)
                context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard())
                return
    
            if step == 12:
                data["نام مدیر تصفیه"] = text
                data["step"] = 13
                label = "کد ملی مدیر تصفیه را وارد کنید (اعداد فارسی):"
                remember_last_question(context, label)
                context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard())
                return
    
            if step == 13:
                if not is_valid_persian_national_id(text):
                    context.bot.send_message(chat_id=chat_id, text="❗️کد ملی باید دقیقاً ۱۰ رقم فارسی باشد.", reply_markup=main_keyboard())
                    return
                data["کد ملی مدیر تصفیه"] = text
                data["step"] = 14
                label = "مدت مدیر تصفیه (سال) را وارد کنید (اعداد فارسی):"
                remember_last_question(context, label)
                context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard())
                return
    
            if step == 14:
                if not is_persian_number(text):
                    context.bot.send_message(chat_id=chat_id, text="❗️مدت را فقط با اعداد فارسی وارد کنید.", reply_markup=main_keyboard())
                    return
                data["مدت مدیر تصفیه"] = text
                data["step"] = 15
                label = "آدرس مدیر تصفیه و محل تصفیه را وارد کنید:"
                remember_last_question(context, label)
                context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard())
                return
    
            if step == 15:
                data["آدرس مدیر تصفیه"] = text
                data["step"] = 16
                label = "تعداد سهامداران حاضر را وارد کنید (عدد):"
                remember_last_question(context, label)
                context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard())
                return
    
            if step == 16:
                if not text.isdigit():
                    context.bot.send_message(chat_id=chat_id, text="❗️عدد وارد کنید.", reply_markup=main_keyboard())
                    return
                data["تعداد سهامداران حاضر"] = int(text)
                data["سهامدار_index"] = 1
                data["step"] = 17
                label = "نام سهامدار ۱ را وارد کنید:"
                remember_last_question(context, label)
                context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard())
                return
    
           # حلقه سهامداران: نام → تعداد
            if step == 17:
                i = data["سهامدار_index"]
                if f"سهامدار {i} نام" not in data:
                    data[f"سهامدار {i} نام"] = text
                    label = f"تعداد سهام سهامدار {i} را وارد کنید (اعداد فارسی):"
                    remember_last_question(context, label)
                    context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard())
                    return
                elif f"سهامدار {i} تعداد" not in data:
                    if not is_persian_number(text):
                        context.bot.send_message(chat_id=chat_id, text="❗️تعداد سهام را فقط با اعداد فارسی وارد کنید.", reply_markup=main_keyboard())
                        return
                    data[f"سهامدار {i} تعداد"] = text
                    if i < data["تعداد سهامداران حاضر"]:
                        data["سهامدار_index"] += 1
                        label = f"نام سهامدار {i+1} را وارد کنید:"
                        remember_last_question(context, label)
                        context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard())
                        return
                    else:
                        data["step"] = 18
                        label = "نام وکیل (ثبت‌کننده صورتجلسه) را وارد کنید:"
                        remember_last_question(context, label)
                        context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard())
                        return
    
            if step == 18:
                data["وکیل"] = text
                send_summary(chat_id, context)
                data["step"] = 19
                return
    
            if step >= 19:
                context.bot.send_message(chat_id=chat_id, text="✅ اطلاعات قبلاً ثبت شده است. برای شروع مجدد /start را ارسال کنید.", reply_markup=main_keyboard())
                return
    
    
    # --- به‌روزرسانی کامل: نقل و انتقال سهم‌الشرکه - مسئولیت محدود ---
    
        # -------------------------------
        # نقل و انتقال سهم الشرکه - مسئولیت محدود
        # -------------------------------
        if موضوع == "نقل و انتقال سهام" and نوع_شرکت == "مسئولیت محدود":
            if step == 1:
                data["نام شرکت"] = text
                data["step"] = 2
                label = "شماره ثبت شرکت را وارد کنید (اعداد فارسی):"
                remember_last_question(context, label)
                context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard())
                return
    
            if step == 2:
                if not is_persian_number(text):
                    context.bot.send_message(chat_id=chat_id, text="❗️شماره ثبت را فقط با اعداد فارسی وارد کنید.", reply_markup=main_keyboard())
                    return
                data["شماره ثبت"] = text
                data["step"] = 3
                label = "شناسه ملی شرکت را وارد کنید (اعداد فارسی):"
                remember_last_question(context, label)
                context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard())
                return
    
            if step == 3:
                if not is_persian_number(text):
                    context.bot.send_message(chat_id=chat_id, text="❗️شناسه ملی را فقط با اعداد فارسی وارد کنید.", reply_markup=main_keyboard())
                    return
                data["شناسه ملی"] = text
                data["step"] = 4
                label = "سرمایه ثبت‌شده شرکت (ریال):"
                remember_last_question(context, label)
                context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard())
                return
    
            if step == 4:
                if not is_persian_number(text):
                    context.bot.send_message(chat_id=chat_id, text="❗️سرمایه را فقط با اعداد فارسی وارد کنید.", reply_markup=main_keyboard())
                    return
                data["سرمایه"] = text
                data["step"] = 5
                label = "تاریخ صورتجلسه را وارد کنید (مثلاً: ۱۴۰۴/۰۶/۰۱):"
                remember_last_question(context, label)
                context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard())
                return
    
            if step == 5:
                if not is_valid_persian_date(text):
                    context.bot.send_message(chat_id=chat_id, text="❗️فرمت تاریخ صحیح نیست. نمونه: ۱۴۰۴/۰۵/۱۵", reply_markup=main_keyboard())
                    return
                data["تاریخ"] = text
                data["step"] = 6
                label = get_label("ساعت")
                remember_last_question(context, label)
                context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard())
                return
    
            if step == 6:
                if not is_persian_number(text):
                    context.bot.send_message(chat_id=chat_id, text="❗️ساعت را فقط با اعداد فارسی وارد کنید.", reply_markup=main_keyboard())
                    return
                data["ساعت"] = text
                data["step"] = 7
                label = "تعداد شرکا را وارد کنید:"
                remember_last_question(context, label)
                context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard())
                return
    
            # شرکا
            if step == 7:
                if not text.isdigit():
                    context.bot.send_message(chat_id=chat_id, text="❗️عدد وارد کنید.", reply_markup=main_keyboard())
                    return
                count = int(text)
                if count < 2:
                    context.bot.send_message(chat_id=chat_id, text="❗️حداقل دو شریک لازم است.", reply_markup=main_keyboard())
                    return
                data["تعداد شرکا"] = count
                data["current_partner"] = 1
                data["step"] = 8
                label = get_label("نام شریک", i=1)
                remember_last_question(context, label)
                context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard())
                return
                
            if step == 8:
                i = data["current_partner"]
                data[f"شریک {i}"] = text
                data["step"] = 9
                label = f"سهم‌الشرکه شریک شماره {i} (ریال، اعداد فارسی):"
                remember_last_question(context, label)
                context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard())
                return
    
            if step == 9:
                i = data["current_partner"]
                if not is_persian_number(text):
                    context.bot.send_message(chat_id=chat_id, text="❗️سهم‌الشرکه را فقط با اعداد فارسی وارد کنید.", reply_markup=main_keyboard())
                    return
                data[f"سهم الشرکه شریک {i}"] = text
                if i < data["تعداد شرکا"]:
                    data["current_partner"] = i + 1
                    data["step"] = 8
                    label = f"نام شریک شماره {i+1} را وارد کنید:"
                    remember_last_question(context, label)
                    context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard())
                    return
                else:
                    data["step"] = 10
                    label = "تعداد فروشندگان را وارد کنید:"
                    remember_last_question(context, label)
                    context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard())
                    return
    
            # فروشندگان
            if step == 10:
                if not text.isdigit():
                    context.bot.send_message(chat_id=chat_id, text="❗️عدد وارد کنید.", reply_markup=main_keyboard())
                    return
                data["تعداد فروشندگان"] = int(text)
                data["فروشنده_index"] = 1
                data["step"] = 11
                label = "نام فروشنده شماره ۱ را وارد کنید:"
                remember_last_question(context, label)
                context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard())
                return
    
            if step == 11:
                i = data["فروشنده_index"]
                data[f"فروشنده {i} نام"] = text
                data["step"] = 12
                label = f"کد ملی فروشنده {i} را وارد کنید (اعداد فارسی):"
                remember_last_question(context, label)
                context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard())
                return
    
            if step == 12:
                i = data["فروشنده_index"]
                if not is_valid_persian_national_id(text):
                    context.bot.send_message(chat_id=chat_id, text="❗️کد ملی باید دقیقاً ۱۰ رقم فارسی باشد.", reply_markup=main_keyboard())
                    return
                data[f"فروشنده {i} کد ملی"] = text
                data["step"] = 13
                label = get_label("سهم کل فروشنده", i=i)
                remember_last_question(context, label)
                context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard())
                return
    
            if step == 13:
                i = data["فروشنده_index"]
                if not is_persian_number(text):
                    context.bot.send_message(chat_id=chat_id, text="❗️مبلغ را فقط با اعداد فارسی وارد کنید.", reply_markup=main_keyboard())
                    return
                data[f"فروشنده {i} سهم کل"] = text
                data["step"] = 14
                label = get_label("شماره سند صلح", i=i)
                remember_last_question(context, label)
                context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard())
                return
    
            if step == 14:
                i = data["فروشنده_index"]
                data[f"فروشنده {i} سند صلح"] = text
                data["step"] = 15
                label = f"تاریخ سند صلح فروشنده {i} را وارد کنید:"
                remember_last_question(context, label)
                context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard())
                return
    
            if step == 15:
                i = data["فروشنده_index"]
                if not is_valid_persian_date(text):
                    context.bot.send_message(chat_id=chat_id, text="❗️فرمت تاریخ صحیح نیست. نمونه: ۱۴۰۴/۰۵/۱۵", reply_markup=main_keyboard())
                    return
                data[f"فروشنده {i} تاریخ سند"] = text
                data["step"] = 16
                label = get_label("شماره دفترخانه", i=i)
                remember_last_question(context, label)
                context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard())
                return
    
            if step == 16:
                i = data["فروشنده_index"]
                data[f"فروشنده {i} دفترخانه"] = text
                data["step"] = 17
                label = f"تعداد خریداران فروشنده {i} را وارد کنید:"
                remember_last_question(context, label)
                context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard())
                return
    
            if step == 17:
                if not text.isdigit():
                    context.bot.send_message(chat_id=chat_id, text="❗️عدد وارد کنید.", reply_markup=main_keyboard())
                    return
                i = data["فروشنده_index"]
                data[f"تعداد خریداران {i}"] = int(text)
                data[f"خریدار_index_{i}"] = 1
                data["step"] = 18
                label = f"نام خریدار ۱ از فروشنده {i} را وارد کنید:"
                remember_last_question(context, label)
                context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard())
                return
    
            if step == 18:
                i = data["فروشنده_index"]
                k = data[f"خریدار_index_{i}"]
                data[f"خریدار {i}-{k} نام"] = text
                data["step"] = 19
                label = f"نام پدر خریدار {k} از فروشنده {i}:"
                remember_last_question(context, label)
                context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard())
                return
    
            if step == 19:
                i = data["فروشنده_index"]
                k = data[f"خریدار_index_{i}"]
                data[f"خریدار {i}-{k} پدر"] = text
                data["step"] = 20
                label = f"تاریخ تولد خریدار {k} از فروشنده {i}:"
                remember_last_question(context, label)
                context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard())
                return
    
            if step == 20:
                i = data["فروشنده_index"]
                k = data[f"خریدار_index_{i}"]
                if not is_valid_persian_date(text):
                    context.bot.send_message(chat_id=chat_id, text="❗️فرمت تاریخ صحیح نیست. نمونه: ۱۴۰۴/۰۵/۱۵", reply_markup=main_keyboard())
                    return
                data[f"خریدار {i}-{k} تولد"] = text
                data["step"] = 21
                label = get_label("کد ملی خریدار", i=i, k=k)
                remember_last_question(context, label)
                context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard())
                return
    
            if step == 21:
                i = data["فروشنده_index"]
                k = data[f"خریدار_index_{i}"]
                if not is_valid_persian_national_id(text):
                    context.bot.send_message(chat_id=chat_id, text="❗️کد ملی باید دقیقاً ۱۰ رقم فارسی باشد.", reply_markup=main_keyboard())
                    return
                data[f"خریدار {i}-{k} کد ملی"] = text
                data["step"] = 22
                label = get_label("آدرس خریدار", i=i, k=k)
                remember_last_question(context, label)
                context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard())
                return
    
            if step == 22:
                i = data["فروشنده_index"]
                k = data[f"خریدار_index_{i}"]
                data[f"خریدار {i}-{k} آدرس"] = text
                data["step"] = 23
                label = f"میزان سهم‌الشرکه منتقل‌شده به خریدار {k} از فروشنده {i} (ریال):"
                remember_last_question(context, label)
                context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard())
                return
    
            if step == 23:
                i = data["فروشنده_index"]
                k = data[f"خریدار_index_{i}"]
                data[f"خریدار {i}-{k} سهم منتقل"] = text
                if k < data[f"تعداد خریداران {i}"]:
                    data[f"خریدار_index_{i}"] = k + 1
                    data["step"] = 18
                    label = f"نام خریدار {k+1} از فروشنده {i} را وارد کنید:"
                    remember_last_question(context, label)
                    context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard())
                    return
                else:
                    if i < data["تعداد فروشندگان"]:
                        data["فروشنده_index"] = i + 1
                        data["step"] = 11
                        label = f"نام فروشنده شماره {i+1} را وارد کنید:"
                        remember_last_question(context, label)
                        context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard())
                        return
                    else:
                        data["step"] = 24
                        label = "نام وکیل (ثبت‌کننده صورتجلسه) را وارد کنید:"
                        remember_last_question(context, label)
                        context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard())
                        return
    
            if step == 24:
                data["وکیل"] = text
                send_summary(chat_id, context)
                data["step"] = 25
                return
    
        # -------------------------------
        # نقل و انتقال سهام - سهامی خاص
        # -------------------------------
        
        if موضوع == "نقل و انتقال سهام" and نوع_شرکت == "سهامی خاص":
            if step == 1:
                data["نام شرکت"] = text
                data["step"] = 2
                label = "شماره ثبت شرکت را وارد کنید:"
                remember_last_question(context, label)
                context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard())
                return
    
            if step == 2:
                if not is_persian_number(text):
                    context.bot.send_message(chat_id=chat_id, text="❗️شماره ثبت را فقط با اعداد فارسی وارد کنید.", reply_markup=main_keyboard())
                    return
                data["شماره ثبت"] = text
                data["step"] = 3
                label = "شناسه ملی شرکت را وارد کنید:"
                remember_last_question(context, label)
                context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard())
                return
    
            if step == 3:
                if not is_persian_number(text):
                    context.bot.send_message(chat_id=chat_id, text="❗️شناسه ملی را فقط با اعداد فارسی وارد کنید.", reply_markup=main_keyboard())
                    return
                data["شناسه ملی"] = text
                data["step"] = 4
                label = "سرمایه شرکت به ریال را وارد کنید (عدد فارسی):"
                remember_last_question(context, label)
                context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard())
                return
    
            if step == 4:
                if not is_persian_number(text):
                    context.bot.send_message(chat_id=chat_id, text="❗️سرمایه را فقط با اعداد فارسی وارد کنید.", reply_markup=main_keyboard())
                    return
                data["سرمایه"] = text
                data["step"] = 5
                label = "تاریخ صورتجلسه را وارد کنید (مثلاً: ۱۴۰۴/۰۵/۱۵):"
                remember_last_question(context, label)
                context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard())
                return
    
            if step == 5:
                if text.count('/') != 2:
                    context.bot.send_message(chat_id=chat_id, text="❗️فرمت تاریخ صحیح نیست.", reply_markup=main_keyboard())
                    return
                data["تاریخ"] = text
                data["step"] = 6
                label = "ساعت جلسه را وارد کنید :"
                remember_last_question(context, label)
                context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard())
                return
    
            if step == 6:
                if not is_persian_number(text):
                    context.bot.send_message(chat_id=chat_id, text="❗️ساعت را فقط با اعداد فارسی وارد کنید.", reply_markup=main_keyboard())
                    return
                saat = int(fa_to_en_number(text))
                if saat < 8 or saat > 17:
                    context.bot.send_message(chat_id=chat_id, text="❗️ساعت جلسه باید بین ۸ تا ۱۷ باشد.", reply_markup=main_keyboard())
                    return
                data["ساعت"] = text
                data["step"] = 7
                label = "مدیر عامل (رئیس جلسه) را وارد کنید:"
                remember_last_question(context, label)
                context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard())
                return
    
            if step == 7:
                data["مدیر عامل"] = text
                data["step"] = 8
                label = "ناظر اول جلسه را وارد کنید (از بین اعضای هیئت مدیره):"
                remember_last_question(context, label)
                context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard())
                return
    
            if step == 8:
                if text == data["مدیر عامل"]:
                    context.bot.send_message(chat_id=chat_id, text="❗️ناظر اول نمی‌تواند با مدیر عامل یکی باشد. لطفاً شخص دیگری را انتخاب کنید.", reply_markup=main_keyboard())
                    return
                data["نایب رییس"] = text
                data["step"] = 9
                label = "ناظر دوم جلسه را وارد کنید (از بین اعضای هیئت مدیره):"
                remember_last_question(context, label)
                context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard())
                return
    
            if step == 9:
                if text == data["مدیر عامل"] or text == data["نایب رییس"]:
                    context.bot.send_message(chat_id=chat_id, text="❗️ناظر دوم نمی‌تواند با مدیر عامل یا ناظر اول یکی باشد. لطفاً شخص دیگری را انتخاب کنید.", reply_markup=main_keyboard())
                    return
                data["رییس"] = text
                data["step"] = 10
                label = "منشی جلسه را وارد کنید:"
                remember_last_question(context, label)
                context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard())
                return
    
            if step == 10:
                data["منشی"] = text
                data["step"] = 11
                label = "تعداد فروشندگان را وارد کنید:"
                remember_last_question(context, label)
                context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard())
                return
    
            
            # شروع دریافت فروشندگان
            if step == 11:
                if not text.isdigit():
                    context.bot.send_message(chat_id=chat_id, text="❗️تعداد فروشندگان را با عدد وارد کنید.", reply_markup=main_keyboard())
                    return
                count = int(text)
                if count < 1:
                    context.bot.send_message(chat_id=chat_id, text="❗️حداقل یک فروشنده باید وجود داشته باشد.", reply_markup=main_keyboard())
                    return
                data["تعداد فروشندگان"] = count
                data["فروشنده_index"] = 1
                data["step"] = 12
                label = "نام فروشنده شماره ۱ را وارد کنید:"
                remember_last_question(context, label)
                context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard())
                return
    
            if step >= 12 and data.get("فروشنده_index", 0) <= data.get("تعداد فروشندگان", 0):
                i = data["فروشنده_index"]
                prefix = f"فروشنده {i}"
    
                if f"{prefix} نام" not in data:
                    data[f"{prefix} نام"] = text
                    label = f"کد ملی {prefix} را وارد کنید:"
                    remember_last_question(context, label)
                    context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard())
                    return
                if f"{prefix} کد ملی" not in data:
                    data[f"{prefix} کد ملی"] = text
                    label = f"تعداد سهام منتقل‌شده توسط {prefix} را وارد کنید:"
                    remember_last_question(context, label)
                    context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard())
                    return
                elif f"{prefix} تعداد" not in data:
                    data[f"{prefix} تعداد"] = text
                    label = "تعداد خریداران برای این فروشنده را وارد کنید:"
                    remember_last_question(context, label)
                    context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard())
                    data["step"] = 13
                    return
    
            # مرحله تعیین تعداد خریداران برای هر فروشنده
    
            if step == 13:
                if not text.isdigit():
                    context.bot.send_message(chat_id=chat_id, text="❗️تعداد خریداران را با عدد وارد کنید.", reply_markup=main_keyboard())
                    return
                count = int(text)
                if count < 1:
                    context.bot.send_message(chat_id=chat_id, text="❗️حداقل یک خریدار لازم است.", reply_markup=main_keyboard())
                    return
                i = data["فروشنده_index"]
                data[f"تعداد خریداران {i}"] = count
                data[f"خریدار_index_{i}"] = 1
                data["step"] = 14
                label = f"نام خریدار شماره ۱ از فروشنده {i} را وارد کنید:"
                remember_last_question(context, label)
                context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard())
                return
    
            if step == 14:
                i = data["فروشنده_index"]
                k = data[f"خریدار_index_{i}"]
            
                if f"خریدار {i}-{k} نام" not in data:
                    data[f"خریدار {i}-{k} نام"] = text
                    label = f"کد ملی خریدار {k} از فروشنده {i} را وارد کنید:"
                    remember_last_question(context, label)
                    context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard())
                    return
                elif f"خریدار {i}-{k} کد ملی" not in data:
                    data[f"خریدار {i}-{k} کد ملی"] = text
                    label = f"آدرس خریدار {k} از فروشنده {i} را وارد کنید:"
                    remember_last_question(context, label)
                    context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard())
                    return
                elif f"خریدار {i}-{k} آدرس" not in data:
                    data[f"خریدار {i}-{k} آدرس"] = text
                    total = data[f"تعداد خریداران {i}"]
                    if k < total:
                        data[f"خریدار_index_{i}"] += 1
                        label = f"نام خریدار شماره {k+1} از فروشنده {i} را وارد کنید:"
                        remember_last_question(context, label)
                        context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard())
                        return
                    else:
                        # همه خریداران ثبت شدن
                        if i < data["تعداد فروشندگان"]:
                            data["فروشنده_index"] += 1
                            data["step"] = 12  # برمی‌گردیم به مرحله نام فروشنده جدید
                            label = f"نام فروشنده شماره {i+1} را وارد کنید:"
                            remember_last_question(context, label)
                            context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard())
                        else:
                            data["step"] = 15  # مرحله بعد از خریداران (مثلاً سهامداران قبل)
                            label = "تعداد سهامداران قبل از نقل و انتقال را وارد کنید:"
                            remember_last_question(context, label)
                            context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard())
                        return
                    
                # مرحله دریافت سهامداران قبل از انتقال
            if step == 15:
                if not text.isdigit():
                    context.bot.send_message(chat_id=chat_id, text="❗️عدد وارد کنید.", reply_markup=main_keyboard())
                    return
                count = int(text)
                data["تعداد سهامداران قبل"] = count
                data["سهامدار_قبل_index"] = 1
                data["step"] = 16
                label = f"نام سهامدار قبل شماره ۱ را وارد کنید:"
                remember_last_question(context, label)
                context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard())
                return
        
            if step == 16:
                i = data["سهامدار_قبل_index"]
                prefix = f"سهامدار قبل {i}"
                if f"{prefix} نام" not in data:
                    data[f"{prefix} نام"] = text
                    label = f"تعداد سهام {prefix} را وارد کنید:"
                    remember_last_question(context, label)
                    context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard())
                    return
                elif f"{prefix} تعداد" not in data:
                    data[f"{prefix} تعداد"] = text
                    if i < data["تعداد سهامداران قبل"]:
                        data["سهامدار_قبل_index"] += 1
                        label = f"نام سهامدار قبل شماره {i+1} را وارد کنید:"
                        remember_last_question(context, label)
                        context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard())
                    else:
                        data["step"] = 17
                        label = "تعداد سهامداران بعد از نقل و انتقال را وارد کنید:"
                        remember_last_question(context, label)
                        context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard())
                    return
        
            # مرحله دریافت سهامداران بعد از انتقال
            if step == 17:
                if not text.isdigit():
                    context.bot.send_message(chat_id=chat_id, text="❗️عدد وارد کنید.", reply_markup=main_keyboard())
                    return
                count = int(text)
                data["تعداد سهامداران بعد"] = count
                data["سهامدار_بعد_index"] = 1
                data["step"] = 18
                label = f"نام سهامدار بعد شماره ۱ را وارد کنید:"
                remember_last_question(context, label)
                context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard())
                return
        
            if step == 18:
                i = data["سهامدار_بعد_index"]
                prefix = f"سهامدار بعد {i}"
                if f"{prefix} نام" not in data:
                    data[f"{prefix} نام"] = text
                    label = f"تعداد سهام {prefix} را وارد کنید:"
                    remember_last_question(context, label)
                    context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard())
                    return
                elif f"{prefix} تعداد" not in data:
                    data[f"{prefix} تعداد"] = text
                    if i < data["تعداد سهامداران بعد"]:
                        data["سهامدار_بعد_index"] += 1
                        label = f"نام سهامدار بعد شماره {i+1} را وارد کنید:"
                        remember_last_question(context, label)
                        context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard())
                    else:
                        data["step"] = 19
                        label = "نام وکیل (شخص ثبت‌کننده صورتجلسه) را وارد کنید:"
                        remember_last_question(context, label)
                        context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard())
                    return
        
            # مرحله آخر: دریافت وکیل
            if step == 19:
                data["وکیل"] = text
                send_summary(chat_id, context)  # ✅ ساخت و ارسال صورتجلسه
                data["step"] = 20
                return
        
            if step >= 20:
                context.bot.send_message(chat_id=chat_id, text="✅ اطلاعات قبلاً ثبت شده است. برای شروع مجدد /start را ارسال کنید.", reply_markup=main_keyboard())
                return
    
     
    # منطق قبلی برای سایر موارد و صورتجلسات
    
        if step == 1:
            data["نام شرکت"] = text
            data["step"] = 2
            next_field = fields[2]
            label = get_label(next_field)
            remember_last_question(context, label)
            context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard())
            return
    
        if step == 0:
            context.bot.send_message(chat_id=chat_id, text="لطفاً نوع شرکت را از گزینه‌های ارائه شده انتخاب کنید.", reply_markup=main_keyboard())
            return
    
        if 2 <= step < len(fields):
            field = fields[step]
    
            if field == "تاریخ":
                if text.count('/') != 2:
                    context.bot.send_message(chat_id=chat_id, text="❗️فرمت تاریخ صحیح نیست. لطفاً به صورت ۱۴۰۴/۰۴/۰۷ وارد کنید (با دو /).", reply_markup=main_keyboard())
                    return
    
            if field in persian_number_fields:
                if not is_persian_number(text):
                    context.bot.send_message(chat_id=chat_id, text=f"لطفاً مقدار '{field}' را فقط با اعداد فارسی وارد کنید.", reply_markup=main_keyboard())
                    return
    
            data[field] = text
            data["step"] += 1
            if data["step"] < len(fields):
                next_field = fields[data["step"]]
                label = get_label(next_field)
                remember_last_question(context, label)
                context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard())
            else:
                send_summary(chat_id, context)
            return

        context.bot.send_message(
            chat_id=chat_id,
            text="دستور نامعتبر یا مرحله ناشناخته است. برای بازگشت از دکمه «⬅️ بازگشت» استفاده کنید یا /start بزنید.",
            reply_markup=main_keyboard()
        )
        
    except Exception as e:
        import traceback
        traceback.print_exc()
        print("handle_message ERROR:", e)
        context.bot.send_message(chat_id=update.effective_chat.id, text="❌ خطای غیرمنتظره.")
        
def handle_back(update: Update, context: CallbackContext):
    chat_id = update.message.chat_id
    data = user_data.setdefault(chat_id, {"step": 0})
    step = data.get("step", 0)
    موضوع = data.get("موضوع صورتجلسه")
    نوع_شرکت = data.get("نوع شرکت")

    # اگر هنوز موضوع انتخاب نشده → منوی موضوعات را دوباره نشان بده
    if not موضوع:
        context.bot.send_message(chat_id=chat_id, text="به منوی موضوعات برگشتید.")
        # همون منوی موضوعات فعلی خودت را صدا بزن (تابعش هر چی اسم گذاشتی)
        send_topic_menu(chat_id, context)
        return

    # اگر در انتخاب «نوع شرکت» هستیم یا باید به آن برگردیم
    if step == 1:  # قبل از سؤال «نام شرکت»
        data.pop("نوع شرکت", None)
        data["step"] = 0
        context.bot.send_message(chat_id=chat_id, text="به انتخاب نوع شرکت برگشتید.")
        send_company_type_menu(chat_id, context)
        return

    # بازگشت پیشنهاد هوشمند نام شرکت
    if data.get("ai_mode") == "name_suggestion":
        step = data.get("step", 0)

        # اگر در مرحله 2 بودیم → پاسخ مرحله 2 را پاک کن و برگرد به مرحله 1
        if step >= 2:
            # پاک‌سازی پاسخ مرحله 2
            if "حوزه فعالیت" in data:
                data.pop("حوزه فعالیت", None)

            data["step"] = 1
            context.bot.send_message(
                chat_id=chat_id,
                text="🔁 دوباره: چه کلمه‌ی اصلی برای نام شرکت در نظر دارید؟\n"
                     "مثال: آتی، پارس، نیک، آراد…",
                reply_markup=back_keyboard()  # دکمه بازگشت ثابت
            )
            return

        # اگر در مرحله 1 یا نامشخص بود → پاسخ مرحله 1 را هم پاک کن و از فلو خارج شو
        if "کلمه اصلی" in data:
            data.pop("کلمه اصلی", None)

        data.pop("ai_mode", None)
        data["step"] = 0
        send_ai_services_menu(chat_id, context)
        return

    # --------------------------------------
    # بازگشت: تغییر نام شرکت - سهامی خاص
    # --------------------------------------
    if موضوع == "تغییر نام شرکت" and نوع_شرکت == "سهامی خاص":
        # 2..6: یک قدم عقب با لیست کلیدها
        if 2 <= step <= 6:
            prev_step = step - 1
            order = ["نام شرکت","شماره ثبت","شناسه ملی","سرمایه","تاریخ","ساعت"]
            key = order[prev_step - 1] if prev_step - 1 < len(order) else None
            if prev_step == 1:
                data.pop("نام شرکت", None)
                data["step"] = 1
                context.bot.send_message(chat_id=chat_id, text=get_label("نام شرکت"))
                return
            if key:
                data.pop(key, None)
                data["step"] = prev_step
                context.bot.send_message(chat_id=chat_id, text=get_label(key))
                return
    
        # 7..10: هیئت‌رئیسه
        if step == 7:
            data["step"] = 6
            context.bot.send_message(chat_id=chat_id, text=get_label("ساعت"))
            return
        if step == 8:
            data.pop("مدیر عامل", None)
            data["step"] = 7
            context.bot.send_message(chat_id=chat_id, text=get_label("مدیر عامل"))
            return
        if step == 9:
            data.pop("نایب رییس", None)
            data["step"] = 8
            context.bot.send_message(chat_id=chat_id, text=get_label("نایب رییس"))
            return
        if step == 10:
            data.pop("رییس", None)
            data["step"] = 9
            context.bot.send_message(chat_id=chat_id, text=get_label("رییس"))
            return
    
        # 11..12: نام جدید ← وکیل
        if step == 11:
            data.pop("منشی", None)
            data["step"] = 10
            context.bot.send_message(chat_id=chat_id, text=get_label("منشی"))
            return
        if step == 12:
            data.pop("نام جدید شرکت", None)
            data["step"] = 11
            context.bot.send_message(chat_id=chat_id, text=get_label("نام جدید شرکت"))
            return
    
        # 1: برگشت به انتخاب نوع شرکت (در صورت نیاز)
        if step == 1:
            data["step"] = 0
            send_company_type_menu(update, context)
            return

    # --------------------------------------
    # بازگشت: تمدید سمت اعضا - سهامی خاص
    # --------------------------------------
    if موضوع == "تمدید سمت اعضا" and نوع_شرکت == "سهامی خاص":
        # مسیر خطی 2..6
        if 2 <= step <= 6:
            prev_step = step - 1
            order = ["نام شرکت","شماره ثبت","شناسه ملی","سرمایه","تاریخ","ساعت"]
            key = order[prev_step - 1] if prev_step - 1 < len(order) else None
            if prev_step == 1:
                data.pop("نام شرکت", None)
                data["step"] = 1
                label = get_label("نام شرکت")
                remember_last_question(context, label)
                context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard())
                return
            if key:
                data.pop(key, None)
                data["step"] = prev_step
                label = get_label(key)
                remember_last_question(context, label)
                context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard())
                return
    
        # هیئت‌رئیسه 7..10
        if step == 7:
            data["step"] = 6
            label = get_label("ساعت")
            remember_last_question(context, label)
            context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard())
            return
        if step == 8:
            data.pop("مدیر عامل", None)
            data["step"] = 7
            label = "نام مدیرعامل را وارد کنید (مثال: آقای ... / خانم ...):"
            remember_last_question(context, label)
            context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard())
            return
        if step == 9:
            data.pop("نایب رییس", None)
            data["step"] = 8
            label = "نام نایب‌رییس (ناظر ۱) را وارد کنید (مثال: آقای ... / خانم ...):"
            remember_last_question(context, label)
            context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard())
            return
        if step == 10:
            data.pop("رییس", None)
            data["step"] = 9
            label = "نام رییس (ناظر ۲) را وارد کنید (مثال: آقای ... / خانم ...):"
            remember_last_question(context, label)
            context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard())
            return
    
        # بازگشت قبل از حلقه هیئت‌مدیره
        if step == 11:
            data.pop("منشی", None)
            data["step"] = 10
            label = "نام منشی جلسه را وارد کنید (مثال: آقای ... / خانم ...):"
            remember_last_question(context, label)
            context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard())
            return
    
        # حلقه هیئت‌مدیره (step=12)
        if step == 12:
            i = data.get("عضو_index", 1)
            fa_i = str(i).translate(str.maketrans("0123456789", "۰۱۲۳۴۵۶۷۸۹"))
            if f"عضو {i} نام" not in data:
                if i == 1:
                    data.pop("تعداد اعضای هیئت مدیره", None)
                    data["step"] = 11
                    label = "تعداد اعضای هیئت‌مدیره را وارد کنید (اعداد فارسی):"
                    remember_last_question(context, label)
                    context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard())
                    return
                prev_i = i - 1
                fa_prev = str(prev_i).translate(str.maketrans("0123456789", "۰۱۲۳۴۵۶۷۸۹"))
                data["عضو_index"] = prev_i
                data.pop(f"عضو {prev_i} کد ملی", None)
                data["step"] = 12
                label = f"کد ملی عضو هیئت‌مدیره {fa_prev} را وارد کنید (اعداد فارسی):"
                remember_last_question(context, label)
                context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard())
                return
            if f"عضو {i} کد ملی" not in data:
                data.pop(f"عضو {i} نام", None)
                data["step"] = 12
                label = f"نام عضو هیئت‌مدیره {fa_i} را وارد کنید (مثال: آقای ... / خانم ...):"
                remember_last_question(context, label)
                context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard())
                return
    
        # 13..18 بازرسین/روزنامه/وکیل
        if step == 13:
            data.pop("بازرس اصلی", None)
            data["step"] = 12
            idx = data.get('عضو_index', 1)
            fa_idx = str(idx).translate(str.maketrans("0123456789", "۰۱۲۳۴۵۶۷۸۹"))
            label = f"نام عضو هیئت‌مدیره {fa_idx} را وارد کنید (مثال: آقای ... / خانم ...):"
            remember_last_question(context, label)
            context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard())
            return
        if step == 14:
            data.pop("کد ملی بازرس اصلی", None)
            data["step"] = 13
            label = "نام بازرس اصلی را وارد کنید (مثال: آقای ... / خانم ...):"
            remember_last_question(context, label)
            context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard())
            return
        if step == 15:
            data.pop("بازرس علی البدل", None)
            data["step"] = 14
            label = "کد ملی بازرس اصلی را وارد کنید (اعداد فارسی):"
            remember_last_question(context, label)
            context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard())
            return
        if step == 16:
            data.pop("کد ملی بازرس علی البدل", None)
            data["step"] = 15
            label = "نام بازرس علی‌البدل را وارد کنید (مثال: آقای ... / خانم ...):"
            remember_last_question(context, label)
            context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard())
            return
        if step == 17:
            data.pop("روزنامه کثیرالانتشار", None)
            data["step"] = 16
            label = "کد ملی بازرس علی‌البدل را وارد کنید (اعداد فارسی):"
            remember_last_question(context, label)
            context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard())
            return
        if step == 18:
            data.pop("وکیل", None)
            data["step"] = 17
            label = "نام روزنامه کثیرالانتشار را وارد کنید:"
            remember_last_question(context, label)
            context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard())
            return
    
        # سهامداران: 19 تعداد → 20 حلقه
        if step == 19:
            data["step"] = 18
            label = "نام وکیل (سهامدار یا وکیل رسمی شرکت) را وارد کنید (مثال: آقای ... / خانم ...):"
            remember_last_question(context, label)
            context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard())
            return
    
        # --- back از حلقه سهامداران (step == 20) ---
        if step == 20:
            i = data.get("سهامدار_index", 1)
            fa_i = str(i).translate(str.maketrans("0123456789", "۰۱۲۳۴۵۶۷۸۹"))
    
            # حالت 1: الان منتظر "نام سهامدار شماره i" هستیم
            if f"سهامدار {i} نام" not in data:
                if i == 1:
                    data.pop("تعداد سهامداران", None)
                    data["step"] = 19
                    label = "تعداد سهامداران حاضر را وارد کنید (عدد فارسی):"
                    remember_last_question(context, label)
                    context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard())
                    return
                prev_i = i - 1
                fa_prev = str(prev_i).translate(str.maketrans("0123456789", "۰۱۲۳۴۵۶۷۸۹"))
                data["سهامدار_index"] = prev_i
                data.pop(f"سهامدار {prev_i} تعداد", None)
                data["step"] = 20
                label = f"تعداد سهام سهامدار {fa_prev} را وارد کنید (اعداد فارسی):"
                remember_last_question(context, label)
                context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard())
                return
    
            # حالت 2: الان منتظر "تعداد سهام سهامدار i" هستیم
            if f"سهامدار {i} تعداد" not in data:
                data.pop(f"سهامدار {i} نام", None)
                data["step"] = 20
                label = f"نام سهامدار شماره {fa_i} را وارد کنید (مثال: آقای ... / خانم ...):"
                remember_last_question(context, label)
                context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard())
                return
    
        if step >= 21:
            maxc = data.get("تعداد سهامداران", 1)
            data["سهامدار_index"] = maxc
            data.pop(f"سهامدار {maxc} تعداد", None)
            data["step"] = 20
            fa_max = str(maxc).translate(str.maketrans("0123456789", "۰۱۲۳۴۵۶۷۸۹"))
            label = f"تعداد سهام سهامدار {fa_max} را وارد کنید (اعداد فارسی):"
            remember_last_question(context, label)
            context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard())
            return



    # --------------------------------------
    # بازگشت: تغییر موضوع فعالیت – سهامی خاص
    # مراحل: 1..10 خطی، 11 تعداد سهامداران، 12 حلقه سهامداران، 13 انتخاب الحاق/جایگزین (callback)، 14 موضوع جدید، 15 وکیل
    # --------------------------------------
    if موضوع == "تغییر موضوع فعالیت" and نوع_شرکت == "سهامی خاص":
        # بازگشت در مسیر خطی 2..10
        if 2 <= step <= 10:
            prev_step = step - 1
            linear_order = {
                1:"نام شرکت", 2:"شماره ثبت", 3:"شناسه ملی", 4:"سرمایه", 5:"تاریخ",
                6:"ساعت", 7:"مدیر عامل", 8:"نایب رییس", 9:"رییس"
            }
            key = linear_order.get(prev_step, None)
            if prev_step == 1:
                data.pop("نام شرکت", None)
                data["step"] = 1
                context.bot.send_message(chat_id=chat_id, text="نام شرکت را وارد کنید:")
                return
            if key:
                data.pop(key, None)
                data["step"] = prev_step
                context.bot.send_message(chat_id=chat_id, text=get_label(key))
                return
            # prev_step == 10 → منشی
            if prev_step == 10:
                data.pop("منشی", None)
                data["step"] = 10
                context.bot.send_message(chat_id=chat_id, text="منشی جلسه را وارد کنید:")
                return

        # 11 → بازگشت به 10 (منشی)
        if step == 11:
            data.pop("تعداد سهامداران", None)
            data["step"] = 10
            context.bot.send_message(chat_id=chat_id, text="منشی جلسه را وارد کنید:")
            return

        # 12 → داخل حلقه سهامداران
        if step == 12:
            i = data.get("سهامدار_index", 1)
            # اگر منتظر نام هستیم
            if f"سهامدار {i} نام" not in data:
                if i == 1:
                    data.pop("تعداد سهامداران", None)
                    data["step"] = 11
                    context.bot.send_message(chat_id=chat_id, text="تعداد سهامداران حاضر را وارد کنید:")
                    return
                else:
                    prev_i = i - 1
                    data["سهامدار_index"] = prev_i
                    data.pop(f"سهامدار {prev_i} تعداد", None)
                    data["step"] = 12
                    context.bot.send_message(chat_id=chat_id, text=f"تعداد سهام سهامدار {prev_i} را وارد کنید (اعداد فارسی):")
                    return
            # اگر منتظر تعداد هستیم
            if f"سهامدار {i} تعداد" not in data:
                data.pop(f"سهامدار {i} نام", None)
                data["step"] = 12
                context.bot.send_message(chat_id=chat_id, text=f"نام سهامدار {i} را وارد کنید:")
                return

        # 13 (انتخاب الحاق/جایگزین) → برگرد به آخرین «تعداد سهام» در حلقه
        if step == 13:
            i = data.get("سهامدار_index", 1)
            data.pop(f"سهامدار {i} تعداد", None)
            data["step"] = 12
            context.bot.send_message(chat_id=chat_id, text=f"تعداد سهام سهامدار {i} را وارد کنید (اعداد فارسی):")
            return

        # 14 (موضوع جدید) → برگرد به دکمه الحاق/جایگزین
        if step == 14:
            data.pop("نوع تغییر موضوع", None)
            data["step"] = 13
            keyboard = [
                [InlineKeyboardButton("➕ اضافه می‌گردد", callback_data='الحاق')],
                [InlineKeyboardButton("🔄 جایگزین می‌گردد", callback_data='جایگزین')]
            ]
            context.bot.send_message(chat_id=chat_id, text="❓آیا موضوعات جدید به موضوع قبلی اضافه می‌شوند یا جایگزین آن؟",
                                     reply_markup=InlineKeyboardMarkup(keyboard))
            return

        # 15 (وکیل) → برگرد به موضوع جدید
        if step == 15:
            data.pop("موضوع جدید", None)
            data["step"] = 14
            context.bot.send_message(chat_id=chat_id, text="موضوع جدید فعالیت شرکت را وارد کنید:")
            return

    # --------------------------------------
    # بازگشت: تغییر نام شرکت - مسئولیت محدود
    # --------------------------------------
    if موضوع == "تغییر نام شرکت" and نوع_شرکت == "مسئولیت محدود":
        # 2..6: یک قدم عقب
        if 2 <= step <= 6:
            prev_step = step - 1
            order = ["نام شرکت","شماره ثبت","شناسه ملی","سرمایه","تاریخ","ساعت"]
            key = order[prev_step - 1] if prev_step - 1 < len(order) else None
            if prev_step == 1:
                data.pop("نام شرکت", None)
                data["step"] = 1
                context.bot.send_message(chat_id=chat_id, text=get_label("نام شرکت"))
                return
            if key:
                data.pop(key, None)
                data["step"] = prev_step
                context.bot.send_message(chat_id=chat_id, text=get_label(key))
                return
    
        # 7 ← برگشت به 6 (ساعت)
        if step == 7:
            data.pop("نام جدید شرکت", None)
            data["step"] = 6
            context.bot.send_message(chat_id=chat_id, text=get_label("ساعت"))
            return
    
        # 8 ← برگشت به 7 (نام جدید شرکت)
        if step == 8:
            data.pop("تعداد شرکا", None)
            data["step"] = 7
            context.bot.send_message(chat_id=chat_id, text=get_label("نام جدید شرکت"))
            return
    
        # حلقه شرکا (9 و 10)
        if step == 9:
            i = data.get("current_partner", 1)
            if i == 1:
                data.pop("تعداد شرکا", None)
                data["step"] = 8
                context.bot.send_message(chat_id=chat_id, text=get_label("تعداد شرکا"))
                return
            prev_i = i - 1
            data["current_partner"] = prev_i
            data.pop(f"سهم الشرکه شریک {prev_i}", None)
            data["step"] = 10
            context.bot.send_message(chat_id=chat_id, text=get_label("سهم الشرکه شریک", i=prev_i))
            return
    
        if step == 10:
            i = data.get("current_partner", 1)
            data.pop(f"شریک {i}", None)
            data["step"] = 9
            context.bot.send_message(chat_id=chat_id, text=get_label("نام شریک", i=i))
            return
    
        # 11 ← برگشت به «سهم‌الشرکه شریک آخر»
        if step == 11:
            last = data.get("تعداد شرکا", 1)
            data["current_partner"] = last
            data.pop(f"سهم الشرکه شریک {last}", None)
            data["step"] = 10
            context.bot.send_message(chat_id=chat_id, text=get_label("سهم الشرکه شریک", i=last))
            return
    
        # 1 ← بازگشت به انتخاب نوع شرکت (در صورت نیاز)
        if step == 1:
            data["step"] = 0
            send_company_type_menu(update, context)
            return

    # -------------------------------
    # تغییر آدرس - مسئولیت محدود
    # steps: 1=نام شرکت، 2..9 فیلدهای common، 10=تعداد شرکا، >10 حلقه شرکا (نام/سهم)
    # -------------------------------
    if موضوع == "تغییر آدرس" and نوع_شرکت == "مسئولیت محدود":
        common_fields = ["نام شرکت","شماره ثبت","شناسه ملی","سرمایه","تاریخ","ساعت","آدرس جدید","کد پستی","وکیل"]

        # برگشت داخل بخش فیلدهای common (2..10)
        if 2 <= step <= 10:
            prev_step = step - 1
            if prev_step == 1:
                data.pop("نام شرکت", None)
                data["step"] = 1
                context.bot.send_message(chat_id=chat_id, text="نام شرکت را وارد کنید:")
                return
            key = common_fields[prev_step - 1]
            data.pop(key, None)
            data["step"] = prev_step
            context.bot.send_message(chat_id=chat_id, text=get_label(key))
            return

        # حلقه شرکا: >10
        if step > 10:
            i = data.get("current_partner", 1)
            count = data.get("تعداد شرکا", 0)

            # اگر منتظر نام شریک i هستیم (پس هنوز کلید سهم‌الشرکه‌اش ثبت نشده)
            if f"شریک {i}" not in data:
                if i == 1:
                    # برگرد به «تعداد شرکا»
                    data.pop("تعداد شرکا", None)
                    data["step"] = 10
                    context.bot.send_message(chat_id=chat_id, text="تعداد شرکا را وارد کنید (بین ۲ تا ۷):")
                    return
                else:
                    # برگرد به «سهم‌الشرکه شریک قبلی»
                    prev_i = i - 1
                    data["current_partner"] = prev_i
                    data.pop(f"سهم الشرکه شریک {prev_i}", None)
                    data["step"] = 10 + prev_i  # همچنان در فاز >10
                    context.bot.send_message(chat_id=chat_id, text=f"میزان سهم الشرکه شریک شماره {prev_i} را به ریال وارد کنید (عدد فارسی):")
                    return

            # اگر منتظر سهم‌الشرکه شریک i هستیم
            if f"سهم الشرکه شریک {i}" not in data:
                data.pop(f"شریک {i}", None)
                data["step"] = 10 + i
                context.bot.send_message(chat_id=chat_id, text=f"نام شریک شماره {i} را وارد کنید:")
                return

            # اگر بعد از اتمام کار هستیم
            context.bot.send_message(chat_id=chat_id, text="برای شروع مجدد /start را ارسال کنید.")
            return

    # --------------------------------------
    # تغییر موضوع فعالیت - مسئولیت محدود
    # steps: 1..7 خطی تا «تعداد شرکا»، 8=نام شریک i، 9=سهم‌الشرکه شریک i،
    # 10=انتخاب الحاق/جایگزین (callback)، 11=موضوع جدید، 12=وکیل
    # --------------------------------------
    if موضوع == "تغییر موضوع فعالیت" and نوع_شرکت == "مسئولیت محدود":
        if 2 <= step <= 7:  # فیلدهای خطی تا قبل از ورود شرکا
            prev_step = step - 1
            order = ["نام شرکت","شماره ثبت","شناسه ملی","سرمایه","تاریخ","ساعت"]
            key = order[prev_step - 1] if prev_step - 1 < len(order) else None
            if prev_step == 1:
                data.pop("نام شرکت", None)
                data["step"] = 1
                context.bot.send_message(chat_id=chat_id, text="نام شرکت را وارد کنید:")
                return
            if key:
                data.pop(key, None)
                data["step"] = prev_step
                context.bot.send_message(chat_id=chat_id, text=get_label(key))
                return

        # 8/9: حلقه شرکا
        if step in (8, 9):
            i = data.get("current_partner", 1)
            if step == 8:
                # منتظر «نام شریک i»
                if i == 1:
                    data.pop("تعداد شرکا", None)
                    data["step"] = 7
                    context.bot.send_message(chat_id=chat_id, text="تعداد شرکا را وارد کنید:")
                    return
                else:
                    # برگرد به «سهم‌الشرکه شریک قبلی»
                    prev_i = i - 1
                    data["current_partner"] = prev_i
                    data.pop(f"سهم الشرکه شریک {prev_i}", None)
                    data["step"] = 9
                    context.bot.send_message(chat_id=chat_id, text=f"سهم الشرکه شریک شماره {prev_i} را وارد کنید (عدد فارسی):")
                    return
            else:  # step == 9 → منتظر «سهم‌الشرکه شریک i»
                data.pop(f"شریک {i}", None)
                data["step"] = 8
                context.bot.send_message(chat_id=chat_id, text=f"نام شریک شماره {i} را وارد کنید:")
                return

        # 10: دکمه الحاق/جایگزین
        if step == 10:
            i = data.get("تعداد شرکا", 1)
            data["current_partner"] = i
            data.pop(f"سهم الشرکه شریک {i}", None)
            data["step"] = 9
            context.bot.send_message(chat_id=chat_id, text=f"سهم الشرکه شریک شماره {i} را وارد کنید (عدد فارسی):")
            return

        # 11: موضوع جدید
        if step == 11:
            data.pop("نوع تغییر موضوع", None)
            data["step"] = 10
            # دوباره همان دکمه‌های الحاق/جایگزین را بفرست
            keyboard = [
                [InlineKeyboardButton("➕ اضافه می‌گردد", callback_data='الحاق')],
                [InlineKeyboardButton("🔄 جایگزین می‌گردد", callback_data='جایگزین')]
            ]
            context.bot.send_message(chat_id=chat_id, text="❓آیا موضوعات جدید به موضوع قبلی اضافه می‌شوند یا جایگزین آن؟",
                                     reply_markup=InlineKeyboardMarkup(keyboard))
            return

        # 12: وکیل
        if step == 12:
            data.pop("موضوع جدید", None)
            data["step"] = 11
            context.bot.send_message(chat_id=chat_id, text="موضوع جدید فعالیت شرکت را وارد کنید:")
            return

    # --------------------------------------
    # نقل و انتقال سهام - سهامی خاص
    # steps: 1..11 خطی
    # 12: فروشنده i (نام/کدملی/تعداد)
    # 13: تعداد خریداران برای فروشنده i
    # 14: خریدار k از فروشنده i (نام/کدملی/آدرس)
    # 15: تعداد سهامداران قبل
    # 16: حلقه سهامداران قبل (نام/تعداد)
    # 17: تعداد سهامداران بعد
    # 18: حلقه سهامداران بعد (نام/تعداد)
    # 19: وکیل
    # --------------------------------------
    if موضوع == "نقل و انتقال سهام" and نوع_شرکت == "سهامی خاص":
        linear_map = {
            1: "نام شرکت", 2: "شماره ثبت", 3: "شناسه ملی", 4: "سرمایه",
            5: "تاریخ", 6: "ساعت", 7: "مدیر عامل", 8: "نایب رییس",
            9: "رییس", 10: "منشی", 11: "تعداد فروشندگان"
        }
    
        # برگشت در مسیر خطی 2..11
        if 2 <= step <= 11:
            prev_step = step - 1
            key = linear_map.get(prev_step)
            if key:
                data.pop(key, None)
                data["step"] = prev_step
                context.bot.send_message(chat_id=chat_id, text=get_label(key))
                return
    
        # 12: فروشنده i
        if step == 12:
            i = data.get("فروشنده_index", 1)
            prefix = f"فروشنده {i}"
    
            # اگر منتظر "نام فروشنده i" هستیم
            if f"{prefix} نام" not in data:
                if i == 1:
                    data.pop("تعداد فروشندگان", None)
                    data["step"] = 11
                    context.bot.send_message(chat_id=chat_id, text="تعداد فروشندگان را وارد کنید:")
                    return
                # برگرد به "آدرس آخرین خریدارِ فروشنده قبلی"
                prev_i = i - 1
                total_k = data.get(f"تعداد خریداران {prev_i}", 1)
                data["فروشنده_index"] = prev_i
                data[f"خریدار_index_{prev_i}"] = total_k
                data.pop(f"خریدار {prev_i}-{total_k} آدرس", None)
                data["step"] = 14
                context.bot.send_message(chat_id=chat_id, text=f"آدرس خریدار {total_k} از فروشنده {prev_i} را وارد کنید:")
                return
    
            # اگر منتظر "کدملی فروشنده i" هستیم
            if f"{prefix} کد ملی" not in data:
                data.pop(f"{prefix} نام", None)
                data["step"] = 12
                context.bot.send_message(chat_id=chat_id, text=f"نام فروشنده شماره {i} را وارد کنید:")
                return
    
            # اگر منتظر "تعداد سهام منتقل‌شده فروشنده i" هستیم
            if f"{prefix} تعداد" not in data:
                data.pop(f"{prefix} کد ملی", None)
                data["step"] = 12
                context.bot.send_message(chat_id=chat_id, text=f"کد ملی فروشنده شماره {i} را وارد کنید:")
                return
    
        # 13: تعداد خریداران برای فروشنده i
        if step == 13:
            i = data.get("فروشنده_index", 1)
            data.pop(f"فروشنده {i} تعداد", None)
            data["step"] = 12
            context.bot.send_message(chat_id=chat_id, text=f"تعداد سهام منتقل‌شده توسط فروشنده {i} را وارد کنید:")
            return
    
        # 14: خریدار k از فروشنده i
        if step == 14:
            i = data.get("فروشنده_index", 1)
            k = data.get(f"خریدار_index_{i}", 1)
    
            if f"خریدار {i}-{k} نام" not in data:
                data.pop(f"تعداد خریداران {i}", None)
                data["step"] = 13
                context.bot.send_message(chat_id=chat_id, text=f"تعداد خریداران برای فروشنده {i} را وارد کنید:")
                return
    
            if f"خریدار {i}-{k} کد ملی" not in data:
                data.pop(f"خریدار {i}-{k} نام", None)
                data["step"] = 14
                context.bot.send_message(chat_id=chat_id, text=f"نام خریدار شماره {k} از فروشنده {i} را وارد کنید:")
                return
    
            if f"خریدار {i}-{k} آدرس" not in data:
                data.pop(f"خریدار {i}-{k} کد ملی", None)
                data["step"] = 14
                context.bot.send_message(chat_id=chat_id, text=f"کد ملی خریدار {k} از فروشنده {i} را وارد کنید:")
                return
    
        # 15: تعداد سهامداران قبل
        if step == 15:
            i = data.get("فروشنده_index", 1)
            total_k = data.get(f"تعداد خریداران {i}", None)
            if total_k:
                data[f"خریدار_index_{i}"] = total_k
                data.pop(f"خریدار {i}-{total_k} آدرس", None)
                data["step"] = 14
                context.bot.send_message(chat_id=chat_id, text=f"آدرس خریدار {total_k} از فروشنده {i} را وارد کنید:")
                return
            data["step"] = 13
            context.bot.send_message(chat_id=chat_id, text=f"تعداد خریداران برای فروشنده {i} را وارد کنید:")
            return
    
        # 16: حلقه سهامداران قبل (نام/تعداد)
        if step == 16:
            i = data.get("سهامدار_قبل_index", 1)
            prefix = f"سهامدار قبل {i}"
    
            # اگر منتظر نام هستیم
            if f"{prefix} نام" not in data:
                if i == 1:
                    data.pop("تعداد سهامداران قبل", None)
                    data["step"] = 15
                    context.bot.send_message(chat_id=chat_id, text="تعداد سهامداران قبل از نقل و انتقال را وارد کنید:")
                    return
                prev_i = i - 1
                data["سهامدار_قبل_index"] = prev_i
                data.pop(f"سهامدار قبل {prev_i} تعداد", None)
                data["step"] = 16
                context.bot.send_message(chat_id=chat_id, text=f"تعداد سهام سهامدار قبل شماره {prev_i} را وارد کنید:")
                return
    
            # اگر منتظر تعداد هستیم
            if f"{prefix} تعداد" not in data:
                data.pop(f"{prefix} نام", None)
                data["step"] = 16
                context.bot.send_message(chat_id=chat_id, text=f"نام سهامدار قبل شماره {i} را وارد کنید:")
                return
    
            # حالت حفاظتی: هر دو مقدار پر است ولی کاربر «بازگشت» زده
            if i > 1:
                prev_i = i - 1
                data["سهامدار_قبل_index"] = prev_i
                data.pop(f"سهامدار قبل {prev_i} تعداد", None)
                data["step"] = 16
                context.bot.send_message(chat_id=chat_id, text=f"تعداد سهام سهامدار قبل شماره {prev_i} را وارد کنید:")
                return
            else:
                data.pop("سهامدار قبل 1 نام", None)
                data.pop("سهامدار قبل 1 تعداد", None)
                data["step"] = 16
                context.bot.send_message(chat_id=chat_id, text="نام سهامدار قبل شماره ۱ را وارد کنید:")
                return
    
        # 17: تعداد سهامداران بعد
        # 17: تعداد سهامداران بعد  ← با Back باید به "تعداد" آخرین سهامدارِ قبل برگردد
        if step == 17:
            maxc = data.get("تعداد سهامداران قبل", 1)
            i = data.get("سهامدار_قبل_index", maxc)
            # اگر به هر دلیلی index از max جلوتر است، روی آخرین نفر قفل کن
            if i > maxc:
                i = maxc
                data["سهامدار_قبل_index"] = i
        
            # فقط یک قدم به عقب: "تعداد" آخرین سهامدار را پاک کن و همان را دوباره بپرس
            data.pop(f"سهامدار قبل {i} تعداد", None)
            data["step"] = 16
            context.bot.send_message(chat_id=chat_id, text=f"تعداد سهام سهامدار قبل شماره {i} را وارد کنید:")
            return
    
        # 18: حلقه سهامداران بعد (نام/تعداد)
        if step == 18:
            i = data.get("سهامدار_بعد_index", 1)
            prefix = f"سهامدار بعد {i}"
    
            # اگر منتظر نام هستیم
            if f"{prefix} نام" not in data:
                if i == 1:
                    data.pop("تعداد سهامداران بعد", None)
                    data["step"] = 17
                    context.bot.send_message(chat_id=chat_id, text="تعداد سهامداران بعد از نقل و انتقال را وارد کنید:")
                    return
                prev_i = i - 1
                data["سهامدار_بعد_index"] = prev_i
                data.pop(f"سهامدار بعد {prev_i} تعداد", None)
                data["step"] = 18
                context.bot.send_message(chat_id=chat_id, text=f"تعداد سهام سهامدار بعد شماره {prev_i} را وارد کنید:")
                return
    
            # اگر منتظر تعداد هستیم
            if f"{prefix} تعداد" not in data:
                data.pop(f"{prefix} نام", None)
                data["step"] = 18
                context.bot.send_message(chat_id=chat_id, text=f"نام سهامدار بعد شماره {i} را وارد کنید:")
                return
    
            # حالت حفاظتی
            if i > 1:
                prev_i = i - 1
                data["سهامدار_بعد_index"] = prev_i
                data.pop(f"سهامدار بعد {prev_i} تعداد", None)
                data["step"] = 18
                context.bot.send_message(chat_id=chat_id, text=f"تعداد سهام سهامدار بعد شماره {prev_i} را وارد کنید:")
                return
            else:
                data.pop("سهامدار بعد 1 نام", None)
                data.pop("سهامدار بعد 1 تعداد", None)
                data["step"] = 18
                context.bot.send_message(chat_id=chat_id, text="نام سهامدار بعد شماره ۱ را وارد کنید:")
                return
    
        # 19: وکیل
        # 19: وکیل  ← با Back باید به "تعداد" آخرین سهامدارِ بعد برگردد
        if step == 19:
            maxc = data.get("تعداد سهامداران بعد", 1)
            i = data.get("سهامدار_بعد_index", maxc)
            if i > maxc:
                i = maxc
                data["سهامدار_بعد_index"] = i
        
            data.pop(f"سهامدار بعد {i} تعداد", None)
            data["step"] = 18
            context.bot.send_message(chat_id=chat_id, text=f"تعداد سهام سهامدار بعد شماره {i} را وارد کنید:")
            return

    # --------------------------------------
    # بازگشت: انحلال شرکت - مسئولیت محدود
    # مراحل: 1..6 خطی، 7=تعداد شرکا، 8/9 حلقه شرکا، 10..15 فیلدهای پایانی
    # --------------------------------------
    if موضوع == "انحلال شرکت" and نوع_شرکت == "مسئولیت محدود":
        # خطی 2..6 → یک قدم عقب
        if 2 <= step <= 6:
            prev_step = step - 1
            order = ["نام شرکت","شماره ثبت","شناسه ملی","سرمایه","تاریخ","ساعت"]
            key = order[prev_step - 1] if prev_step - 1 < len(order) else None
            if prev_step == 1:
                data.pop("نام شرکت", None)
                data["step"] = 1
                context.bot.send_message(chat_id=chat_id, text="نام شرکت را وارد کنید:")
                return
            if key:
                data.pop(key, None)
                data["step"] = prev_step
                context.bot.send_message(chat_id=chat_id, text=get_label(key))
                return

        # 7 → برگرد به 6 (ساعت)
        if step == 7:
            data.pop("تعداد شرکا", None)
            data["step"] = 6
            context.bot.send_message(chat_id=chat_id, text=get_label("ساعت"))
            return

        # 8/9: حلقه شرکا (نام ← سهم)
        if step in (8, 9):
            i = data.get("current_partner", 1)
            if step == 8:
                # منتظر «نام شریک i»
                if i == 1:
                    data.pop("تعداد شرکا", None)
                    data["step"] = 7
                    context.bot.send_message(chat_id=chat_id, text="تعداد شرکا را وارد کنید (عدد):")
                    return
                prev_i = i - 1
                data["current_partner"] = prev_i
                data.pop(f"سهم الشرکه شریک {prev_i}", None)
                data["step"] = 9
                context.bot.send_message(chat_id=chat_id, text=f"سهم‌الشرکه شریک شماره {prev_i} را به ریال وارد کنید (اعداد فارسی):")
                return
            else:  # step == 9 → منتظر «سهم‌الشرکه»
                data.pop(f"شریک {i}", None)
                data["step"] = 8
                context.bot.send_message(chat_id=chat_id, text=f"نام شریک شماره {i} را وارد کنید:")
                return

        # 10: علت انحلال ← برگرد به سهم‌الشرکه آخرین شریک
        if step == 10:
            i = data.get("current_partner", data.get("تعداد شرکا", 1))
            if i and i >= 1 and f"سهم الشرکه شریک {i}" in data:
                data.pop(f"سهم الشرکه شریک {i}", None)
                data["step"] = 9
                context.bot.send_message(chat_id=chat_id, text=f"سهم‌الشرکه شریک شماره {i} را به ریال وارد کنید (اعداد فارسی):")
            else:
                data.pop("تعداد شرکا", None)
                data["step"] = 7
                context.bot.send_message(chat_id=chat_id, text="تعداد شرکا را وارد کنید (عدد):")
            return

        # 11..15: یک قدم به عقب در مسیر پایانی
        if step == 11:
            data.pop("علت انحلال", None)
            data["step"] = 10
            context.bot.send_message(chat_id=chat_id, text="علت انحلال را وارد کنید (مثلاً: مشکلات اقتصادی، توافق شرکا و ...):")
            return

        if step == 12:
            data.pop("نام مدیر تصفیه", None)
            data["step"] = 11
            context.bot.send_message(chat_id=chat_id, text="نام مدیر تصفیه را وارد کنید:")
            return

        if step == 13:
            data.pop("کد ملی مدیر تصفیه", None)
            data["step"] = 12
            context.bot.send_message(chat_id=chat_id, text="کد ملی مدیر تصفیه را وارد کنید (اعداد فارسی):")
            return

        if step == 14:
            data.pop("مدت مدیر تصفیه", None)
            data["step"] = 13
            context.bot.send_message(chat_id=chat_id, text="مدت مدیر تصفیه (سال) را وارد کنید (اعداد فارسی):")
            return

        if step == 15:
            data.pop("آدرس مدیر تصفیه", None)
            data["step"] = 14
            context.bot.send_message(chat_id=chat_id, text="آدرس مدیر تصفیه و محل تصفیه را وارد کنید:")
            return



    # --------------------------------------
    # بازگشت: انتخاب مدیران - سهامی خاص (بازنویسی‌شده و همسان با فلو رفت)
    # --------------------------------------
    if موضوع == "انتخاب مدیران" and نوع_شرکت == "سهامی خاص":
    
        # مراحل پایه: 2..6 (یک‌قدم عقب)
        if 2 <= step <= 6:
            prev_step = step - 1
            order = ["نام شرکت","شماره ثبت","شناسه ملی","سرمایه","تاریخ","ساعت"]
            key = order[prev_step - 1] if (prev_step - 1) < len(order) else None
    
            if prev_step == 1:
                data.pop("نام شرکت", None)
                data["step"] = 1
                label = get_label("نام شرکت")
                if 'remember_last_question' in globals(): remember_last_question(context, label)
                context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard()); return
    
            if key:
                data.pop(key, None)
                data["step"] = prev_step
                label = get_label(key)
                if 'remember_last_question' in globals(): remember_last_question(context, label)
                context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard()); return
    
        # از ۷ به ۶ (قبل از حلقهٔ اعضا)
        if step == 7:
            data.pop("ساعت", None)
            data["step"] = 6
            label = get_label("ساعت")
            if 'remember_last_question' in globals(): remember_last_question(context, label)
            context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard()); return
    
        # حلقهٔ اعضا: step=8
        if step == 8:
            i = int(data.get("board_index", 1))
            fa_i = str(i).translate(str.maketrans("0123456789","۰۱۲۳۴۵۶۷۸۹"))
            prefix = f"عضو {i}"
    
            # الف) اگر هنوز «نام عضو i» نگرفته‌ایم → روی نام i هستیم
            if f"{prefix} نام" not in data:
                # اگر i=1 → برگرد به «تعداد اعضا»
                if i == 1:
                    data.pop("تعداد اعضای هیئت مدیره", None)
                    data["step"] = 7
                    label = "تعداد اعضای هیئت‌مدیره را وارد کنید (اعداد فارسی):"
                    if 'remember_last_question' in globals(): remember_last_question(context, label)
                    context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard()); return
    
                # i>1 → برگرد یک عضو عقب و آن عضو را از نو از «نام» بپرس
                j = i - 1
                data["board_index"] = j
                # پاک‌سازی کامل عضو j (نام/کدملی/سمت/حق‌امضا/سؤال مدیرعامل)
                for k in (
                    f"عضو {j} نام",
                    f"عضو {j} کد ملی",
                    f"عضو {j} سمت",
                    f"عضو {j} سمت کد",
                    f"عضو {j} حق‌امضا",
                    f"عضو {j} مدیرعامل بیرون سهامداران؟",
                ):
                    data.pop(k, None)
    
                fa_j = str(j).translate(str.maketrans("0123456789","۰۱۲۳۴۵۶۷۸۹"))
                label = f"نام عضو هیئت‌مدیره {fa_j} را وارد کنید (مثال: آقای ... / خانم ...):"
                if 'remember_last_question' in globals(): remember_last_question(context, label)
                context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard()); return
    
            # ب) اگر «نام» داریم ولی «کد ملی عضو i» نداریم → برگرد به «نام عضو i»
            if f"{prefix} کد ملی" not in data:
                data.pop(f"{prefix} نام", None)
                label = f"نام عضو هیئت‌مدیره {fa_i} را وارد کنید (مثال: آقای ... / خانم ...):"
                if 'remember_last_question' in globals(): remember_last_question(context, label)
                context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard()); return
    
            # ج) اگر نام و کدملی هر دو ثبت شده‌اند (و منتظر دکمه‌های سمت/حق‌امضا هستیم)
            #    برگرد به «کد ملی عضو i»
            data.pop(f"{prefix} کد ملی", None)
            label = f"کد ملی عضو هیئت‌مدیره {fa_i} را وارد کنید (اعداد فارسی):"
            if 'remember_last_question' in globals(): remember_last_question(context, label)
            context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard()); return
    
        # از «وکیل» (step=9) به آخرین عضو برگرد
        if step == 9:
            data.pop("وکیل", None)
            total = 0
            try:
                total = int(fa_to_en_number(str(data.get("تعداد اعضای هیئت مدیره", 0)) or "0"))
            except Exception:
                total = 1
    
            j = total if total > 0 else 1
            data["board_index"] = j
            # پاک‌سازی کامل عضو آخر تا از «نام عضو j» شروع شود
            for k in (
                f"عضو {j} نام",
                f"عضو {j} کد ملی",
                f"عضو {j} سمت",
                f"عضو {j} سمت کد",
                f"عضو {j} حق‌امضا",
                f"عضو {j} مدیرعامل بیرون سهامداران؟",
            ):
                data.pop(k, None)
    
            fa_j = str(j).translate(str.maketrans("0123456789","۰۱۲۳۴۵۶۷۸۹"))
            data["step"] = 8
            label = f"نام عضو هیئت‌مدیره {fa_j} را وارد کنید (مثال: آقای ... / خانم ...):"
            if 'remember_last_question' in globals(): remember_last_question(context, label)
            context.bot.send_message(chat_id=chat_id, text=label, reply_markup=main_keyboard()); return



    
    # --------------------------------------
    # بازگشت: نقل و انتقال سهم‌الشرکه - مسئولیت محدود
    # مراحل:
    # 1..6 خطی پایه، 7=تعداد شرکا، 8/9 حلقه شرکا،
    # 10=تعداد فروشندگان، 11..16 خطی فروشنده،
    # 17=تعداد خریداران فروشنده i، 18..23 حلقه خریدار،
    # 24=وکیل
    # --------------------------------------
    if موضوع == "نقل و انتقال سهام" and نوع_شرکت == "مسئولیت محدود":
        # خطی پایه: 2..6 ← یک قدم عقب
        if step == 1:
            # برگشت به انتخاب نوع شرکت برای موضوع نقل و انتقال
            data["step"] = 0
            send_company_type_menu(update, context)  # همان تابعی که در پروژه‌ات داری
            return
            
        if 2 <= step <= 6:
            prev_step = step - 1
            order = ["نام شرکت","شماره ثبت","شناسه ملی","سرمایه","تاریخ","ساعت"]
            key = order[prev_step - 1] if prev_step - 1 < len(order) else None
            if prev_step == 1:
                data.pop("نام شرکت", None)
                data["step"] = 1
                context.bot.send_message(chat_id=chat_id, text="نام شرکت را وارد کنید:")
                return
            if key:
                data.pop(key, None)
                data["step"] = prev_step
                # از برچسب‌های آماده استفاده می‌کنیم اگر موجود باشد
                lbl = get_label(key) if key in order else f"{key} را وارد کنید:"
                context.bot.send_message(chat_id=chat_id, text=lbl)
                return

        # 7 ← برگشت به 6 (ساعت)
        if step == 7:
            data.pop("تعداد شرکا", None)
            data["step"] = 6
            context.bot.send_message(chat_id=chat_id, text=get_label("ساعت"))
            return

        # حلقه شرکا (8/9)
        if step in (8, 9):
            i = data.get("current_partner", 1)
            # اگر منتظر «نام شریک i» هستیم
            if step == 8:
                if i == 1:
                    data.pop("تعداد شرکا", None)
                    data["step"] = 7
                    context.bot.send_message(chat_id=chat_id, text="تعداد شرکا را وارد کنید:")
                    return
                # برگرد به «سهم‌الشرکه شریک قبلی»
                prev_i = i - 1
                data["current_partner"] = prev_i
                data.pop(f"سهم الشرکه شریک {prev_i}", None)
                data["step"] = 9
                context.bot.send_message(chat_id=chat_id, text=f"سهم‌الشرکه شریک شماره {prev_i} را به ریال وارد کنید (اعداد فارسی):")
                return
            # اگر منتظر «سهم‌الشرکه شریک i» هستیم
            if step == 9:
                data.pop(f"شریک {i}", None)
                data["step"] = 8
                context.bot.send_message(chat_id=chat_id, text=f"نام شریک شماره {i} را وارد کنید:")
                return

        # 10 ← برگرد به «سهم‌الشرکه شریک آخر»
        if step == 10:
            last = data.get("تعداد شرکا", 1)
            data["current_partner"] = last
            data.pop(f"سهم الشرکه شریک {last}", None)
            data["step"] = 9
            context.bot.send_message(chat_id=chat_id, text=f"سهم‌الشرکه شریک شماره {last} را به ریال وارد کنید (اعداد فارسی):")
            return

        # فروشنده (11..16) و تعداد خریداران (17)
        if step == 11:
            i = data.get("فروشنده_index", 1)
            if i == 1:
                data.pop("تعداد فروشندگان", None)
                data["step"] = 10
                context.bot.send_message(chat_id=chat_id, text="تعداد فروشندگان را وارد کنید:")
                return
            # برگشت به آخرین فیلد خریدارِ فروشنده قبلی (سهم منتقل)
            prev_i = i - 1
            total_k = data.get(f"تعداد خریداران {prev_i}", 1)
            data["فروشنده_index"] = prev_i
            data[f"خریدار_index_{prev_i}"] = total_k
            data.pop(f"خریدار {prev_i}-{total_k} سهم منتقل", None)
            data["step"] = 23
            context.bot.send_message(chat_id=chat_id, text=f"میزان سهم‌الشرکه منتقل‌شده به خریدار {total_k} از فروشنده {prev_i} (ریال):")
            return

        if step == 12:
            i = data.get("فروشنده_index", 1)
            data.pop(f"فروشنده {i} نام", None)
            data["step"] = 11
            context.bot.send_message(chat_id=chat_id, text=f"نام فروشنده شماره {i} را وارد کنید:")
            return

        if step == 13:
            i = data.get("فروشنده_index", 1)
            data.pop(f"فروشنده {i} کد ملی", None)
            data["step"] = 12
            context.bot.send_message(chat_id=chat_id, text=f"کد ملی فروشنده {i} را وارد کنید (اعداد فارسی):")
            return

        if step == 14:
            i = data.get("فروشنده_index", 1)
            data.pop(f"فروشنده {i} سهم کل", None)
            data["step"] = 13
            context.bot.send_message(chat_id=chat_id, text=f"کل سهم‌الشرکه فروشنده {i} (ریال):")
            return

        if step == 15:
            i = data.get("فروشنده_index", 1)
            data.pop(f"فروشنده {i} سند صلح", None)
            data["step"] = 14
            context.bot.send_message(chat_id=chat_id, text=f"شماره سند صلح فروشنده {i} را وارد کنید:")
            return

        if step == 16:
            i = data.get("فروشنده_index", 1)
            data.pop(f"فروشنده {i} تاریخ سند", None)
            data["step"] = 15
            context.bot.send_message(chat_id=chat_id, text=f"تاریخ سند صلح فروشنده {i} را وارد کنید:")
            return

        if step == 17:
            i = data.get("فروشنده_index", 1)
            data.pop(f"فروشنده {i} دفترخانه", None)
            data["step"] = 16
            context.bot.send_message(chat_id=chat_id, text=f"شماره دفترخانه فروشنده {i} را وارد کنید:")
            return

        # حلقه خریداران (18..23)
        if step == 18:
            i = data.get("فروشنده_index", 1)
            k = data.get(f"خریدار_index_{i}", 1)
            if k == 1:
                data.pop(f"تعداد خریداران {i}", None)
                data["step"] = 17
                context.bot.send_message(chat_id=chat_id, text=f"تعداد خریداران فروشنده {i} را وارد کنید:")
                return
            # برگرد به «سهم منتقلِ» خریدار قبلی
            prev_k = k - 1
            data[f"خریدار_index_{i}"] = prev_k
            data.pop(f"خریدار {i}-{prev_k} سهم منتقل", None)
            data["step"] = 23
            context.bot.send_message(chat_id=chat_id, text=f"میزان سهم‌الشرکه منتقل‌شده به خریدار {prev_k} از فروشنده {i} (ریال):")
            return

        if step == 19:
            i = data.get("فروشنده_index", 1)
            k = data.get(f"خریدار_index_{i}", 1)
            data.pop(f"خریدار {i}-{k} نام", None)
            data["step"] = 18
            context.bot.send_message(chat_id=chat_id, text=f"نام خریدار {k} از فروشنده {i} را وارد کنید:")
            return

        if step == 20:
            i = data.get("فروشنده_index", 1)
            k = data.get(f"خریدار_index_{i}", 1)
            data.pop(f"خریدار {i}-{k} پدر", None)
            data["step"] = 19
            context.bot.send_message(chat_id=chat_id, text=f"نام پدر خریدار {k} از فروشنده {i}:")
            return

        if step == 21:
            i = data.get("فروشنده_index", 1)
            k = data.get(f"خریدار_index_{i}", 1)
            data.pop(f"خریدار {i}-{k} تولد", None)
            data["step"] = 20
            context.bot.send_message(chat_id=chat_id, text=f"تاریخ تولد خریدار {k} از فروشنده {i}:")
            return

        if step == 22:
            i = data.get("فروشنده_index", 1)
            k = data.get(f"خریدار_index_{i}", 1)
            data.pop(f"خریدار {i}-{k} کد ملی", None)
            data["step"] = 21
            context.bot.send_message(chat_id=chat_id, text=f"کد ملی خریدار {k} از فروشنده {i} (اعداد فارسی):")
            return

        if step == 23:
            i = data.get("فروشنده_index", 1)
            k = data.get(f"خریدار_index_{i}", 1)
            data.pop(f"خریدار {i}-{k} آدرس", None)
            data["step"] = 22
            context.bot.send_message(chat_id=chat_id, text=f"آدرس خریدار {k} از فروشنده {i}:")
            return

        # 24 ← برگرد به «سهم منتقلِ» آخرین خریدارِ آخرین فروشنده
        if step == 24:
            i = data.get("فروشنده_index", data.get("تعداد فروشندگان", 1))
            if i > data.get("تعداد فروشندگان", 1):
                i = data.get("تعداد فروشندگان", 1)
            total_k = data.get(f"تعداد خریداران {i}", 1)
            data[f"خریدار_index_{i}"] = total_k
            data.pop(f"خریدار {i}-{total_k} سهم منتقل", None)
            data["step"] = 23
            context.bot.send_message(chat_id=chat_id, text=f"میزان سهم‌الشرکه منتقل‌شده به خریدار {total_k} از فروشنده {i} (ریال):")
            return

    # --------------------------------------
    # بازگشت: انحلال شرکت - سهامی خاص
    # --------------------------------------
    if موضوع == "انحلال شرکت" and نوع_شرکت == "سهامی خاص":
        # مراحل خطی تا قبل از حلقه سهامداران
        linear_map = {
            1: "نام شرکت", 2: "شماره ثبت", 3: "شناسه ملی", 4: "سرمایه",
            5: "تاریخ", 6: "ساعت", 7: "مدیر عامل", 8: "نایب رییس",
            9: "رییس", 10: "منشی", 11: "علت انحلال", 12: "نام مدیر تصفیه",
            13: "کد ملی مدیر تصفیه", 14: "مدت مدیر تصفیه", 15: "آدرس مدیر تصفیه",
            16: "تعداد سهامداران حاضر"
        }

        # برگشت در مسیر خطی: برگرد به سؤال قبلی و همان را بپرس
        if 2 <= step <= 16:
            prev_step = step - 1
            key = linear_map.get(prev_step)
            if key:
                data.pop(key, None)
                data["step"] = prev_step
                # اگر key در get_label نیست، متن سؤال را خودمان می‌دهیم
                label = get_label(key) if key in fields else {
                    "علت انحلال": "علت انحلال را وارد کنید (مثلاً: مشکلات اقتصادی):",
                    "نام مدیر تصفیه": "نام مدیر تصفیه را وارد کنید:",
                    "کد ملی مدیر تصفیه": "کد ملی مدیر تصفیه را وارد کنید (اعداد فارسی):",
                    "مدت مدیر تصفیه": "مدت مدیر تصفیه (سال) را وارد کنید (اعداد فارسی):",
                    "آدرس مدیر تصفیه": "آدرس مدیر تصفیه و محل تصفیه را وارد کنید:",
                    "تعداد سهامداران حاضر": "تعداد سهامداران حاضر را وارد کنید (عدد):",
                }.get(key, f"{key} را وارد کنید:")
                context.bot.send_message(chat_id=chat_id, text=label)
                return

        # حلقه سهامداران: step == 17  (نام ← تعداد)
        if step == 17:
            i = data.get("سهامدار_index", 1)
        
            # اگر هنوز نامِ سهامدار i ثبت نشده:
            if f"سهامدار {i} نام" not in data:
                if i == 1:
                    # فقط وقتی روی «نام سهامدار 1» هستیم به مرحله 16 برگرد
                    data.pop("تعداد سهامداران حاضر", None)
                    data["step"] = 16
                    context.bot.send_message(chat_id=chat_id, text=get_label("تعداد سهامداران حاضر"))
                else:
                    # برگرد به تعدادِ سهام سهامدار قبلی
                    prev_i = i - 1
                    data["سهامدار_index"] = prev_i
                    data.pop(f"سهامدار {prev_i} تعداد", None)
                    data["step"] = 17
                    context.bot.send_message(chat_id=chat_id, text=f"تعداد سهام سهامدار {prev_i} را وارد کنید (اعداد فارسی):")
                return
        
            # اگر نام ثبت شده ولی تعداد نه → برگرد به نام همان i
            if f"سهامدار {i} تعداد" not in data:
                data.pop(f"سهامدار {i} نام", None)
                data["step"] = 17
                context.bot.send_message(chat_id=chat_id, text=f"نام سهامدار {i} را وارد کنید:")
                return
        
            # هر دو مقدارِ i پر است → برو به سهامدار قبلی و تعدادش را بپرس
            if i > 1:
                data.pop(f"سهامدار {i} نام", None)
                data.pop(f"سهامدار {i} تعداد", None)
                data["سهامدار_index"] = i - 1
                data["step"] = 17
                context.bot.send_message(chat_id=chat_id, text=f"تعداد سهام سهامدار {i-1} را وارد کنید (اعداد فارسی):")
                return
            else:
                # i == 1 → برگرد ابتدای حلقه
                data.pop("سهامدار 1 نام", None)
                data.pop("سهامدار 1 تعداد", None)
                data["step"] = 17
                context.bot.send_message(chat_id=chat_id, text="نام سهامدار ۱ را وارد کنید:")
                return
        
        # وکیل: step == 18 → برگرد به آخرین سهامدار (تعداد)
        if step == 18:
            i = data.get("سهامدار_index", 1)
            data.pop("وکیل", None)
            data.pop(f"سهامدار {i} تعداد", None)  # 🔧 اضافه شد
            data["step"] = 17
            context.bot.send_message(chat_id=chat_id, text=f"تعداد سهام سهامدار {i} را وارد کنید (اعداد فارسی):")
            return

            
    # -------------------------------
    # حالت عمومی پیش‌فرض (مسیرهای ساده)
    # -------------------------------
    if step == 0:
        data.pop("موضوع صورتجلسه", None)
        data.pop("نوع شرکت", None)
        context.bot.send_message(chat_id=chat_id, text="به انتخاب موضوع برگشتید.")
        send_topic_menu(chat_id, context)
        return
    
    # فقط اگر step در محدوده‌ی فرم ساده است
    if 2 <= step < len(fields):
        prev_step = step - 1
        key = fields[prev_step]
        data.pop(key, None)
        data["step"] = prev_step
        context.bot.send_message(chat_id=chat_id, text=get_label(key))
        return
    
    # در غیر این‌صورت، هیچ برگشت عمومی نزن؛ مسیرهای تخصصی بالاتر کار را انجام داده‌اند
    context.bot.send_message(chat_id=chat_id, text="یک مرحله به عقب برگشتید.")


def button_handler(update: Update, context: CallbackContext):
    query = update.callback_query

    # ۱) رشته‌ی کال‌بک را جدا نگه دار
    payload = query.data or ""
    if payload.startswith("newspaper:"):
        return  # بگذار handle_newspaper_choice رسیدگی کند


    chat_id = query.message.chat_id
    query.answer()

    # ۲) از اینجا به بعد، 'data' دوباره همان دیکشنری وضعیت کاربر است
    data = user_data.setdefault(chat_id, {})


    
    # اگر کال‌بکِ مخصوص خروج از AI بود یا هنوز داخل AI هستیم، این هندلر کاری نکند
    if data == AI_RESUME or context.user_data.get("ai_mode"):
        return



    if "موضوع صورتجلسه" not in user_data.get(chat_id, {}):
        # اولین کلیک روی دکمه‌ی موضوع
        if query.data == "topic:extend_roles":
            # موضوع مخصوص تمدید سمت اعضا (فقط سهامی خاص)
            user_data[chat_id]["موضوع صورتجلسه"] = TOPIC_EXTEND_ROLES
            user_data[chat_id]["step"] = 0
    
            # حالت‌های این سناریو در context.user_data
            context.user_data["topic"] = TOPIC_EXTEND_ROLES
            context.user_data["company_type"] = "سهامی خاص"
    
            # پاک‌سازی وضعیت قبلی سناریو (اگر بود)
            context.user_data.pop("extend_roles", None)
            context.user_data.pop("extend_state", None)
    
            # شروع سناریو اختصاصی تمدید سمت اعضا
            start_extend_roles_flow(update, context)
            return
        else:
            # سایر موضوع‌ها طبق روال قبلی → انتخاب نوع شرکت
            user_data[chat_id]["موضوع صورتجلسه"] = query.data
            user_data[chat_id]["step"] = 0
            send_company_type_menu(chat_id, context)
            return



    if user_data[chat_id].get("step") == 0:
        user_data[chat_id]["نوع شرکت"] = query.data
        # اگر موضوع = نقل و انتقال سهام است
        if user_data[chat_id]["موضوع صورتجلسه"] == "نقل و انتقال سهام":
            if query.data == "مسئولیت محدود":
                # 👇 اول اطلاعیه ماده ۱۰۳، بعد سوال نام شرکت
                context.bot.send_message(chat_id=chat_id, text=get_label("اطلاعیه_ماده103", سند="سند صلح"))

                user_data[chat_id]["step"] = 1
                context.bot.send_message(chat_id=chat_id, text="نام شرکت را وارد کنید:")
                return
            else:
                # سهامی خاص یا سایر انواع → بدون اطلاعیه
                user_data[chat_id]["step"] = 1
                context.bot.send_message(chat_id=chat_id, text="نام شرکت را وارد کنید:")
                return

        # شروع: تغییر نام شرکت - مسئولیت محدود
        if user_data[chat_id].get("موضوع صورتجلسه") == "تغییر نام شرکت" and query.data == "مسئولیت محدود":
            user_data[chat_id]["step"] = 1
            context.bot.send_message(chat_id=chat_id, text=get_label("نام شرکت"))
            return

        # شروع: تغییر نام شرکت - سهامی خاص
        if user_data[chat_id].get("موضوع صورتجلسه") == "تغییر نام شرکت" and query.data == "سهامی خاص":
            user_data[chat_id]["step"] = 1
            context.bot.send_message(chat_id=chat_id, text=get_label("نام شرکت"))
            return
    
        # سایر موضوع‌ها
        user_data[chat_id]["step"] = 1
        context.bot.send_message(chat_id=chat_id, text="نام شرکت را وارد کنید:")
        return

    if data.get("موضوع صورتجلسه") == "تغییر موضوع فعالیت" and data.get("step") in (10, 13):
        انتخاب = query.data
        query.answer()

        if انتخاب == "الحاق":
            data["نوع تغییر موضوع"] = "الحاق"
        elif انتخاب == "جایگزین":
            data["نوع تغییر موضوع"] = "جایگزین"
        else:
            context.bot.send_message(chat_id=chat_id, text="❗️انتخاب نامعتبر بود.")
            return

        # اگر قبلاً در مسئولیت محدود بود step=10 → بعدش 11
        # اگر در سهامی خاص هستیم step=13 → بعدش 14
        if data.get("step") == 10:
            data["step"] = 11
        else:
            data["step"] = 14

        context.bot.send_message(chat_id=chat_id, text="موضوع جدید فعالیت شرکت را وارد کنید:")
        return



def render_board_election_text(d: dict) -> str:
    # لیست اعضا
    total = int(fa_to_en_number(str(d.get("تعداد اعضای هیئت مدیره", 0)) or "0"))
    lines = []
    for i in range(1, total + 1):
        nm   = d.get(f"عضو {i} نام","")
        nid  = d.get(f"عضو {i} کد ملی","")
        rol  = d.get(f"عضو {i} سمت","")          # برچسب فارسی سمت
        code = d.get(f"عضو {i} سمت کد")           # کُد سمت (ceo / chair / ...)
    
        # اگر مدیرعامل و پاسخ «بله» بوده:
        ceo_out = bool(d.get(f"عضو {i} مدیرعامل بیرون سهامداران؟"))
        suffix  = " (خارج از اعضا)" if (code == "ceo" and ceo_out) else ""
    
        if nm or nid or rol:
            lines.append(f"{nm} به شماره ملی {nid} به سمت {rol}{suffix}")


    members_block = "\n".join(lines).strip()

    # بند حق‌امضا هوشمند
    sig_clause = build_signature_clause_roles(d)
    sig_clause = f"\n{sig_clause}\n" if sig_clause else ""

    # ⚠️ پرانتزها را نرمال کردم به فرم استاندارد (… (سهامی خاص))
    text_out = f"""
صورتجلسه هیئت مدیره شرکت {d.get("نام شرکت","")} ({d.get("نوع شرکت","")})
شماره ثبت شرکت :     {d.get("شماره ثبت","")}
شناسه ملی :      {d.get("شناسه ملی","")}
سرمایه ثبت شده : {d.get("سرمایه","")} ریال

جلسه هیئت مدیره شرکت {d.get("نام شرکت","")} ({d.get("نوع شرکت","")}) ثبت شده به شماره {d.get("شماره ثبت","")} در تاریخ  {d.get("تاریخ","")} ساعت {d.get("ساعت","")} با حضور کلیه سهامداران در محل قانونی شرکت تشکیل و نسبت به تعیین سمت و تعیین دارندگان حق امضاء اتخاذ تصمیم شد. 

{members_block}

{sig_clause}
ج: اینجانبان اعضاء هیات مدیره ضمن قبولی سمت خود اقرار می نمائیم که هیچگونه سوء پیشینه کیفری نداشته و ممنوعیت اصل 141 قانون اساسی و مواد 111 و 126 لایحه اصلاحی قانون تجارت را نداریم .

هیئت مدیره به {d.get("وکیل","")} احدی از اعضاء شرکت وکالت داده می شود که ضمن مراجعه به اداره ثبت شرکتها نسبت به ثبت صورتجلسه و پرداخت حق الثبت و امضاء ذیل دفاتر ثبت اقدام نماید. 

امضاء اعضای هیات مدیره

{build_signatures_block(d)}
""".strip()
    return text_out



def send_summary(chat_id, context):
    data = user_data[chat_id]
    موضوع = data.get("موضوع صورتجلسه")
    نوع_شرکت = data.get("نوع شرکت")

        # ✅ خروجی: تغییر موضوع فعالیت – سهامی خاص
    if موضوع == "تغییر موضوع فعالیت" and نوع_شرکت == "سهامی خاص":
        # خطوط عمل بر اساس الحاق/جایگزین
        action_line = (
            "صورتجلسه مجمع عمومی فوق العاده شرکت "
            f"{data['نام شرکت']} ){نوع_شرکت} (ثبت شده به شماره {data['شماره ثبت']} در تاریخ  {data['تاریخ']} ساعت {data['ساعت']} "
            "با حضور کلیه سهامداران در محل قانونی شرکت تشکیل و نسبت به الحاق مواردی به موضوع شرکت اتخاذ تصمیم شد."
            if data.get("نوع تغییر موضوع") == "الحاق"
            else
            "صورتجلسه مجمع عمومی فوق العاده شرکت "
            f"{data['نام شرکت']} ){نوع_شرکت} (ثبت شده به شماره {data['شماره ثبت']} در تاریخ  {data['تاریخ']} ساعت {data['ساعت']} "
            "با حضور کلیه سهامداران در محل قانونی شرکت تشکیل و نسبت به تغییر موضوع شرکت اتخاذ تصمیم شد."
        )

        subject_intro = (
            "ب: مواردی به شرح ذیل به موضوع شرکت الحاق شد:"
            if data.get("نوع تغییر موضوع") == "الحاق"
            else
            "ب: موضوع شرکت به شرح ذیل تغییر یافت:"
        )

        # جدول سهامداران حاضر
        rows = ""
        for i in range(1, data.get("تعداد سهامداران", 0) + 1):
            rows += f"{i}\n\t{data.get(f'سهامدار {i} نام', '')}\t{data.get(f'سهامدار {i} تعداد', '')}\t\n"

        text = f"""صورتجلسه مجمع عمومی فوق العاده شرکت {data['نام شرکت']} ){نوع_شرکت}(
شماره ثبت شرکت :     {data['شماره ثبت']}
شناسه ملی :      {data['شناسه ملی']}
سرمایه ثبت شده : {data['سرمایه']} ریال

{action_line}
الف: در اجرای ماده 101 لایحه اصلاحی قانون تجارت: 

ـ  {data['مدیر عامل']}                                   به سمت رئیس جلسه 
ـ  {data['نایب رییس']}                                  به سمت ناظر 1 جلسه 
ـ  {data['رییس']}                                        به سمت ناظر 2 جلسه 
ـ  {data['منشی']}                                        به سمت منشی جلسه انتخاب شدند

{subject_intro}
{data['موضوع جدید']} 
و ماده مربوطه اساسنامه به شرح فوق اصلاح می گردد. 
ج: مجمع به {data['وکیل']} از سهامداران شرکت وکالت داده می شود که ضمن مراجعه به اداره ثبت شرکت ها نسبت به ثبت صورتجلسه و پرداخت حق الثبت و امضاء ذیل دفاتر ثبت اقدام نماید.

امضاء اعضاء هیات رئیسه: 
رئیس جلسه :  {data['مدیر عامل']}                                   ناظر1 جلسه : {data['نایب رییس']}                               


ناظر2جلسه : {data['رییس']}                                       منشی جلسه: {data['منشی']}





صورت سهامداران حاضر در مجمع عمومی (فوق العاده) مورخه {data['تاریخ']}
{data['نام شرکت']}
ردیف\tنام و نام خانوادگی\tتعداد سهام\tامضا سهامداران
{rows}
"""

        context.bot.send_message(chat_id=chat_id, text=text)

        # فایل Word
        file_path = generate_word_file(text)
        with open(file_path, 'rb') as f:
            context.bot.send_document(chat_id=chat_id, document=f, filename="صورتجلسه تغییر موضوع سهامی خاص.docx")
        os.remove(file_path)
        return

    # کد صورتجلسه تغییر آدرس مسئولیت محدود
    
    if موضوع == "تغییر آدرس" and نوع_شرکت == "مسئولیت محدود":
        # صورتجلسه مسئولیت محدود با لیست شرکا
        partners_lines = ""
        count = data.get("تعداد شرکا", 0)
        for i in range(1, count + 1):
            name = data.get(f"شریک {i}", "")
            share = data.get(f"سهم الشرکه شریک {i}", "")
            partners_lines += f"{name}                                              {share} ریال\n"
        text = f"""صورتجلسه مجمع عمومی فوق العاده شرکت {data['نام شرکت']} {data['نوع شرکت']}
شماره ثبت شرکت : {data['شماره ثبت']}
شناسه ملی : {data['شناسه ملی']}
سرمایه ثبت شده : {data['سرمایه']} ریال

صورتجلسه مجمع عمومی فوق العاده شرکت {data['نام شرکت']} {data['نوع شرکت']} ثبت شده به شماره {data['شماره ثبت']} در تاریخ {data['تاریخ']} ساعت {data['ساعت']} با حضور کلیه شرکا در محل قانونی شرکت تشکیل و نسبت به تغییر محل شرکت اتخاذ تصمیم شد. 

اسامی شرکا                                                     میزان سهم الشرکه
{partners_lines}
محل شرکت از آدرس قبلی به آدرس {data['آدرس جدید']} به کدپستی {data['کد پستی']} انتقال یافت.

به {data['وکیل']} احدی از شرکاء وکالت داده می شود تا ضمن مراجعه به اداره ثبت شرکتها نسبت به ثبت صورتجلسه و امضاء ذیل دفتر ثبت اقدام نماید.

امضاء شرکا : 

"""
        # فاصله بین اسامی امضاءها به سبک نمونه
        signers = ""
        for i in range(1, count + 1):
            signers += f"{data.get(f'شریک {i}', '')}     "
        text += signers
        context.bot.send_message(chat_id=chat_id, text=text)
        
        # ✅ ساخت فایل Word و ارسال
        file_path = generate_word_file(text)
        with open(file_path, 'rb') as f:
            context.bot.send_document(chat_id=chat_id, document=f, filename="صورتجلسه.docx")
    
        os.remove(file_path)  # ← حذف فایل پس از ارسال (اختیاری)
        return


    # ---------------------------
    # ۱) تمدید سمت اعضا — فقط سهامی خاص (داینامیک هیئت‌مدیره + سهامداران)
    # ---------------------------
    if موضوع == "تمدید سمت اعضا" and نوع_شرکت == "سهامی خاص":
        meeting_title = _meeting_title_by_jalali_date(data.get("تاریخ", ""))

        # بلوک هیئت‌مدیره (داینامیک)
        board_parts = []
        total_board = int(fa_to_en_number(str(data.get("تعداد اعضای هیئت مدیره", "0"))) or 0)
        for i in range(1, total_board + 1):
            nm  = data.get(f"عضو {i} نام", "")
            nid = data.get(f"عضو {i} کد ملی", "")
            board_parts.append(nm if not nid else f"{nm} به شماره ملی {nid}")
        board_block = " ".join([p for p in board_parts if p.strip()])

        # جدول سهامداران (داینامیک)
        holders_lines = []
        total_holders = int(fa_to_en_number(str(data.get("تعداد سهامداران", "0"))) or 0)
        for j in range(1, total_holders + 1):
            nm = data.get(f"سهامدار {j} نام", "")
            sh = data.get(f"سهامدار {j} تعداد", "")
            holders_lines.append(f"{j}\n\t{nm}\t{sh}\t")
        holders_block = "\n".join(holders_lines)

        # متن نهایی (طبق قالبی که خودت دادی)
        text_out = f"""
{meeting_title} شرکت {data.get("نام شرکت","")} ){نوع_شرکت}(
شماره ثبت شرکت :     {data.get("شماره ثبت","")}
شناسه ملی :      {data.get("شناسه ملی","")}
سرمایه ثبت شده : {data.get("سرمایه","")} ریال

{meeting_title} شرکت {data.get("نام شرکت","")} ){نوع_شرکت}( ثبت شده به شماره {data.get("شماره ثبت","")} در تاریخ {data.get("تاریخ","")} ساعت {data.get("ساعت","")} با حضور کلیه سهامداران در محل قانونی شرکت تشکیل گردید.
الف: در اجرای ماده 101 لایحه اصلاحی قانون تجارت
ـ  {data.get("مدیر عامل","")}                                   به سمت رئیس جلسه 
ـ  {data.get("نایب رییس","")}                                  به سمت ناظر 1 جلسه 
ـ  {data.get("رییس","")}                                        به سمت ناظر 2 جلسه 
ـ  {data.get("منشی","")}                                        به سمت منشی جلسه انتخاب شدند
ب: در خصوص دستور جلسه، 1ـ انتخاب مدیران 2ـ انتخاب بازرسین 3ـ انتخاب روزنامه کثیرالانتشار
ب ـ 1ـ اعضای هیات مدیره عبارتند از {board_block} برای مدت دو سال انتخاب و با امضاء ذیل صورتجلسه قبولی خود را اعلام می دارند. 
ب ـ 2ـ با رعایت ماده 147 لایحه اصلاحی قانون تجارت {data.get("بازرس اصلی","")} به شماره ملی {data.get("کد ملی بازرس اصلی","")} به سمت بازرس اصلی و {data.get("بازرس علی البدل","")} به شماره ملی {data.get("کد ملی بازرس علی البدل","")} به سمت بازرس علی البدل برای مدت یک سال مالی انتخاب شدند.
ب ـ 3ـ روزنامه کثیرالانتشار {data.get("روزنامه کثیرالانتشار","")} جهت نشر آگهی های شرکت انتخاب شد.
ج: اینجانبان اعضاء هیات مدیره و بازرسین ضمن قبولی سمت خود اقرار می نمائیم که هیچگونه سوء پیشینه کیفری نداشته و ممنوعیت اصل 141 قانون اساسی و مواد 111 و 147 لایحه اصلاحی قانون تجارت را نداریم. 
د: به {data.get("وکیل","")} احدی از سهامداران یا وکیل رسمی شرکت وکالت داده می شود که ضمن مراجعه به اداره ثبت شرکت ها نسبت به ثبت صورتجلسه و پرداخت حق الثبت و امضاء ذیل دفاتر ثبت اقدام نماید.
امضاء اعضاء هیات رئیسه: 
رئیس جلسه :  {data.get("مدیر عامل","")}                                   ناظر1 جلسه : {data.get("نایب رییس","")}                               


ناظر2جلسه : {data.get("رییس","")}                                       منشی جلسه: {data.get("منشی","")}

امضاء اعضای هیات مدیره:
{ "                           ".join([data.get(f"عضو {k} نام","") for k in range(1, total_board+1)]) }
امضاء بازرسین:
{data.get("بازرس اصلی","")}                                    {data.get("بازرس علی البدل","")}



صورت سهامداران حاضر در {meeting_title} مورخه {data.get("تاریخ","")}
{data.get("نام شرکت","")}
ردیف\tنام و نام خانوادگی\tتعداد سهام\tامضا سهامداران
{holders_block}
""".strip()

        # ارسال متن بلند در چند تکه (برای محدودیت تلگرام)
        for i in range(0, len(text_out), 3500):
            context.bot.send_message(chat_id=chat_id, text=text_out[i:i+3500])

        # فایل Word (با همان تابع پروژهٔ خودت)
        try:
            filepath = generate_word_file(text_out)
            with open(filepath, "rb") as f:
                context.bot.send_document(chat_id=chat_id, document=f, filename=os.path.basename(filepath))
        except Exception as e:
            context.bot.send_message(chat_id=chat_id, text=f"⚠️ ساخت فایل Word ناموفق بود: {e}")

        return  # پایان این سناریو

    
    if موضوع == "نقل و انتقال سهام" and نوع_شرکت == "سهامی خاص":
        text = f"""صورتجلسه مجمع عمومی فوق العاده شرکت {data['نام شرکت']} ({نوع_شرکت})  
    شماره ثبت شرکت :     {data['شماره ثبت']}
    شناسه ملی :      {data['شناسه ملی']}
    سرمایه ثبت شده : {data['سرمایه']} ریال

    صورتجلسه مجمع عمومی فوق العاده شرکت {data['نام شرکت']} ({نوع_شرکت}) ثبت شده به شماره {data['شماره ثبت']} در تاریخ  {data['تاریخ']} ساعت {data['ساعت']} با حضور کلیه سهامداران در محل قانونی شرکت تشکیل گردید و تصمیمات ذیل اتخاذ گردید.

    الف: در اجرای ماده 101 لایحه اصلاحی قانون تجارت: 
    ـ  {data['مدیر عامل']}                                   به سمت رئیس جلسه 
    ـ  {data['نایب رییس']}                                  به سمت ناظر 1 جلسه 
    ـ  {data['رییس']}                                        به سمت ناظر 2 جلسه 
    ـ  {data['منشی']}                         به سمت منشی جلسه انتخاب شدند

    ب: دستور جلسه اتخاذ تصمیم در خصوص نقل و انتقال سهام، مجمع موافقت و تصویب نمود که:"""

        foroshandeha_tajmi = defaultdict(list)

        for i in range(1, data["تعداد فروشندگان"] + 1):
            nam = data[f'فروشنده {i} نام']
            kodmeli = data[f'فروشنده {i} کد ملی']
            tedad = data[f'فروشنده {i} تعداد']
            for j in range(1, data.get(f"تعداد خریداران {i}", 0) + 1):
                foroshandeha_tajmi[nam].append({
                    "کد ملی": kodmeli,
                    "تعداد": tedad,
                    "خریدار": data.get(f'خریدار {i}-{j} نام', ''),
                    "کد ملی خریدار": data.get(f'خریدار {i}-{j} کد ملی', ''),
                    "آدرس خریدار": data.get(f'خریدار {i}-{j} آدرس', '')
                })

        for nam_forooshande, vaghzari_ha in foroshandeha_tajmi.items():
            kod_meli_forooshande = vaghzari_ha[0]["کد ملی"]
            matn = f"\n    {nam_forooshande} به شماره ملی {kod_meli_forooshande} "

            jomalat = []
            majmoo_montaghel = 0
            for item in vaghzari_ha:
                tedad = int(fa_to_en_number(item["تعداد"]))
                majmoo_montaghel += tedad
                jomalat.append(
                    f"تعداد {item['تعداد']} سهم به {item['خریدار']} به شماره ملی {item['کد ملی خریدار']} به آدرس {item['آدرس خریدار']}"
                )

            matn += " و همچنین ".join(jomalat)
            matn += " واگذار کرد"

            majmoo_saham_qabl = 0
            for j in range(1, data["تعداد سهامداران قبل"] + 1):
                if data[f"سهامدار قبل {j} نام"] == nam_forooshande:
                    majmoo_saham_qabl = int(fa_to_en_number(data[f"سهامدار قبل {j} تعداد"]))
                    break

            if majmoo_montaghel == majmoo_saham_qabl:
                matn += " و از شرکت خارج شد و دیگر هیچ گونه حق و سمتی ندارد."

            text += matn

        text += f"""

    مجمع به {data['وکیل']} احدی از سهامداران شرکت وکالت داده می شود که ضمن مراجعه به اداره ثبت شرکتها نسبت به ثبت صورتجلسه و پرداخت حق الثبت و امضاء ذیل دفاتر ثبت اقدام نماید. 

    امضاء اعضاء هیات رئیسه: 
    رئیس جلسه :  {data['مدیر عامل']}                                   ناظر1 جلسه : {data['نایب رییس']}                                
    ناظر2جلسه : {data['رییس']}                                       منشی جلسه: {data['منشی']}


    فروشندگان :"""
        for nam_forooshande in foroshandeha_tajmi:
            text += f" {nam_forooshande}     "

        text += "\nخریداران :"
        for vaghzari_ha in foroshandeha_tajmi.values():
            for item in vaghzari_ha:
                text += f" {item['خریدار']}     "
    
    
        # جدول سهامداران قبل
        text += f"\n\nصورت سهامداران حاضر در مجمع عمومی (فوق العاده) مورخه {data['تاریخ']}\n{data['نام شرکت']} قبل از نقل و انتقال سهام\n"
        text += "ردیف\tنام و نام خانوادگی\tتعداد سهام\tامضا سهامداران\n"
        for i in range(1, data["تعداد سهامداران قبل"] + 1):
            text += f"{i}\t{data[f'سهامدار قبل {i} نام']}\t{data[f'سهامدار قبل {i} تعداد']}\t\n"

        # جدول سهامداران بعد
        text += f"\nصورت سهامداران حاضر در مجمع عمومی (فوق العاده) مورخه {data['تاریخ']}\n{data['نام شرکت']} بعد از نقل و انتقال سهام\n"
        text += "ردیف\tنام و نام خانوادگی\tتعداد سهام\tامضا سهامداران\n"
        for i in range(1, data["تعداد سهامداران بعد"] + 1):
            text += f"{i}\t{data[f'سهامدار بعد {i} نام']}\t{data[f'سهامدار بعد {i} تعداد']}\t\n"

        # ارسال متن و فایل Word
        context.bot.send_message(chat_id=chat_id, text=text)

        file_path = generate_word_file(text)
        with open(file_path, 'rb') as f:
            context.bot.send_document(chat_id=chat_id, document=f, filename="صورتجلسه نقل و انتقال.docx")

        os.remove(file_path)
        return

    # کد صورتجلسه تغییر آدرس سهامی خاص
    
    if موضوع == "تغییر آدرس" and نوع_شرکت == "سهامی خاص":
        # فقط در این حالت صورتجلسه سهامی خاص را بفرست
        text = f"""صورتجلسه مجمع عمومی فوق العاده شرکت {data['نام شرکت']} {data['نوع شرکت']}
شماره ثبت شرکت : {data['شماره ثبت']}
شناسه ملی : {data['شناسه ملی']}
سرمایه ثبت شده : {data['سرمایه']} ریال

صورتجلسه مجمع عمومی فوق العاده شرکت {data['نام شرکت']} {data['نوع شرکت']} ثبت شده به شماره {data['شماره ثبت']} در تاریخ {data['تاریخ']} ساعت {data['ساعت']} با حضور کلیه سهامداران در محل قانونی شرکت تشکیل گردید و تصمیمات ذیل اتخاذ گردید.

الف: در اجرای ماده 101 لایحه اصلاحی قانون تجارت: 
ـ  {data['مدیر عامل']} به سمت رئیس جلسه 
ـ  {data['نایب رییس']} به سمت ناظر 1 جلسه 
ـ  {data['رییس']} به سمت ناظر 2 جلسه 
ـ  {data['منشی']} به سمت منشی جلسه انتخاب شدند

ب: دستور جلسه اتخاذ تصمیم در خصوص تغییر محل شرکت، مجمع موافقت و تصویب نمود که:
محل شرکت از آدرس قبلی به آدرس جدید {data['آدرس جدید']} کد پستی {data['کد پستی']} انتقال یافت.

مجمع به {data['وکیل']} احدی از سهامداران شرکت وکالت داده می شود که ضمن مراجعه به اداره ثبت شرکتها نسبت به ثبت صورتجلسه و پرداخت حق الثبت و امضاء ذیل دفاتر ثبت اقدام نماید.

امضاء اعضاء هیات رئیسه: 
رئیس جلسه : {data['مدیر عامل']}     ناظر1 جلسه : {data['نایب رییس']}     
ناظر2 جلسه : {data['رییس']}         منشی جلسه: {data['منشی']}"""
        context.bot.send_message(chat_id=chat_id, text=text)

        # ✅ ساخت فایل Word و ارسال
        file_path = generate_word_file(text)
        with open(file_path, 'rb') as f:
            context.bot.send_document(chat_id=chat_id, document=f, filename="صورتجلسه.docx")
    
        os.remove(file_path)  # ← حذف فایل پس از ارسال (اختیاری)
        return

    if موضوع == "تغییر موضوع فعالیت" and نوع_شرکت == "مسئولیت محدود":
        count = data.get("تعداد شرکا", 0)
        partners_lines = ""
        for i in range(1, count + 1):
            name = data.get(f"شریک {i}", "")
            share = data.get(f"سهم الشرکه شریک {i}", "")
            partners_lines += f"{name}                                              {share} ریال\n"

        action_line = (
            "نسبت به الحاق مواردی به موضوع شرکت اتخاذ تصمیم شد."
            if data["نوع تغییر موضوع"] == "الحاق"
            else "نسبت به تغییر موضوع شرکت اتخاذ تصمیم شد."
        )
        subject_line = (
            "مواردی به شرح ذیل به موضوع شرکت الحاق شد:"
            if data["نوع تغییر موضوع"] == "الحاق"
            else "موضوع شرکت به شرح ذیل تغییر یافت:"
        )

        text = f"""صورتجلسه مجمع عمومی فوق العاده شرکت {data['نام شرکت']} ({نوع_شرکت})
شماره ثبت شرکت :     {data['شماره ثبت']}
شناسه ملی :      {data['شناسه ملی']}
سرمایه ثبت شده : {data['سرمایه']} ریال

صورتجلسه مجمع عمومی فوق العاده شرکت {data['نام شرکت']} ({نوع_شرکت}) ثبت شده به شماره {data['شماره ثبت']} در تاریخ  {data['تاریخ']} ساعت {data['ساعت']} با حضور کلیه شرکا در محل قانونی شرکت تشکیل و {action_line}

اسامی شرکا                                                        میزان سهم الشرکه
{partners_lines}
{subject_line}
{data['موضوع جدید']} 
و ماده مربوطه اساسنامه به شرح فوق اصلاح می گردد. 
به {data['وکیل']} از شرکاء شرکت وکالت داده می شود که ضمن مراجعه به اداره ثبت شرکت ها نسبت به ثبت صورتجلسه و پرداخت حق الثبت و امضاء ذیل دفاتر ثبت اقدام نماید.

امضاء شرکاء: 
"""

        for i in range(1, count + 1):
            text += f"{data.get(f'شریک {i}', '')}     "
        context.bot.send_message(chat_id=chat_id, text=text)

        # فایل Word
        file_path = generate_word_file(text)
        with open(file_path, 'rb') as f:
            context.bot.send_document(chat_id=chat_id, document=f, filename="صورتجلسه تغییر موضوع فعالیت.docx")
        os.remove(file_path)
        return

    # -------------------------------
    # خروجی: تغییر نام شرکت - سهامی خاص
    # -------------------------------
    if موضوع == "تغییر نام شرکت" and نوع_شرکت == "سهامی خاص":
        text = f"""صورتجلسه مجمع عمومی فوق العاده شرکت {data['نام شرکت']} ({نوع_شرکت})
    شماره ثبت شرکت :     {data['شماره ثبت']}
    شناسه ملی :     {data['شناسه ملی']}
    سرمایه ثبت شده : {data['سرمایه']} ریال
    
    صورتجلسه مجمع عمومی فوق العاده شرکت {data['نام شرکت']} ({نوع_شرکت}) ثبت شده به شماره {data['شماره ثبت']} در تاریخ  {data['تاریخ']} ساعت {data['ساعت']} با حضور کلیه سهامداران در محل قانونی شرکت تشکیل و نسبت به تغییر نام شرکت اتخاذ تصمیم شد: 
    الف: در اجرای ماده 101 لایحه اصلاحی قانون تجارت: 
    
    ـ  {data['مدیر عامل']}                                   به سمت رئیس جلسه 
    ـ  {data['نایب رییس']}                                  به سمت ناظر 1 جلسه 
    ـ  {data['رییس']}                                        به سمت ناظر 2 جلسه 
    ـ  {data['منشی']}                                        به سمت منشی جلسه انتخاب شدند
    
    ب: پس از شور و بررسی مقرر گردید نام شرکت از {data['نام شرکت']} به {data['نام جدید شرکت']} تغییر یابد در نتیجه ماده مربوطه اساسنامه بشرح مذکور اصلاح می گردد.
    
    ج: مجمع به {data['وکیل']} احدی از سهامداران یا وکیل رسمی شرکت وکالت داده می شود که ضمن مراجعه به اداره ثبت شرکت ها نسبت به ثبت صورتجلسه و پرداخت حق الثبت و امضاء ذیل دفاتر ثبت اقدام نماید.
    
    امضاء اعضاء هیات رئیسه: 
    رئیس جلسه :  {data['مدیر عامل']}                                   ناظر1 جلسه : {data['نایب رییس']}                               
    
    
    ناظر2جلسه : {data['رییس']}                                       منشی جلسه: {data['منشی']}
    """
    
        context.bot.send_message(chat_id=chat_id, text=text)
        file_path = generate_word_file(text)
        with open(file_path, 'rb') as f:
            context.bot.send_document(chat_id=chat_id, document=f, filename="صورتجلسه تغییر نام شرکت سهامی خاص.docx")
        os.remove(file_path)
        return

    # -------------------------------
    # خروجی: تغییر نام شرکت - مسئولیت محدود
    # -------------------------------
    if موضوع == "تغییر نام شرکت" and نوع_شرکت == "مسئولیت محدود":
        count = data.get("تعداد شرکا", 0)
    
        # جدول شرکا
        partners_lines = ""
        for i in range(1, count + 1):
            nm = data.get(f"شریک {i}", "")
            sh = data.get(f"سهم الشرکه شریک {i}", "")
            partners_lines += f"{nm}                                              {sh} ریال\n"
    
        # امضاها: هر دو نام در یک خط بعدی خط جدید
        signer_lines = ""
        for i in range(1, count + 1):
            signer_lines += data.get(f"شریک {i}", "")
            if i % 2 == 1 and i != count:
                signer_lines += "\t"
            else:
                signer_lines += "\n"
    
        text = f"""صورتجلسه مجمع عمومی فوق العاده شرکت {data['نام شرکت']} ({نوع_شرکت})
    شماره ثبت شرکت :     {data['شماره ثبت']}
    شناسه ملی :     {data['شناسه ملی']}
    سرمایه ثبت شده : {data['سرمایه']} ریال
    
    صورتجلسه مجمع عمومی فوق العاده شرکت {data['نام شرکت']} ({نوع_شرکت}) ثبت شده به شماره {data['شماره ثبت']} در تاریخ  {data['تاریخ']} ساعت {data['ساعت']} با حضور کلیه شرکا در محل قانونی شرکت تشکیل و نسبت به تغییر نام شرکت اتخاذ تصمیم شد: 
    
    اسامی شرکا                                                        میزان سهم الشرکه
    {partners_lines}
    پس از شور و بررسی مقرر گردید نام شرکت از {data['نام شرکت']} به {data['نام جدید شرکت']} تغییر یابد در نتیجه ماده مربوطه اساسنامه بشرح مذکور اصلاح می گردد.
    
    به {data['وکیل']} احدی از شرکاء یا وکیل رسمی شرکت وکالت داده می شود که ضمن مراجعه به اداره ثبت شرکت ها نسبت به ثبت صورتجلسه و پرداخت حق الثبت و امضاء ذیل دفاتر ثبت اقدام نماید.
    
    امضاء شرکاء: 
    
    {signer_lines}"""
    
        context.bot.send_message(chat_id=chat_id, text=text)
        file_path = generate_word_file(text)
        with open(file_path, 'rb') as f:
            context.bot.send_document(chat_id=chat_id, document=f, filename="صورتجلسه تغییر نام شرکت مسئولیت محدود.docx")
        send_thank_you_message_chatid(chat_id, context)
        os.remove(file_path)
        return

    # -------------------------------
    # خروجی: انحلال شرکت - مسئولیت محدود
    # -------------------------------
    if موضوع == "انحلال شرکت" and نوع_شرکت == "مسئولیت محدود":
        # ساخت لیست شرکا
        partners_lines = ""
        count = data.get("تعداد شرکا", 0)
        for i in range(1, count + 1):
            name = data.get(f"شریک {i}", "")
            share = data.get(f"سهم الشرکه شریک {i}", "")
            partners_lines += f"{name}                                              {share} ریال\n"

        # امضاها: هر دو نام در یک خط، بعدی خط بعد (برای خوانایی)
        signer_lines = ""
        for i in range(1, count + 1):
            signer_lines += data.get(f"شریک {i}", "")
            if i % 2 == 1 and i != count:
                signer_lines += "\t"
            else:
                signer_lines += "\n"

        text = f"""صورتجلسه انحلال شرکت {data['نام شرکت']} ({نوع_شرکت})
شماره ثبت شرکت :     {data['شماره ثبت']}
شناسه ملی :      {data['شناسه ملی']}
سرمایه ثبت شده : {data['سرمایه']} ریال

صورتجلسه مجمع عمومی فوق العاده شرکت {data['نام شرکت']} ({نوع_شرکت}) ثبت شده به شماره {data['شماره ثبت']} در تاریخ  {data['تاریخ']} ساعت {data['ساعت']} با حضور کلیه شرکا در محل قانونی شرکت تشکیل و تصمیمات ذیل اتخاذ گردید.

اسامی شرکا                                                        میزان سهم الشرکه
{partners_lines}
دستور جلسه، اتخاذ تصمیم در خصوص انحلال شرکت {data['نام شرکت']} ){نوع_شرکت}( پس از بحث و بررسی شرکت بعلت {data['علت انحلال']} منحل گردید و آقای {data['نام مدیر تصفیه']} به شماره ملی {data['کد ملی مدیر تصفیه']} به سمت مدیر تصفیه برای مدت {data['مدت مدیر تصفیه']} سال انتخاب شد. آدرس مدیر تصفیه و محل تصفیه {data['آدرس مدیر تصفیه']} می باشد.
مدیر تصفیه اقرار به دریافت کلیه اموال دارایی ها و دفاتر و اوراق و اسناد مربوط به شرکت را نمود.

به {data['وکیل']} از شرکاء یا وکیل رسمی شرکت وکالت داده می شود که ضمن مراجعه به اداره ثبت شرکت ها نسبت به ثبت صورتجلسه و پرداخت حق الثبت و امضاء ذیل دفاتر ثبت اقدام نماید.

امضاء شرکاء: 

{signer_lines}"""

        # ارسال متن و فایل Word
        context.bot.send_message(chat_id=chat_id, text=text)
        file_path = generate_word_file(text)
        with open(file_path, 'rb') as f:
            context.bot.send_document(chat_id=chat_id, document=f, filename="صورتجلسه انحلال مسئولیت محدود.docx")
        send_thank_you_message_chatid(chat_id, context)
        os.remove(file_path)
        return

    # -------------------------------
    # خروجی: نقل و انتقال سهم الشرکه - مسئولیت محدود
    # -------------------------------
    if موضوع == "نقل و انتقال سهام" and نوع_شرکت == "مسئولیت محدود":
        # جدول شرکا (بالای متن)
        partners_lines = ""
        count = data.get("تعداد شرکا", 0)
        for i in range(1, count + 1):
            name = data.get(f"شریک {i}", "")
            share = data.get(f"سهم الشرکه شریک {i}", "")
            partners_lines += f"{name}                                              {share} ریال\n"

        text = f"""صورتجلسه مجمع عمومی فوق العاده شرکت {data['نام شرکت']} ({نوع_شرکت})
شماره ثبت شرکت :     {data['شماره ثبت']}
شناسه ملی :      {data['شناسه ملی']}
سرمایه ثبت شده : {data['سرمایه']} ریال

صورتجلسه مجمع عمومی فوق العاده شرکت {data['نام شرکت']} ({نوع_شرکت}) ثبت شده به شماره {data['شماره ثبت']} در تاریخ  {data['تاریخ']} ساعت {data['ساعت']} با حضور کلیه شرکا در محل قانونی شرکت تشکیل و نسبت به نقل و انتقال سهم الشرکه بشرح ذیل اتخاذ تصمیم شد:

اسامی شرکا                                                        میزان سهم الشرکه
{partners_lines}
"""

        # پاراگراف‌های واگذاری برای هر فروشنده
        for i in range(1, data.get("تعداد فروشندگان", 0) + 1):
            seller_name = data.get(f"فروشنده {i} نام", "")
            seller_nid = data.get(f"فروشنده {i} کد ملی", "")
            seller_total = data.get(f"فروشنده {i} سهم کل", "")
            senad_no = data.get(f"فروشنده {i} سند صلح", "")
            senad_date = data.get(f"فروشنده {i} تاریخ سند", "")
            daftar_no = data.get(f"فروشنده {i} دفترخانه", "")

            sentence = (
                f"پس از مذاکره مقرر شد که {seller_name} به شماره ملی {seller_nid} "
                f"که دارای {seller_total} ریال سهم الشرکه می باشد "
                f"با رعایت مفاد ماده 103 قانون تجارت و بموجب سند صلح به شماره {senad_no} "
                f"مورخ {senad_date} صادره از دفتراسناد رسمی {daftar_no} "
            )

            # خریداران مرتبط با این فروشنده
            total_transferred = 0
            buyers_cnt = data.get(f"تعداد خریداران {i}", 0)
            first = True
            for k in range(1, buyers_cnt + 1):
                b_name = data.get(f"خریدار {i}-{k} نام", "")
                b_father = data.get(f"خریدار {i}-{k} پدر", "")
                b_birth = data.get(f"خریدار {i}-{k} تولد", "")
                b_nid = data.get(f"خریدار {i}-{k} کد ملی", "")
                b_addr = data.get(f"خریدار {i}-{k} آدرس", "")
                b_share = data.get(f"خریدار {i}-{k} سهم منتقل", "")

                # جمع کل منتقل‌شده برای تعیین خروج/عدم‌خروج فروشنده
                try:
                    total_transferred += int(fa_to_en_number(b_share))
                except Exception:
                    pass

                prefix = "معادل" if first else "و همچنین معادل"
                sentence += (
                    f"{prefix} {b_share} ریال سهم الشرکه خود را به {b_name} "
                    f"فرزند {b_father} متولد {b_birth} "
                    f"به شماره ملی {b_nid} آدرس محل سکونت {b_addr} منتقل "
                )
                first = False

            # اگر به اندازه کل سهم‌الشرکه‌اش منتقل کرده باشد → خروج از شرکت
            try:
                seller_total_int = int(fa_to_en_number(seller_total))
            except Exception:
                seller_total_int = None

            if seller_total_int is not None and seller_total_int == total_transferred:
                sentence += "و از شرکت خارج  شد و دیگر هیچ گونه حق و سمتی در شرکت ندارد."
            else:
                sentence += "نمود."
                
            text += sentence + "\n"

        text += "\nاین نقل و انتقال سهم الشرکه مورد موافقت کلیه شرکاء با رعایت مفاد ماده 102 قانون تجارت قرار گرفت.\n\n"
        text += f"به {data['وکیل']} احدی از شرکاء یا وکیل رسمی شرکت وکالت داده شد که ضمن مراجعه به اداره ثبت شرکتها نسبت به ثبت صورتجلسه و پرداخت حق الثبت و امضاء ذیل دفتر ثبت اقدام نماید. \n\n"

        # جدول امضاء پایانی
        text += "    نام شرکاء                                        میزان سهم الشرکه                                     امضاء\n"
        for i in range(1, count + 1):
            text += f" {data.get(f'شریک {i}', '')}                                   {data.get(f'سهم الشرکه شریک {i}', '')} ریال\n"

        # ارسال متن و فایل Word
        context.bot.send_message(chat_id=chat_id, text=text)
        file_path = generate_word_file(text)
        with open(file_path, 'rb') as f:
            context.bot.send_document(chat_id=chat_id, document=f, filename="صورتجلسه نقل و انتقال سهم‌الشرکه مسئولیت محدود.docx")
        os.remove(file_path)
        return

    
    # -------------------------------
    # خروجی: انحلال شرکت - سهامی خاص
    # -------------------------------
    if موضوع == "انحلال شرکت" and نوع_شرکت == "سهامی خاص":
        # ساخت جدول سهامداران حاضر
        count = data.get("تعداد سهامداران حاضر", 0)
        rows = ""
        for i in range(1, count + 1):
            rows += f"{i}\n\t{data.get(f'سهامدار {i} نام','')}\t{data.get(f'سهامدار {i} تعداد','')}\t\n"

        # متن اصلی مطابق قالب شما (با اصلاح برچسب‌های متنیِ منطقی)
        text = f"""صورتجلسه انحلال شرکت {data['نام شرکت']} ){نوع_شرکت}(
شماره ثبت شرکت :     {data['شماره ثبت']}
شناسه ملی :      {data['شناسه ملی']}
سرمایه ثبت شده : {data['سرمایه']} ریال

صورتجلسه مجمع عمومی فوق العاده شرکت {data['نام شرکت']} ){نوع_شرکت}( ثبت شده به شماره {data['شماره ثبت']} در تاریخ  {data['تاریخ']} ساعت {data['ساعت']} با حضور کلیه سهامداران در محل قانونی شرکت تشکیل گردید و تصمیمات ذیل اتخاذ گردید.
الف: در اجرای ماده 101 لایحه اصلاحی قانون تجارت: 

ـ  {data['مدیر عامل']}                                   به سمت رئیس جلسه 
ـ  {data['نایب رییس']}                                  به سمت ناظر 1 جلسه 
ـ  {data['رییس']}                                        به سمت ناظر 2 جلسه 
ـ  {data['منشی']}                                       به سمت منشی جلسه انتخاب شدند

ب: دستور جلسه، اتخاذ تصمیم در خصوص انحلال شرکت {data['نام شرکت']} ){نوع_شرکت}( پس از بحث و بررسی شرکت بعلت {data['علت انحلال']} منحل گردید و  {data['نام مدیر تصفیه']} به شماره ملی {data['کد ملی مدیر تصفیه']} به سمت مدیر تصفیه برای مدت {data['مدت مدیر تصفیه']} سال انتخاب شد. آدرس مدیر تصفیه و محل تصفیه {data['آدرس مدیر تصفیه']} می باشد.
مدیر تصفیه اقرار به دریافت کلیه اموال دارایی ها و دفاتر و اوراق و اسناد مربوط به شرکت را نمود.

ج: مجمع به {data['وکیل']} از سهامداران یا وکیل رسمی شرکت وکالت داده می شود که ضمن مراجعه به اداره ثبت شرکتها نسبت به ثبت صورتجلسه و پرداخت حق الثبت و امضاء ذیل دفاتر ثبت اقدام نماید. 
امضاء اعضاء هیات رئیسه: 

رئیس جلسه :  {data['مدیر عامل']}                                   ناظر1 جلسه : {data['نایب رییس']}                               


ناظر2جلسه : {data['رییس']}                                       منشی جلسه: {data['منشی']}





صورت سهامداران حاضر در مجمع عمومی (فوق العاده) مورخه {data['تاریخ']}
{data['نام شرکت']}
ردیف\tنام و نام خانوادگی\tتعداد سهام\tامضا سهامداران
{rows}"""

        # ارسال متن
        context.bot.send_message(chat_id=chat_id, text=text)

        # فایل Word
        file_path = generate_word_file(text)
        with open(file_path, 'rb') as f:
            context.bot.send_document(chat_id=chat_id, document=f, filename="صورتجلسه انحلال.docx")
        os.remove(file_path)
        return

    else:
        # اگر هیچ‌کدام از حالت‌های بالا نبود:
        context.bot.send_message(chat_id=chat_id, text="✅ اطلاعات با موفقیت دریافت شد.\nدر حال حاضر صورتجلسه‌ای برای این ترکیب تعریف نشده است.")

@app.route('/webhook', methods=['POST'])
def webhook():
    update = telegram.Update.de_json(request.get_json(force=True), bot)
    dispatcher.process_update(update)
    return 'ok'
# updater = Updater(...)  # disabled for webhook mode

dispatcher = Dispatcher(bot, None, workers=4, use_context=True)
dispatcher.add_handler(CallbackQueryHandler(handle_newspaper_choice, pattern=r"^newspaper:"))
dispatcher.add_handler(CallbackQueryHandler(handle_inline_callbacks), group=0)

# ===== گروه 0: مربوط به AI =====
dispatcher.add_handler(MessageHandler(Filters.text & Filters.regex(f"^{re.escape(AI_ASK_TEXT)}$"), enter_ai_mode_reply), group=0)

# دکمه‌ی اینلاین «بازگشت از AI»
dispatcher.add_handler(CallbackQueryHandler(resume_from_ai, pattern=f"^{AI_RESUME}$"), group=0)

# ===== گروه 1: هندلرهای عمومی =====
dispatcher.add_handler(MessageHandler(Filters.contact, handle_contact), group=1)
dispatcher.add_handler(CommandHandler("ai", cmd_ai), group=1)
dispatcher.add_handler(CommandHandler("start", start), group=1)
dispatcher.add_handler(CallbackQueryHandler(button_handler, pattern=fr"^(?!{AI_RESUME}$).+"),group=1)
dispatcher.add_handler(MessageHandler(Filters.text & ~Filters.command, handle_message), group=1)



def remember_last_question(context, label: str):
    """
    ذخیره‌ی آخرین برچسب سؤال برای اینکه در حالت AI
    بتوانیم بعد از خروج دوباره به همان مرحله برگردیم.
    """
    try:
        context.user_data["last_question_text"] = label
    except Exception as e:
        print("remember_last_question error:", e)

if __name__ == "__main__":
    import os
    port = int(os.environ.get("PORT", 5000))
    app.run(host="0.0.0.0", port=port)

